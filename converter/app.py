"""
GovAI 文档转换微服务
===================
基于 LibreOffice 的独立文档转换服务，提供：
  - POST /convert-to-pdf   — 将上传文件转为 PDF
  - POST /extract-text     — 从上传文件提取纯文本
  - GET  /health           — 健康检查
"""

import asyncio
import logging
import os
import shutil
import tempfile
import uuid
from pathlib import Path

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import FileResponse, JSONResponse

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("converter")

app = FastAPI(title="GovAI Converter Service", version="1.0.0")

# 共享卷挂载点 — 用于后端直接取文件，也可通过 HTTP 返回
SHARED_DIR = Path(os.getenv("SHARED_DIR", "/shared"))
SHARED_DIR.mkdir(parents=True, exist_ok=True)

# LibreOffice 用户配置目录（避免多进程冲突）
LO_USER_DIR = Path("/tmp/lo_user")
LO_USER_DIR.mkdir(parents=True, exist_ok=True)

# 支持的文件后缀
CONVERTIBLE_EXTENSIONS = {
    "docx", "doc", "xlsx", "xls", "pptx", "ppt",
    "odt", "ods", "odp", "rtf", "html", "htm", "txt", "csv",
}


async def _run_libreoffice(input_path: Path, output_dir: Path, convert_to: str = "pdf") -> Path:
    """调用 LibreOffice headless 转换文件"""
    user_profile = LO_USER_DIR / str(uuid.uuid4())
    user_profile.mkdir(parents=True, exist_ok=True)

    cmd = [
        "libreoffice",
        "--headless",
        "--norestore",
        "--nofirststartwizard",
        f"-env:UserInstallation=file://{user_profile}",
        "--convert-to", convert_to,
        "--outdir", str(output_dir),
        str(input_path),
    ]

    logger.info(f"执行 LibreOffice: {' '.join(cmd)}")

    proc = await asyncio.create_subprocess_exec(
        *cmd,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    stdout, stderr = await proc.communicate()

    # 清理临时 profile
    shutil.rmtree(user_profile, ignore_errors=True)

    if proc.returncode != 0:
        err_msg = stderr.decode(errors="replace")
        logger.error(f"LibreOffice 转换失败 (rc={proc.returncode}): {err_msg}")
        raise RuntimeError(f"LibreOffice conversion failed: {err_msg}")

    # 查找输出文件
    stem = input_path.stem
    candidates = list(output_dir.glob(f"{stem}.*"))
    result_files = [f for f in candidates if f.suffix.lower() == f".{convert_to}"]

    if not result_files:
        # 尝试匹配任何新文件
        all_files = list(output_dir.iterdir())
        logger.warning(f"未找到 .{convert_to} 输出, 目录中有: {[f.name for f in all_files]}")
        if all_files:
            return all_files[0]
        raise RuntimeError(f"转换后未找到 .{convert_to} 文件")

    return result_files[0]


async def _extract_text_pdf(file_path: Path) -> str:
    """使用 pdfplumber 提取 PDF 文本（比 LibreOffice 精准得多）"""
    try:
        import pdfplumber
        pages_text: list[str] = []
        with pdfplumber.open(str(file_path)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages_text.append(text)
                # 也尝试提取表格
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            row_text = "\t".join(
                                cell.strip() if cell else "" for cell in row
                            )
                            if row_text.strip():
                                pages_text.append(row_text)
        result = "\n\n".join(pages_text)
        if result.strip():
            return result
        logger.warning("pdfplumber 未提取到任何文本，降级到 LibreOffice")
        return ""
    except ImportError:
        logger.warning("pdfplumber 未安装，降级到 LibreOffice")
        return ""
    except Exception as e:
        logger.warning(f"pdfplumber 提取失败: {e}，降级到 LibreOffice")
        return ""


async def _extract_text_docx(file_path: Path) -> str:
    """使用 python-docx 提取 DOCX 文本（比 LibreOffice 更精准）"""
    try:
        from docx import Document
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except Exception as e:
        logger.warning(f"python-docx 提取失败, 降级到 LibreOffice: {e}")
        return ""


def _read_text_with_encoding_detection(file_path: Path) -> str:
    """多编码探测读取文本文件"""
    for encoding in ("utf-8", "utf-8-sig", "gbk", "gb2312", "gb18030", "latin-1"):
        try:
            return file_path.read_text(encoding=encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return file_path.read_text(encoding="utf-8", errors="replace")


async def _extract_text_libreoffice(file_path: Path) -> str:
    """通过 LibreOffice 转为 txt 来提取文本"""
    with tempfile.TemporaryDirectory(prefix="lo_txt_") as tmp_dir:
        output_dir = Path(tmp_dir)
        try:
            txt_file = await _run_libreoffice(file_path, output_dir, convert_to="txt")
            return _read_text_with_encoding_detection(txt_file)
        except Exception as e:
            logger.error(f"LibreOffice 文本提取失败: {e}")
            # 最后兜底: 尝试直接读取
            try:
                return _read_text_with_encoding_detection(file_path)
            except Exception:
                return ""


def _save_upload_temp(upload: UploadFile) -> tuple[Path, str]:
    """将上传文件保存到临时目录，返回 (路径, 扩展名)"""
    original_name = upload.filename or "unknown.bin"
    ext = Path(original_name).suffix.lstrip(".").lower()
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}" if ext else "")
    content = upload.file.read()
    tmp.write(content)
    tmp.close()
    return Path(tmp.name), ext


# ── API 路由 ──

@app.get("/health")
async def health():
    """健康检查"""
    # 验证 LibreOffice 是否可用
    proc = await asyncio.create_subprocess_exec(
        "libreoffice", "--headless", "--version",
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    stdout, _ = await proc.communicate()
    lo_version = stdout.decode(errors="replace").strip() if proc.returncode == 0 else "unavailable"
    return {"status": "ok", "libreoffice": lo_version}


@app.post("/convert-to-pdf")
async def convert_to_pdf(file: UploadFile = File(...)):
    """
    将上传文件转换为 PDF。
    
    - 如果上传的已经是 PDF，直接返回
    - 否则使用 LibreOffice headless 转换
    
    返回:
      - PDF 文件流 (application/pdf)
      - Header X-PDF-Filename: 输出文件名
    """
    tmp_path, ext = _save_upload_temp(file)

    try:
        if ext == "pdf":
            # 已经是 PDF，直接返回
            stem = Path(file.filename or 'document').stem
            from urllib.parse import quote
            safe_name = quote(f"{stem}.pdf")
            return FileResponse(
                str(tmp_path),
                media_type="application/pdf",
                filename="document.pdf",
                headers={
                    "X-PDF-Filename": safe_name,
                    "Content-Disposition": f"attachment; filename=\"document.pdf\"; filename*=UTF-8''{safe_name}",
                },
            )

        if ext not in CONVERTIBLE_EXTENSIONS and ext != "pdf":
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件格式: .{ext}。支持: {', '.join(sorted(CONVERTIBLE_EXTENSIONS))}",
            )

        # LibreOffice 转换
        output_dir = Path(tempfile.mkdtemp(prefix="lo_pdf_"))
        try:
            pdf_path = await _run_libreoffice(tmp_path, output_dir, convert_to="pdf")

            # 复制到共享目录（可选，供后端直接读取）
            shared_name = f"{uuid.uuid4().hex}.pdf"
            shared_path = SHARED_DIR / shared_name
            shutil.copy2(pdf_path, shared_path)

            stem = Path(file.filename or 'document').stem
            encoded = quote(f"{stem}.pdf")
            return FileResponse(
                str(shared_path),
                media_type="application/pdf",
                filename="document.pdf",
                headers={
                    "X-PDF-Filename": encoded,
                    "X-Shared-Path": str(shared_path),
                },
            )
        finally:
            shutil.rmtree(output_dir, ignore_errors=True)
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("PDF 转换失败")
        raise HTTPException(status_code=500, detail=f"PDF 转换失败: {str(e)}")
    finally:
        # 延迟清理原始临时文件（FileResponse 需要文件存在）
        # 通过 background task 清理
        pass


@app.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    """
    从上传文件提取纯文本。

    策略:
      1. DOCX → 优先用 python-docx 精准提取
      2. 其他格式 → LibreOffice 转 txt
      3. 兜底 → 直接读取

    返回:
      { "text": "...", "char_count": 1234, "method": "python-docx|libreoffice|raw" }
    """
    tmp_path, ext = _save_upload_temp(file)

    try:
        text = ""
        method = "unknown"

        if ext == "pdf":
            text = await _extract_text_pdf(tmp_path)
            method = "pdfplumber"

        if not text and ext == "docx":
            text = await _extract_text_docx(tmp_path)
            method = "python-docx"

        if not text and ext in {"txt", "md", "csv"}:
            text = tmp_path.read_text(encoding="utf-8", errors="replace")
            method = "raw"

        if not text:
            text = await _extract_text_libreoffice(tmp_path)
            method = "libreoffice"

        if not text:
            # 最终兜底
            try:
                text = tmp_path.read_text(encoding="utf-8", errors="replace")
                method = "raw-fallback"
            except Exception:
                text = ""
                method = "failed"

        return JSONResponse({
            "text": text,
            "char_count": len(text),
            "method": method,
        })
    except Exception as e:
        logger.exception("文本提取失败")
        raise HTTPException(status_code=500, detail=f"文本提取失败: {str(e)}")
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


@app.post("/convert-and-extract")
async def convert_and_extract(file: UploadFile = File(...)):
    """
    一次性完成 PDF 转换 + 文本提取（减少上传次数）。

    返回:
      { "pdf_path": "/shared/xxx.pdf", "text": "...", "char_count": 1234 }
    """
    tmp_path, ext = _save_upload_temp(file)

    try:
        # 1. 提取文本
        text = ""
        if ext == "pdf":
            text = await _extract_text_pdf(tmp_path)
        if not text and ext == "docx":
            text = await _extract_text_docx(tmp_path)
        if not text and ext in {"txt", "md", "csv"}:
            text = tmp_path.read_text(encoding="utf-8", errors="replace")
        if not text:
            text = await _extract_text_libreoffice(tmp_path)
        if not text:
            try:
                text = tmp_path.read_text(encoding="utf-8", errors="replace")
            except Exception:
                text = ""

        # 2. 转为 PDF
        pdf_shared_path = ""
        if ext == "pdf":
            shared_name = f"{uuid.uuid4().hex}.pdf"
            shared_path = SHARED_DIR / shared_name
            shutil.copy2(tmp_path, shared_path)
            pdf_shared_path = str(shared_path)
        elif ext in CONVERTIBLE_EXTENSIONS:
            output_dir = Path(tempfile.mkdtemp(prefix="lo_both_"))
            try:
                pdf_file = await _run_libreoffice(tmp_path, output_dir, convert_to="pdf")
                shared_name = f"{uuid.uuid4().hex}.pdf"
                shared_path = SHARED_DIR / shared_name
                shutil.copy2(pdf_file, shared_path)
                pdf_shared_path = str(shared_path)
            finally:
                shutil.rmtree(output_dir, ignore_errors=True)
        else:
            # 不可转 PDF 的格式（如纯文本），不生成 PDF
            pdf_shared_path = ""

        return JSONResponse({
            "pdf_path": pdf_shared_path,
            "text": text,
            "char_count": len(text),
        })
    except Exception as e:
        logger.exception("转换+提取失败")
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
