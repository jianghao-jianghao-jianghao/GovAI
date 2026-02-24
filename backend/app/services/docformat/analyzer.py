"""格式诊断模块 — 适配自 docformat-gui (MIT License)"""

import re
import logging
from collections import defaultdict
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger('docformat.analyzer')

NO_INDENT_PATTERNS = [
    r'^附件[：:]',
    r'^联系人[：:]',
    r'^抄送[：:]',
    r'^主送[：:]',
]


def is_no_indent_para(text, alignment):
    """检查是否不需要首行缩进的段落"""
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return True
    for pattern in NO_INDENT_PATTERNS:
        if re.match(pattern, text.strip()):
            return True
    return False


def analyze_punctuation(doc) -> list:
    """分析标点符号问题"""
    issues = []
    patterns = [
        ('英文括号', r'[\(\)]'),
        ('英文引号', r'["\']'),
        ('英文冒号', r'(?<=[^\d\s]):(?=[^\d/\\])'),
        ('英文逗号', r'(?<=[^\d]),(?=[^\d])'),
        ('英文分号', r';'),
        ('英文问号', r'\?'),
        ('英文叹号', r'!'),
    ]
    ellipsis_pattern = r'\.{2,}'
    dash_pattern = r'--+'
    period_pattern = r'(?<=[\u4e00-\u9fff])\.(?!\.)'

    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        if not re.search(r'[\u4e00-\u9fff]', text):
            continue
        for name, pattern in patterns:
            for match in re.finditer(pattern, text):
                issues.append({'para': i + 1, 'type': name, 'char': match.group()})
        for match in re.finditer(ellipsis_pattern, text):
            issues.append({'para': i + 1, 'type': '不规范省略号', 'char': match.group()})
        for match in re.finditer(dash_pattern, text):
            issues.append({'para': i + 1, 'type': '不规范破折号', 'char': match.group()})
        for match in re.finditer(period_pattern, text):
            issues.append({'para': i + 1, 'type': '英文句号', 'char': match.group()})

    return issues


def analyze_numbering(doc) -> list:
    """分析序号问题"""
    issues = []
    numbering_patterns = {
        'chinese_1': r'^[一二三四五六七八九十]+、',
        'chinese_2': r'^（[一二三四五六七八九十]+）',
        'arabic_dot': r'^\d+\.',
        'arabic_comma': r'^\d+、',
        'arabic_paren': r'^\d+[）\)]',
        'arabic_paren_full': r'^（\d+）',
    }
    found_styles = defaultdict(list)
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for style_name, pattern in numbering_patterns.items():
            if re.match(pattern, text):
                found_styles[style_name].append(i + 1)
                break
    arabic_styles = [k for k in found_styles if k.startswith('arabic')]
    if len(arabic_styles) > 1:
        issues.append({
            'type': '序号格式不统一',
            'detail': f"同时存在: {', '.join(arabic_styles)}",
        })
    return issues


def analyze_paragraph_format(doc) -> list:
    """分析段落格式问题"""
    issues = []
    indent_issues = []
    line_spacing_values = defaultdict(list)

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text or len(text) < 10:
            continue
        alignment = para.paragraph_format.alignment
        if is_no_indent_para(text, alignment):
            continue
        pf = para.paragraph_format
        indent = pf.first_line_indent
        if indent is None or indent == Pt(0) or (hasattr(indent, 'pt') and indent.pt == 0):
            indent_issues.append(i + 1)
        if pf.line_spacing is not None:
            line_spacing_values[str(pf.line_spacing)].append(i + 1)

    if indent_issues:
        issues.append({'type': '缺少首行缩进', 'paras': indent_issues})
    if len(line_spacing_values) > 1:
        issues.append({
            'type': '行距不统一',
            'detail': f"存在 {len(line_spacing_values)} 种不同行距",
        })
    return issues


def analyze_font(doc) -> list:
    """分析字体问题"""
    issues = []
    font_names = set()
    font_sizes = set()
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        for run in para.runs:
            if run.font.name:
                font_names.add(run.font.name)
            if run.font.size:
                font_sizes.add(str(run.font.size))
    if len(font_names) > 4:
        issues.append({
            'type': '字体种类过多',
            'detail': f"检测到 {len(font_names)} 种字体: {', '.join(list(font_names)[:5])}..."
        })
    if len(font_sizes) > 4:
        issues.append({
            'type': '字号不统一',
            'detail': f"检测到 {len(font_sizes)} 种字号"
        })
    return issues


def analyze_document(input_path: str) -> dict:
    """完整诊断，返回结构化结果

    Returns:
        {
            "punctuation": [...],
            "numbering": [...],
            "paragraph": [...],
            "font": [...],
            "summary": {"total_issues": N, "suggestions": [...]}
        }
    """
    doc = Document(input_path)
    results = {
        'punctuation': analyze_punctuation(doc),
        'numbering': analyze_numbering(doc),
        'paragraph': analyze_paragraph_format(doc),
        'font': analyze_font(doc),
    }

    total = (
        len(results['punctuation'])
        + len(results['numbering'])
        + len(results['paragraph'])
        + len(results['font'])
    )
    suggestions = []
    if results['punctuation']:
        suggestions.append('建议运行标点修复以修正标点问题')
    if results['paragraph'] or results['font']:
        suggestions.append('建议运行格式化以统一段落和字体格式')

    results['summary'] = {
        'total_issues': total,
        'suggestions': suggestions,
    }
    return results
