"""标点符号修复模块 — 适配自 docformat-gui (MIT License)
保留所有原始格式（字体、大小、颜色），只修正标点"""

import re
import logging
from docx import Document

logger = logging.getLogger('docformat.punctuation')

LEFT_DOUBLE_QUOTE = '\u201c'
RIGHT_DOUBLE_QUOTE = '\u201d'
LEFT_SINGLE_QUOTE = '\u2018'
RIGHT_SINGLE_QUOTE = '\u2019'

REPLACEMENTS = {
    "(": "（",
    ")": "）",
    ":": "：",
    ";": "；",
    "?": "？",
    "!": "！",
}

_PLACEHOLDER_PREFIX = "\x02PROT"


def _protect_special_patterns(text):
    """保护不应被替换的特殊模式"""
    protected = []
    counter = [0]

    def _replace_with_placeholder(match):
        placeholder = f"{_PLACEHOLDER_PREFIX}{counter[0]}\x03"
        protected.append((placeholder, match.group()))
        counter[0] += 1
        return placeholder

    result = text
    result = re.sub(r"(?:https?|ftp)://\S+", _replace_with_placeholder, result)
    result = re.sub(r"[\w.+-]+@[\w-]+\.[\w.-]+", _replace_with_placeholder, result)
    result = re.sub(r"[A-Za-z]:\\", _replace_with_placeholder, result)
    result = re.sub(r"[A-Za-z]+[\s-]?\d+:\d{2,}", _replace_with_placeholder, result)
    result = re.sub(r"(?<!\d)(\d{1,2}:\d{2}(?::\d{2})?)(?!\d)", _replace_with_placeholder, result)

    return result, protected


def _restore_protected(text, protected):
    result = text
    for placeholder, original in protected:
        result = result.replace(placeholder, original)
    return result


def has_chinese(text):
    return bool(re.search(r"[\u4e00-\u9fff]", text))


def fix_text(text):
    """修复文本中的标点（完整 7 步管线）"""
    if not text:
        return text
    result, protected = _protect_special_patterns(text)

    # 1. 省略号
    result = re.sub(r"\.{2,}", "……", result)
    result = re.sub(r"。{2,}", "……", result)
    # 2. 破折号
    result = re.sub(r"--+", "——", result)
    result = re.sub(r"—(?!—)", "——", result)
    # 3. 基本替换
    if has_chinese(result):
        for en, cn in REPLACEMENTS.items():
            result = result.replace(en, cn)
    # 4. 逗号
    result = re.sub(r"([\u4e00-\u9fff]),", r"\1，", result)
    result = re.sub(r",([\u4e00-\u9fff])", r"，\1", result)
    # 5. 句号
    result = re.sub(r"([\u4e00-\u9fff])\.(\s|$)", r"\1。\2", result)

    # 6. 双引号配对
    double_quote_chars = ['"', '\u201c', '\u201d', '\u201e', '\u201f', '\u300c', '\u300d']
    temp_result = result
    for q in double_quote_chars:
        temp_result = temp_result.replace(q, "\x00")
    if "\x00" in temp_result:
        chars = list(temp_result)
        quote_count = 0
        for i, c in enumerate(chars):
            if c == "\x00":
                chars[i] = LEFT_DOUBLE_QUOTE if quote_count % 2 == 0 else RIGHT_DOUBLE_QUOTE
                quote_count += 1
        result = "".join(chars)

    # 7. 单引号配对
    single_quote_chars = ["'", '\u2018', '\u2019', '\u201a', '\u201b']
    temp_result = result
    for q in single_quote_chars:
        temp_result = temp_result.replace(q, "\x01")
    if "\x01" in temp_result:
        chars = list(temp_result)
        quote_count = 0
        for i, c in enumerate(chars):
            if c == "\x01":
                chars[i] = LEFT_SINGLE_QUOTE if quote_count % 2 == 0 else RIGHT_SINGLE_QUOTE
                quote_count += 1
        result = "".join(chars)

    result = _restore_protected(result, protected)
    return result


def _fix_simple_punctuation(text):
    """不涉及配对的简单标点替换"""
    if not text:
        return text
    result, protected = _protect_special_patterns(text)
    result = re.sub(r"\.{2,}", "……", result)
    result = re.sub(r"。{2,}", "……", result)
    result = re.sub(r"--+", "——", result)
    result = re.sub(r"—(?!—)", "——", result)
    if has_chinese(result):
        for en, cn in REPLACEMENTS.items():
            result = result.replace(en, cn)
    result = re.sub(r"([\u4e00-\u9fff]),", r"\1，", result)
    result = re.sub(r",([\u4e00-\u9fff])", r"，\1", result)
    result = re.sub(r"([\u4e00-\u9fff])\.(\s|$)", r"\1。\2", result)
    result = _restore_protected(result, protected)
    return result


def _fix_quotes_whole_text(text):
    """对完整文本做引号配对替换"""
    result = text
    double_quote_chars = ['"', "\u201c", "\u201d", "\u201e", "\u201f", "\u300c", "\u300d"]
    temp = result
    for q in double_quote_chars:
        temp = temp.replace(q, "\x00")
    if "\x00" in temp:
        chars = list(temp)
        qi = 0
        for i, c in enumerate(chars):
            if c == "\x00":
                chars[i] = LEFT_DOUBLE_QUOTE if qi % 2 == 0 else RIGHT_DOUBLE_QUOTE
                qi += 1
        result = "".join(chars)

    single_quote_chars = ["'", "\u2018", "\u2019", "\u201a", "\u201b"]
    temp = result
    for q in single_quote_chars:
        temp = temp.replace(q, "\x01")
    if "\x01" in temp:
        chars = list(temp)
        qi = 0
        for i, c in enumerate(chars):
            if c == "\x01":
                chars[i] = LEFT_SINGLE_QUOTE if qi % 2 == 0 else RIGHT_SINGLE_QUOTE
                qi += 1
        result = "".join(chars)
    return result


def _redistribute_text_to_runs(runs, new_full_text):
    """重新分配文本到 runs，保留格式"""
    run_lengths = [len(run.text) for run in runs]
    total_original = sum(run_lengths)
    if len(new_full_text) == total_original:
        pos = 0
        for i, run in enumerate(runs):
            run.text = new_full_text[pos:pos + run_lengths[i]]
            pos += run_lengths[i]
    else:
        runs[0].text = new_full_text
        for run in runs[1:]:
            run.text = ""


def process_paragraph(para) -> bool:
    """处理段落标点，返回是否有修改"""
    full_text = para.text
    if not full_text.strip():
        return False
    runs = para.runs
    if not runs:
        return False

    changed = False
    for run in runs:
        original = run.text
        fixed = _fix_simple_punctuation(original)
        if fixed != original:
            run.text = fixed
            changed = True

    full_after_simple = para.text
    full_after_quotes = _fix_quotes_whole_text(full_after_simple)
    if full_after_quotes != full_after_simple:
        _redistribute_text_to_runs(runs, full_after_quotes)
        changed = True

    return changed


def process_document(input_path: str, output_path: str) -> dict:
    """处理文档标点，返回统计信息

    Returns:
        {"paragraphs_fixed": N, "table_cells_fixed": M}
    """
    logger.info(f"Reading: {input_path}")
    doc = Document(input_path)

    changes = 0
    for i, para in enumerate(doc.paragraphs):
        if process_paragraph(para):
            changes += 1
            preview = para.text[:50] + "..." if len(para.text) > 50 else para.text
            logger.debug(f"  Para {i + 1}: {preview}")

    table_changes = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if process_paragraph(para):
                        table_changes += 1

    logger.info(f"Total: {changes} paragraphs + {table_changes} table cells fixed")
    doc.save(output_path)
    logger.info(f"Saved: {output_path}")

    return {"paragraphs_fixed": changes, "table_cells_fixed": table_changes}
