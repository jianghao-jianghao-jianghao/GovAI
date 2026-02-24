"""
文档格式统一引擎 — 适配自 docformat-gui (MIT License)
移除 CLI / tkinter / PyInstaller 依赖，适配 FastAPI 服务端调用
"""

import json
import re
import copy
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger('docformat.formatter')

# ==================== python-docx 模板补丁 ====================
def _patch_docx_templates():
    """将 python-docx 模板嵌入代码，消除对文件系统路径的依赖 (Docker 兼容)"""
    from docx.parts.hdrftr import FooterPart, HeaderPart

    _DEFAULT_FOOTER_XML = b"""\
<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:ftr
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
    xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
    xmlns:mo="http://schemas.microsoft.com/office/mac/office/2008/main"
    xmlns:mv="urn:schemas-microsoft-com:mac:vml"
    xmlns:o="urn:schemas-microsoft-com:office:office"
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    xmlns:v="urn:schemas-microsoft-com:vml"
    xmlns:w10="urn:schemas-microsoft-com:office:word"
    xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
    xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
    xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
    xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
    xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    mc:Ignorable="w14 wp14"
    >
  <w:p>
    <w:pPr>
      <w:pStyle w:val="Footer"/>
    </w:pPr>
  </w:p>
</w:ftr>
"""
    _DEFAULT_HEADER_XML = b"""\
<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:hdr
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
    xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
    xmlns:mo="http://schemas.microsoft.com/office/mac/office/2008/main"
    xmlns:mv="urn:schemas-microsoft-com:mac:vml"
    xmlns:o="urn:schemas-microsoft-com:office:office"
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    xmlns:v="urn:schemas-microsoft-com:vml"
    xmlns:w10="urn:schemas-microsoft-com:office:word"
    xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
    xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
    xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
    xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
    xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    mc:Ignorable="w14 wp14"
    >
  <w:p>
    <w:pPr>
      <w:pStyle w:val="Header"/>
    </w:pPr>
  </w:p>
</w:hdr>
"""

    @classmethod
    def _patched_footer_xml(cls):
        return _DEFAULT_FOOTER_XML

    @classmethod
    def _patched_header_xml(cls):
        return _DEFAULT_HEADER_XML

    FooterPart._default_footer_xml = _patched_footer_xml
    HeaderPart._default_header_xml = _patched_header_xml
    logger.debug("python-docx 模板补丁已应用")

_patch_docx_templates()

# ==================== 预设 ====================

PRESETS = {
    'official': {
        'name': '公文格式',
        'page': {'top': 3.7, 'bottom': 3.5, 'left': 2.8, 'right': 2.6},
        # 主标题：二号方正小标宋简体，居中
        'title': {
            'font_cn': '方正小标宋简体', 'font_en': 'Times New Roman',
            'size': 22, 'bold': False, 'align': 'center', 'indent': 0,
            'line_spacing': 28, 'space_before': 0, 'space_after': 0,
        },
        # 主送机关：三号仿宋，顶格
        'recipient': {
            'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
            'size': 16, 'bold': False, 'align': 'left', 'indent': 0,
            'line_spacing': 28, 'space_before': 0, 'space_after': 0,
        },
        # 一级标题：三号黑体
        'heading1': {
            'font_cn': '黑体', 'font_en': 'Times New Roman',
            'size': 16, 'bold': False, 'align': 'left', 'indent': 32,
            'line_spacing': 28, 'space_before': 0, 'space_after': 0,
        },
        # 二级标题：三号楷体
        'heading2': {
            'font_cn': '楷体_GB2312', 'font_en': 'Times New Roman',
            'size': 16, 'bold': False, 'align': 'left', 'indent': 32,
            'line_spacing': 28, 'space_before': 0, 'space_after': 0,
        },
        # 三级标题：三号仿宋加粗
        'heading3': {
            'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
            'size': 16, 'bold': True, 'align': 'left', 'indent': 32,
            'line_spacing': 28, 'space_before': 0, 'space_after': 0,
        },
        # 四级标题：三号仿宋
        'heading4': {
            'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
            'size': 16, 'bold': False, 'align': 'left', 'indent': 32,
            'line_spacing': 28, 'space_before': 0, 'space_after': 0,
        },
        # 正文：三号仿宋，两端对齐，首行缩进2字符
        'body': {
            'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
            'size': 16, 'bold': False, 'align': 'justify',
            'indent': 32, 'line_spacing': 28, 'space_before': 0, 'space_after': 0,
        },
        # 落款单位：三号仿宋，右对齐
        'signature': {
            'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
            'size': 16, 'bold': False, 'align': 'right', 'indent': 0,
            'line_spacing': 28, 'space_before': 0, 'space_after': 0,
        },
        # 落款日期：三号仿宋，右对齐
        'date': {
            'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
            'size': 16, 'bold': False, 'align': 'right', 'indent': 0,
            'line_spacing': 28, 'space_before': 0, 'space_after': 0,
        },
        # 附件行：三号仿宋，首行缩进
        'attachment': {
            'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
            'size': 16, 'bold': False, 'align': 'justify', 'indent': 32,
            'line_spacing': 28, 'space_before': 0, 'space_after': 0,
        },
        # 结束语：三号仿宋，首行缩进
        'closing': {
            'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
            'size': 16, 'bold': False, 'align': 'left', 'indent': 32,
            'line_spacing': 28, 'space_before': 0, 'space_after': 0,
        },
        # 表格：小四仿宋
        'table': {
            'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
            'size': 12, 'bold': False, 'line_spacing': 22,
            'first_line_indent': 0, 'header_bold': True,
        },
        'first_line_bold': False,
        'page_number': True,
        'page_number_font': '宋体',
    },
    'academic': {
        'name': '学术论文格式',
        'page': {'top': 2.5, 'bottom': 2.5, 'left': 2.5, 'right': 2.5},
        'title': {
            'font_cn': '黑体', 'font_en': 'Times New Roman',
            'size': 18, 'bold': True, 'align': 'center', 'indent': 0,
            'line_spacing': 28, 'space_before': 0, 'space_after': 0,
        },
        'recipient': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 12, 'bold': False, 'align': 'left', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'heading1': {
            'font_cn': '黑体', 'font_en': 'Times New Roman',
            'size': 15, 'bold': True, 'align': 'left', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'heading2': {
            'font_cn': '黑体', 'font_en': 'Times New Roman',
            'size': 14, 'bold': True, 'align': 'left', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'heading3': {
            'font_cn': '黑体', 'font_en': 'Times New Roman',
            'size': 12, 'bold': False, 'align': 'left', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'heading4': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 12, 'bold': False, 'align': 'left', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'body': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 12, 'bold': False, 'align': 'justify',
            'indent': 24, 'line_spacing': None, 'space_before': 0, 'space_after': 0,
        },
        'signature': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 12, 'bold': False, 'align': 'right', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'date': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 12, 'bold': False, 'align': 'right', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'attachment': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 12, 'bold': False, 'align': 'left', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'closing': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 12, 'bold': False, 'align': 'left', 'indent': 24,
            'space_before': 0, 'space_after': 0,
        },
        'table': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 10, 'bold': False, 'line_spacing': 18,
            'first_line_indent': 0, 'header_bold': True,
        },
        'first_line_bold': False,
        'page_number': True,
        'page_number_font': 'Times New Roman',
    },
    'legal': {
        'name': '法律文书格式',
        'page': {'top': 3.0, 'bottom': 2.5, 'left': 3.0, 'right': 2.5},
        'title': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 22, 'bold': True, 'align': 'center', 'indent': 0,
            'line_spacing': 28, 'space_before': 0, 'space_after': 0,
        },
        'recipient': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 14, 'bold': False, 'align': 'left', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'heading1': {
            'font_cn': '黑体', 'font_en': 'Times New Roman',
            'size': 14, 'bold': False, 'align': 'left', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'heading2': {
            'font_cn': '黑体', 'font_en': 'Times New Roman',
            'size': 14, 'bold': False, 'align': 'left', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'heading3': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 14, 'bold': False, 'align': 'left', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'heading4': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 14, 'bold': False, 'align': 'left', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'body': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 14, 'bold': False, 'align': 'justify',
            'indent': 28, 'line_spacing': None, 'space_before': 0, 'space_after': 0,
        },
        'signature': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 14, 'bold': False, 'align': 'right', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'date': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 14, 'bold': False, 'align': 'right', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'attachment': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 14, 'bold': False, 'align': 'left', 'indent': 0,
            'space_before': 0, 'space_after': 0,
        },
        'closing': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 14, 'bold': False, 'align': 'left', 'indent': 28,
            'space_before': 0, 'space_after': 0,
        },
        'table': {
            'font_cn': '宋体', 'font_en': 'Times New Roman',
            'size': 12, 'bold': False, 'line_spacing': 22,
            'first_line_indent': 0, 'header_bold': True,
        },
        'first_line_bold': False,
        'page_number': True,
        'page_number_font': '宋体',
    },
}


# ==================== 辅助函数 ====================

def remove_background(doc):
    """移除页面背景颜色"""
    body = doc._body._body
    document = body.getparent()
    for elem in list(document):
        tag_name = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag_name == 'background':
            document.remove(elem)

    for para in doc.paragraphs:
        pPr = para._p.get_or_add_pPr()
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            pPr.remove(shd)
        for run in para.runs:
            run.font.highlight_color = None
            rPr = run._r.get_or_add_rPr()
            shd = rPr.find(qn('w:shd'))
            if shd is not None:
                rPr.remove(shd)


def _iter_block_items(doc):
    """Yield paragraphs and tables in document order."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith('}p'):
            yield Paragraph(child, doc)
        elif child.tag.endswith('}tbl'):
            yield Table(child, doc)


# ==================== 表格辅助 ====================

def _set_table_borders(table, size_pt=0.5, color="000000"):
    size = max(1, int(size_pt * 8))
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), str(size))
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), color)
        borders.append(elem)


def _set_table_cell_margins(table, top_cm=0.0, bottom_cm=0.0, left_cm=0.05, right_cm=0.05):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    cell_mar = tbl_pr.find(qn('w:tblCellMar'))
    if cell_mar is None:
        cell_mar = OxmlElement('w:tblCellMar')
        tbl_pr.append(cell_mar)

    def _set_side(tag, cm_value):
        node = cell_mar.find(qn(f'w:{tag}'))
        if node is None:
            node = OxmlElement(f'w:{tag}')
            cell_mar.append(node)
        node.set(qn('w:type'), 'dxa')
        node.set(qn('w:w'), str(int(Cm(cm_value).twips)))

    _set_side('top', top_cm)
    _set_side('bottom', bottom_cm)
    _set_side('left', left_cm)
    _set_side('right', right_cm)


def _set_table_width_percent(table, percent=100):
    percent = max(1, min(100, int(percent)))
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:type'), 'pct')
    tbl_w.set(qn('w:w'), str(percent * 50))


def _set_table_indent(table, indent_twips=0):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:type'), 'dxa')
    tbl_ind.set(qn('w:w'), str(int(indent_twips)))


def _text_weight(text):
    weight = 0.0
    for ch in text:
        if ord(ch) < 128:
            weight += 0.5
        else:
            weight += 1.0
    return weight


def _normalize_pcts(weights, min_pct, max_pct):
    total = sum(weights) or 1.0
    pcts = [w / total * 100 for w in weights]
    for i, v in enumerate(pcts):
        if v < min_pct:
            pcts[i] = min_pct
    for i, v in enumerate(pcts):
        if v > max_pct:
            pcts[i] = max_pct
    total = sum(pcts) or 1.0
    return [v / total * 100 for v in pcts]


def _set_table_col_widths_by_content(table, min_pct=8, max_pct=45):
    if not table.rows:
        return
    col_count = max(len(row.cells) for row in table.rows)
    if col_count == 0:
        return
    max_weights = [1.0] * col_count
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            text = ''.join(p.text for p in cell.paragraphs).strip()
            if text:
                max_weights[c_idx] = max(max_weights[c_idx], _text_weight(text))
    pcts = _normalize_pcts(max_weights, min_pct, max_pct)

    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement('w:tblGrid')
        tbl.insert(0, tbl_grid)
    else:
        for child in list(tbl_grid):
            tbl_grid.remove(child)
    for pct in pcts:
        grid_col = OxmlElement('w:gridCol')
        grid_col.set(qn('w:w'), str(int(pct * 50)))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            tc = cell._tc
            tc_pr = tc.tcPr
            if tc_pr is None:
                tc_pr = OxmlElement('w:tcPr')
                tc.insert(0, tc_pr)
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:type'), 'pct')
            tc_w.set(qn('w:w'), str(int(pcts[c_idx] * 50)))


def _insert_paragraph_after_table(table, text=""):
    p = OxmlElement("w:p")
    table._tbl.addnext(p)
    para = Paragraph(p, table._parent)
    if text:
        para.add_run(text)
    return para


def _insert_paragraph_before_table(table, text=""):
    p = OxmlElement("w:p")
    table._tbl.addprevious(p)
    para = Paragraph(p, table._parent)
    if text:
        para.add_run(text)
    return para


def _insert_paragraph_after_paragraph(paragraph, text=""):
    p = OxmlElement("w:p")
    paragraph._p.addnext(p)
    para = Paragraph(p, paragraph._parent)
    if text:
        para.add_run(text)
    return para


def _insert_paragraph_before_paragraph(paragraph, text=""):
    p = OxmlElement("w:p")
    paragraph._p.addprevious(p)
    para = Paragraph(p, paragraph._parent)
    if text:
        para.add_run(text)
    return para


def _is_numeric_text(text):
    text = text.replace(',', '').replace('％', '%').strip()
    if not text:
        return False
    return re.match(r'^[-+]?\d+(?:\.\d+)?%?$', text) is not None


def _is_short_text(text, max_len=4):
    text = text.strip()
    return 0 < len(text) <= max_len


def _is_table_title(text):
    text = text.strip()
    if not text or len(text) > 30:
        return False
    return re.match(r'^表\s*(?:\d+|[一二三四五六七八九十]+)(?:[-—._、]\d+)?', text) is not None


def _is_table_unit(text):
    text = text.strip()
    if not text or len(text) > 20:
        return False
    return re.match(r'^单位\s*[:：]', text) is not None


def _set_cell_borders(cell, size_pt=0.5, color="000000"):
    size = max(1, int(size_pt * 8))
    tc = cell._tc
    tc_pr = tc.tcPr
    if tc_pr is None:
        tc_pr = OxmlElement('w:tcPr')
        tc.insert(0, tc_pr)
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)
    for edge in ('top', 'left', 'bottom', 'right'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), str(size))
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), color)
        borders.append(elem)


# ==================== 段落类型检测 ====================

def detect_para_type(text, index, total, alignment, all_texts, all_texts_index=None):
    """检测段落类型 → title/recipient/heading1-4/body/signature/date/attachment/closing"""
    text = text.strip()
    if not text:
        return 'empty'

    # 一级标题
    if re.match(r'^[一二三四五六七八九十]+、', text):
        return 'heading1'

    # 二级标题
    if re.match(r'^（[一二三四五六七八九十]+）', text) or re.match(r'^\([一二三四五六七八九十]+\)', text):
        return 'heading2'

    # 三级标题
    if re.match(r'^\d+\.\s*\S', text) and len(text) < 60:
        return 'heading3'

    # 四级标题
    if (re.match(r'^（\d+）', text) or re.match(r'^\(\d+\)', text)) and len(text) < 60:
        return 'heading4'

    # 主送机关
    if re.match(r'^[\u4e00-\u9fff\d、，,（）()\s]+[：:]$', text) and len(text) < 30:
        body_indicators = (
            r'(现将|为了|根据|按照|经研究|为贯彻|为落实|为进一步|为深入|'
            r'如下|以下|特此|兹将|报告如下|说明如下|通知如下|汇报如下|'
            r'的意见|的通知|的报告|的决定|的请示|的函)'
        )
        if not re.search(body_indicators, text):
            return 'recipient'

    # 附件行
    if re.match(r'^附件[：:]\s*', text):
        return 'attachment'
    if re.match(r'^附件\d*[：:．.\s]', text):
        return 'attachment'
    if re.match(r'^附件$', text):
        return 'attachment'

    # 结束语
    closing_patterns = [
        r'^特此(说明|通知|报告|函复|函告|批复|公告|通报)。?$',
        r'^此致$',
        r'^敬礼[！!]?$',
        r'^以上(报告|意见|方案).{0,10}$',
        r'^妥否.{0,10}$',
        r'^请.{0,15}(批示|审批|审议|指示|核准)。?$',
    ]
    for pattern in closing_patterns:
        if re.match(pattern, text):
            return 'closing'

    # 落款日期
    date_patterns = [
        r'^\d{4}年\d{1,2}月\d{1,2}日$',
        r'^\d{4}\.\d{1,2}\.\d{1,2}$',
        r'^\d{4}/\d{1,2}/\d{1,2}$',
        r'^\d{4}-\d{1,2}-\d{1,2}$',
        r'^二[○〇零oO0][一二三四五六七八九零〇○oO0]{2}年.{1,3}月.{1,3}日$',
    ]
    for pattern in date_patterns:
        if re.match(pattern, text):
            return 'date'

    # 落款单位
    if index >= total - 10 and len(text) < 30:
        if re.search(r'(公司|局|委|部|厅|院|所|中心|办公室|集团|银行|学校|大学|医院)$', text):
            return 'signature'
        if all_texts_index is not None:
            remaining_texts = all_texts[all_texts_index + 1:]
        else:
            remaining_texts = []
        for next_text in remaining_texts[:3]:
            for pattern in date_patterns:
                if re.match(pattern, next_text.strip()):
                    return 'signature'

    # 主标题
    if index < 5:
        _check_idx = all_texts_index if all_texts_index is not None else 0
        _title_region_ended = False
        for pt in all_texts[:_check_idx]:
            pt_s = pt.strip()
            if re.search(r'[：:]\s*$', pt_s) and len(pt_s) < 50:
                _title_region_ended = True
                break
            if re.match(r'^[一二三四五六七八九十]+、', pt_s):
                _title_region_ended = True
                break

        if not _title_region_ended:
            title_patterns = [
                r'^关于.+的(通知|报告|请示|函|意见|决定|公告|通报|批复|说明|方案|总结|汇报|复函|答复|建议)$',
                r'^.{2,30}(通知|报告|请示|函|意见|决定|公告|通报|批复|工作方案|工作总结|实施方案|管理办法|暂行规定)$',
            ]
            for pattern in title_patterns:
                if re.match(pattern, text):
                    return 'title'
            if 15 < len(text) < 80 and not re.search(r'[。！？，、；：]$', text):
                if not re.match(r'^[一二三四五六七八九十\d（(]', text):
                    return 'title'
            if alignment == WD_ALIGN_PARAGRAPH.CENTER and len(text) < 60:
                return 'title'

    return 'body'


def _split_heading_by_punct(paragraph):
    """拆分 '（三）xxx：正文' 为标题段 + 正文段"""
    text = paragraph.text.strip()
    if not text:
        return False
    if not (
        re.match(r'^[一二三四五六七八九十]+、', text) or
        re.match(r'^（[一二三四五六七八九十]+）', text) or
        re.match(r'^\([一二三四五六七八九十]+\)', text) or
        re.match(r'^\d+\.\s*\S', text) or
        re.match(r'^（\d+）', text) or
        re.match(r'^\(\d+\)', text)
    ):
        return False
    punct_positions = []
    for ch in ('：', ':', '。'):
        pos = text.find(ch)
        if pos != -1:
            punct_positions.append(pos)
    if not punct_positions:
        return False
    split_idx = min(punct_positions)
    head = text[:split_idx + 1].strip()
    tail = text[split_idx + 1:].strip()
    if not tail:
        return False
    paragraph.text = head
    _insert_paragraph_after_paragraph(paragraph, text=tail)
    return True


# ==================== 字体设置 ====================

def set_font(run, font_cn, font_en, size, bold=False):
    """设置字体，同时清除斜体/下划线/颜色/删除线"""
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = False
    run.font.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.strike = False
    run.font.double_strike = False
    run.font.subscript = False
    run.font.superscript = False

    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:cs'), font_en)


# ==================== 段落格式化 ====================

def format_paragraph(para, fmt, para_type, line_spacing_pt=28, first_line_bold=False):
    """格式化单个段落"""
    pf = para.paragraph_format

    align_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    pf.alignment = align_map.get(fmt.get('align', 'justify'), WD_ALIGN_PARAGRAPH.JUSTIFY)

    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)

    indent = fmt.get('indent', 0)
    pf.first_line_indent = Pt(indent) if indent > 0 else Pt(0)

    ls = fmt.get('line_spacing', line_spacing_pt)
    if ls:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(ls)
    else:
        pf.line_spacing = 1.5

    pf.space_before = Pt(fmt.get('space_before', 0))
    pf.space_after = Pt(fmt.get('space_after', 0))

    # 首句加粗
    if first_line_bold and para_type == 'body':
        full_text = para.text
        first_sentence_end = full_text.find('。')
        if first_sentence_end != -1:
            split_idx = first_sentence_end + 1
            first_part = full_text[:split_idx]
            rest_part = full_text[split_idx:]
            for run in list(para.runs):
                para._p.remove(run._r)
            run1 = para.add_run(first_part)
            set_font(run1, fmt['font_cn'], fmt['font_en'], fmt['size'], bold=True)
            if rest_part:
                run2 = para.add_run(rest_part)
                set_font(run2, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False))
        else:
            for run in para.runs:
                set_font(run, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False))
    else:
        # "一是/二是..." 加粗前缀
        if para_type == 'body':
            m = re.match(r'^([一二三四五六七八九十]{1,3}是)([：:、]?)', para.text)
            if m:
                lead = m.group(1) + (m.group(2) or '')
                rest = para.text[len(lead):]
                for run in list(para.runs):
                    para._p.remove(run._r)
                run1 = para.add_run(lead)
                set_font(run1, fmt['font_cn'], fmt['font_en'], fmt['size'], bold=True)
                if rest:
                    run2 = para.add_run(rest)
                    set_font(run2, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False))
                return

        for run in para.runs:
            set_font(run, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False))


# ==================== 页码 ====================

def add_page_number(doc, font_name="宋体"):
    """添加页码（四号宋体，左右一字线，奇右偶左）"""
    try:
        doc.settings.odd_and_even_pages_header_footer = True
    except Exception:
        settings_el = doc.settings._element
        if settings_el.find(qn('w:evenAndOddHeaders')) is None:
            settings_el.append(OxmlElement('w:evenAndOddHeaders'))

    for section in doc.sections:
        section.odd_and_even_pages_header_footer = True
        section.footer_distance = Cm(0.7)

        odd_footer = section.footer
        even_footer = section.even_page_footer
        odd_footer.is_linked_to_previous = False
        even_footer.is_linked_to_previous = False

        for f in (odd_footer, even_footer):
            for para in f.paragraphs:
                para.clear()

        def _build_footer_line(footer, align, pad_fullwidth):
            if footer.paragraphs:
                para = footer.paragraphs[0]
            else:
                para = footer.add_paragraph()
            para.alignment = align

            if pad_fullwidth:
                run0 = para.add_run("　")
                set_font(run0, font_name, font_name, 14, bold=False)

            run1 = para.add_run("— ")
            set_font(run1, font_name, font_name, 14, bold=False)

            run2 = para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run2._r.append(fldChar1)
            set_font(run2, font_name, font_name, 14, bold=False)

            run3 = para.add_run()
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            run3._r.append(instrText)
            set_font(run3, font_name, font_name, 14, bold=False)

            run4 = para.add_run()
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run4._r.append(fldChar2)
            set_font(run4, font_name, font_name, 14, bold=False)

            run5 = para.add_run(" —")
            set_font(run5, font_name, font_name, 14, bold=False)

            if not pad_fullwidth:
                run6 = para.add_run("　")
                set_font(run6, font_name, font_name, 14, bold=False)

        _build_footer_line(odd_footer, WD_ALIGN_PARAGRAPH.RIGHT, pad_fullwidth=True)
        _build_footer_line(even_footer, WD_ALIGN_PARAGRAPH.LEFT, pad_fullwidth=False)


# ==================== 主流程 ====================

def format_document(input_path: str, output_path: str,
                    preset_name: str = 'official',
                    custom_preset: dict = None,
                    progress_callback=None) -> dict:
    """格式化文档（服务端版本）

    Args:
        input_path: 输入 .docx 文件路径
        output_path: 输出 .docx 文件路径
        preset_name: 预设名 official / academic / legal / custom
        custom_preset: 当 preset_name='custom' 时传入的自定义预设字典
        progress_callback: 可选 callback(current, total, stage_text)

    Returns:
        dict: 统计信息 {para_type: count, ...}

    Raises:
        ValueError: 未知预设名
        FileNotFoundError: 输入文件不存在
    """
    # 选择预设
    if preset_name == 'custom':
        if custom_preset:
            preset = custom_preset
        else:
            logger.warning('Custom preset not provided, falling back to official')
            preset = PRESETS['official']
    elif preset_name not in PRESETS:
        raise ValueError(f'未知预设: {preset_name}，可选: {", ".join(PRESETS.keys())}, custom')
    else:
        preset = PRESETS[preset_name]

    logger.info(f'Preset: {preset.get("name", preset_name)}')
    logger.info(f'Input: {input_path}')

    first_line_bold = preset.get('first_line_bold', False)
    doc = Document(input_path)

    # 拆分标题+正文
    for para in list(doc.paragraphs):
        _split_heading_by_punct(para)

    total_paras = len(doc.paragraphs)
    all_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    all_texts_idx_map = {}
    at_idx = 0
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            all_texts_idx_map[i] = at_idx
            at_idx += 1

    def _progress(current, total, stage):
        if progress_callback:
            progress_callback(current, total, stage)

    # 1. 移除背景
    logger.info('1. Removing background...')
    _progress(0, 100, '移除背景...')
    remove_background(doc)

    # 2. 页边距
    logger.info('2. Setting page margins...')
    _progress(5, 100, '设置页面边距...')
    page = preset['page']
    for section in doc.sections:
        section.top_margin = Cm(page['top'])
        section.bottom_margin = Cm(page['bottom'])
        section.left_margin = Cm(page['left'])
        section.right_margin = Cm(page['right'])

    # 3. 格式化段落
    logger.info('3. Formatting paragraphs...')
    _progress(10, 100, '格式化段落...')
    stats = {
        'title': 0, 'recipient': 0, 'heading1': 0, 'heading2': 0,
        'heading3': 0, 'heading4': 0, 'body': 0, 'signature': 0,
        'date': 0, 'attachment': 0, 'closing': 0,
    }

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        para_type = detect_para_type(
            text, i, total_paras,
            para.paragraph_format.alignment,
            all_texts,
            all_texts_index=all_texts_idx_map.get(i),
        )
        fmt_key = para_type if para_type in preset else 'body'
        fmt = preset.get(fmt_key, preset['body'])
        format_paragraph(para, fmt, para_type, first_line_bold=first_line_bold)
        stats[para_type] = stats.get(para_type, 0) + 1

        preview = text[:35] + '...' if len(text) > 35 else text
        logger.debug(f'   [{para_type:10}] {preview}')

        if total_paras > 0:
            pct = 10 + int(70 * (i + 1) / total_paras)
            _progress(pct, 100, f'格式化段落 ({i + 1}/{total_paras})')

    # 4. 表格
    logger.info('4. Formatting tables...')
    _progress(82, 100, '格式化表格...')
    body_fmt = preset.get('body', {})
    table_fmt = preset.get('table', {})
    table_defaults = {
        'optimize': True, 'border_size_pt': 0.5, 'width_percent': 100,
        'auto_col_width': True, 'col_min_pct': 8, 'col_max_pct': 45,
        'row_height_cm': 0.7, 'cell_margin_top_cm': 0.0, 'cell_margin_bottom_cm': 0.0,
        'cell_margin_left_cm': 0.05, 'cell_margin_right_cm': 0.05,
        'paragraph_single': True, 'after_table_blank_line': True,
        'title_align': 'center', 'unit_align': 'right',
        'unit_space_before_lines': 0.5, 'short_text_len': 4,
    }
    table_cfg = {**table_defaults, **table_fmt}

    tbl_font_cn = table_fmt.get('font_cn', body_fmt.get('font_cn', '仿宋_GB2312'))
    tbl_font_en = table_fmt.get('font_en', body_fmt.get('font_en', 'Times New Roman'))
    tbl_size = table_fmt.get('size', body_fmt.get('size', 16))
    tbl_bold = table_fmt.get('bold', False)
    tbl_line_spacing = table_fmt.get('line_spacing', body_fmt.get('line_spacing', 28))
    tbl_header_bold = table_fmt.get('header_bold', False)
    tbl_first_line_indent = table_fmt.get('first_line_indent', 0)

    blocks = list(_iter_block_items(doc))
    for idx, block in enumerate(blocks):
        if not isinstance(block, Table):
            continue
        table = block
        if table_cfg.get('optimize', True):
            table.autofit = not table_cfg.get('auto_col_width', True)
            _set_table_width_percent(table, table_cfg.get('width_percent', 100))
            _set_table_indent(table, 0)
            _set_table_borders(table, size_pt=table_cfg.get('border_size_pt', 0.5))
            _set_table_cell_margins(
                table,
                top_cm=table_cfg.get('cell_margin_top_cm', 0.0),
                bottom_cm=table_cfg.get('cell_margin_bottom_cm', 0.0),
                left_cm=table_cfg.get('cell_margin_left_cm', 0.05),
                right_cm=table_cfg.get('cell_margin_right_cm', 0.05),
            )
            if table_cfg.get('auto_col_width', True):
                _set_table_col_widths_by_content(
                    table,
                    min_pct=table_cfg.get('col_min_pct', 8),
                    max_pct=table_cfg.get('col_max_pct', 45),
                )

        prev_block = blocks[idx - 1] if idx - 1 >= 0 else None
        prev_para_is_title = isinstance(prev_block, Paragraph) and _is_table_title(prev_block.text)
        prev_para_is_unit = isinstance(prev_block, Paragraph) and _is_table_unit(prev_block.text)

        if isinstance(prev_block, Paragraph):
            if prev_block.text.strip():
                if prev_para_is_title or prev_para_is_unit:
                    _insert_paragraph_before_paragraph(prev_block, text="")
                else:
                    _insert_paragraph_before_table(table, text="")
        elif isinstance(prev_block, Table):
            _insert_paragraph_after_table(prev_block, text="")
        else:
            if idx == 0:
                _insert_paragraph_before_table(table, text="")

        if prev_para_is_title:
            if table_cfg.get('title_align', 'center') == 'center':
                prev_block.alignment = WD_ALIGN_PARAGRAPH.CENTER
            prev_block.paragraph_format.space_before = Pt(0)
            prev_block.paragraph_format.space_after = Pt(0)
            prev_block.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
        unit_para = None
        if isinstance(next_block, Paragraph) and _is_table_unit(next_block.text):
            unit_para = next_block
            if table_cfg.get('unit_align', 'right') == 'right':
                unit_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            unit_space_lines = table_cfg.get('unit_space_before_lines', 0.5)
            unit_para.paragraph_format.space_before = Pt(tbl_size * unit_space_lines)
            unit_para.paragraph_format.space_after = Pt(0)
            unit_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        serial_col_idx = None
        if table.rows:
            header_cells = table.rows[0].cells
            for c_idx, cell in enumerate(header_cells):
                head_text = ''.join(p.text for p in cell.paragraphs).strip()
                if '序号' in head_text or head_text == '序':
                    serial_col_idx = c_idx
                    break

        for row_idx, row in enumerate(table.rows):
            if table_cfg.get('row_height_cm'):
                row.height = Cm(table_cfg.get('row_height_cm'))
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

            for col_idx, cell in enumerate(row.cells):
                if table_cfg.get('optimize', True):
                    _set_cell_borders(cell, size_pt=table_cfg.get('border_size_pt', 0.5))

                cell_text = ''.join(p.text for p in cell.paragraphs).strip()
                for para in cell.paragraphs:
                    if para.text.strip():
                        is_header = (row_idx == 0 and tbl_header_bold)
                        for run in para.runs:
                            set_font(run, tbl_font_cn, tbl_font_en, tbl_size, bold=(tbl_bold or is_header))

                    para.paragraph_format.first_line_indent = Pt(tbl_first_line_indent)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    if table_cfg.get('paragraph_single', True):
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    else:
                        if tbl_line_spacing:
                            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                            para.paragraph_format.line_spacing = Pt(tbl_line_spacing)
                        else:
                            para.paragraph_format.line_spacing = 1.5

                    if row_idx == 0:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif '合计' in cell_text or '总计' in cell_text:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif serial_col_idx is not None and col_idx == serial_col_idx:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif _is_numeric_text(cell_text):
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif _is_short_text(cell_text, table_cfg.get('short_text_len', 4)):
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if table_cfg.get('after_table_blank_line', True):
            next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
            if unit_para is not None:
                after_unit = blocks[idx + 2] if idx + 2 < len(blocks) else None
                if not (isinstance(after_unit, Paragraph) and not after_unit.text.strip()):
                    _insert_paragraph_after_paragraph(unit_para, text="")
            else:
                if not (isinstance(next_block, Paragraph) and not next_block.text.strip()):
                    _insert_paragraph_after_table(table, text="")

    # 5. 页码
    _progress(90, 100, '添加页码...')
    if preset.get('page_number', True):
        logger.info('5. Adding page numbers...')
        add_page_number(doc, font_name=preset.get('page_number_font', '宋体'))
    else:
        logger.info('5. Skipping page numbers...')

    # 保存
    _progress(95, 100, '保存文件...')
    doc.save(output_path)
    _progress(100, 100, '格式化完成')

    logger.info(f'Statistics: {stats}')
    logger.info(f'Output: {output_path}')
    return stats
