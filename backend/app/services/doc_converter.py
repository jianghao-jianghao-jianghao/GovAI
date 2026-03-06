"""
统一文档转换服务（微服务客户端）
================================
通过 HTTP 调用独立的 converter 微服务完成文档转换和文本提取。
converter 微服务基于 LibreOffice headless，支持 PDF、DOCX、XLSX、PPTX 等格式。

主入口:
    - convert_file_to_markdown(file_path, file_name)      —— 从磁盘文件提取文本
    - convert_bytes_to_markdown(content_bytes, file_name)  —— 从内存字节提取文本
    - convert_to_pdf(file_path_or_bytes, file_name)        —— 转为 PDF
    - convert_and_extract(content_bytes, file_name)        —— 同时转 PDF + 提取文本

降级策略:
    converter 微服务不可用时自动降级到内置简单解析器。
"""

import asyncio
import logging
import os
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path

import httpx

logger = logging.getLogger(__name__)

# ── converter 微服务地址 ──
CONVERTER_URL = os.getenv("CONVERTER_URL", "http://converter:8001")

# ── 格式常量 ──

SUPPORTED_EXTENSIONS: set[str] = {
    "pdf", "docx", "doc", "pptx", "ppt",
    "xlsx", "xls", "csv",
    "html", "htm", "txt", "md",
    "json", "xml",
    "odt", "ods", "odp", "rtf",
}

_PLAINTEXT_EXTENSIONS: set[str] = {"txt", "md"}

KB_ALLOWED_EXTENSIONS: set[str] = {
    "pdf", "docx", "doc", "txt", "md", "csv", "xlsx", "xls",
    "pptx", "ppt", "html", "htm", "json", "xml",
}

DOC_IMPORT_EXTENSIONS: set[str] = {
    "pdf", "docx", "doc", "txt", "md", "csv", "xlsx",
    "pptx", "html", "htm",
}


# ── 结果数据类 ──

@dataclass
class DocumentConvertResult:
    """文档转换结果"""
    markdown: str = ""
    title: str = ""
    success: bool = True
    error_message: str = ""
    source_format: str = ""
    char_count: int = 0
    pdf_path: str = ""          # converter 微服务返回的 PDF 共享路径

    def __post_init__(self):
        self.char_count = len(self.markdown)


# ── HTTP 客户端 ──

async def _call_converter(
    endpoint: str,
    file_bytes: bytes,
    file_name: str,
    timeout: float = 120.0,
) -> httpx.Response:
    """调用 converter 微服务"""
    url = f"{CONVERTER_URL}{endpoint}"
    async with httpx.AsyncClient(timeout=timeout) as client:
        resp = await client.post(
            url,
            files={"file": (file_name, file_bytes, "application/octet-stream")},
        )
        resp.raise_for_status()
        return resp


# ── 公开 API ──

async def convert_file_to_markdown(
    file_path: str | Path,
    file_name: str = "",
) -> DocumentConvertResult:
    """
    将本地文件转换为纯文本（通过 converter 微服务提取文本）。

    Args:
        file_path: 本地文件绝对路径
        file_name: 原始文件名（用于推断格式和标题）

    Returns:
        DocumentConvertResult
    """
    file_path = Path(file_path)
    if not file_path.exists():
        return DocumentConvertResult(
            success=False,
            error_message=f"文件不存在: {file_path}",
        )

    ext = file_path.suffix.lstrip(".").lower()
    title = Path(file_name).stem if file_name else file_path.stem

    # 纯文本直接读取
    if ext in _PLAINTEXT_EXTENSIONS:
        content = _read_text_safe(file_path)
        return DocumentConvertResult(markdown=content, title=title, source_format=ext)

    # 调用 converter 微服务
    try:
        file_bytes = file_path.read_bytes()
        resp = await _call_converter(
            "/extract-text",
            file_bytes,
            file_name or file_path.name,
        )
        data = resp.json()
        text = _post_process_text(data.get("text", ""))
        return DocumentConvertResult(
            markdown=text,
            title=title,
            source_format=ext,
        )
    except Exception as e:
        logger.warning(f"converter 微服务调用失败 [{file_name or file_path.name}]: {e}")
        # 降级：尝试本地简单提取
        fallback_text = _local_fallback_extract(file_path, ext)
        if fallback_text:
            return DocumentConvertResult(
                markdown=_post_process_text(fallback_text),
                title=title,
                source_format=ext,
            )
        return DocumentConvertResult(
            success=False,
            title=title,
            error_message=f"文档转换失败: {str(e)}",
            source_format=ext,
        )


async def convert_bytes_to_markdown(
    content_bytes: bytes,
    file_name: str,
) -> DocumentConvertResult:
    """
    将文件字节内容提取为纯文本。

    Args:
        content_bytes: 文件二进制内容
        file_name:     原始文件名（含扩展名）

    Returns:
        DocumentConvertResult
    """
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else "bin"
    title = Path(file_name).stem

    # 纯文本直接解码
    if ext in _PLAINTEXT_EXTENSIONS:
        text = _decode_bytes_safe(content_bytes)
        return DocumentConvertResult(markdown=text, title=title, source_format=ext)

    # 调用 converter 微服务
    try:
        resp = await _call_converter("/extract-text", content_bytes, file_name)
        data = resp.json()
        text = _post_process_text(data.get("text", ""))
        return DocumentConvertResult(
            markdown=text,
            title=title,
            source_format=ext,
        )
    except Exception as e:
        logger.warning(f"converter 微服务文本提取失败 [{file_name}]: {e}")
        # 降级处理
        fallback_text = _local_fallback_extract_bytes(content_bytes, ext)
        if fallback_text:
            return DocumentConvertResult(
                markdown=_post_process_text(fallback_text),
                title=title,
                source_format=ext,
            )
        return DocumentConvertResult(
            success=False,
            title=title,
            error_message=f"文档转换失败: {str(e)}",
            source_format=ext,
        )


async def convert_and_extract(
    content_bytes: bytes,
    file_name: str,
) -> DocumentConvertResult:
    """
    同时完成 PDF 转换 + 文本提取（调用 converter 微服务一次完成）。

    Returns:
        DocumentConvertResult（包含 pdf_path 字段）
    """
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else "bin"
    title = Path(file_name).stem

    try:
        resp = await _call_converter("/convert-and-extract", content_bytes, file_name)
        data = resp.json()
        text = _post_process_text(data.get("text", ""))
        return DocumentConvertResult(
            markdown=text,
            title=title,
            source_format=ext,
            pdf_path=data.get("pdf_path", ""),
        )
    except Exception as e:
        logger.warning(f"converter 微服务 convert-and-extract 失败 [{file_name}]: {e}")
        # 降级：仅提取文本
        result = await convert_bytes_to_markdown(content_bytes, file_name)
        return result


async def convert_to_pdf_bytes(
    content_bytes: bytes,
    file_name: str,
) -> bytes | None:
    """
    将文件转为 PDF，返回 PDF 字节。

    Returns:
        PDF bytes 或 None（失败时）
    """
    try:
        resp = await _call_converter("/convert-to-pdf", content_bytes, file_name)
        return resp.content
    except Exception as e:
        logger.error(f"PDF 转换失败 [{file_name}]: {e}")
        return None


async def save_markdown_file(
    md_content: str,
    target_dir: str | Path,
    file_id: str,
) -> Path:
    """将 Markdown 内容保存到指定目录"""
    target_dir = Path(target_dir)
    target_dir.mkdir(parents=True, exist_ok=True)
    md_path = target_dir / f"{file_id}.md"

    loop = asyncio.get_running_loop()
    await loop.run_in_executor(
        None,
        lambda: md_path.write_text(md_content, encoding="utf-8"),
    )
    return md_path


# ── 工具函数 ──

def is_supported_format(ext: str) -> bool:
    """检查扩展名是否受支持"""
    return ext.lower().lstrip(".") in SUPPORTED_EXTENSIONS


def get_supported_formats() -> list[str]:
    """返回所有支持的文件扩展名（已排序）"""
    return sorted(SUPPORTED_EXTENSIONS)


# ── 内部工具 ──

def _read_text_safe(file_path: Path) -> str:
    """安全读取文本文件，自动探测编码"""
    for encoding in ("utf-8", "utf-8-sig", "gbk", "gb2312", "gb18030", "latin-1"):
        try:
            return file_path.read_text(encoding=encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return file_path.read_text(encoding="utf-8", errors="replace")


def _decode_bytes_safe(data: bytes) -> str:
    """安全解码字节内容，自动探测编码（支持中文 GBK/GB2312/GB18030）"""
    for encoding in ("utf-8", "utf-8-sig", "gbk", "gb2312", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


def _post_process_text(text: str) -> str:
    """清理/规范化提取的文本"""
    if not text:
        return ""

    # 移除 null 字节和其他控制字符（保留 \n \r \t）
    # .doc 等格式转换后可能包含 null 字节，PostgreSQL 不允许存储
    text = text.replace("\x00", "")
    text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

    lines = text.split("\n")
    cleaned: list[str] = []
    consecutive_blank = 0

    for line in lines:
        if line.strip() == "":
            consecutive_blank += 1
            if consecutive_blank <= 2:
                cleaned.append("")
        else:
            consecutive_blank = 0
            cleaned.append(line)

    result = "\n".join(cleaned).strip()
    return result


def _local_fallback_extract(file_path: Path, ext: str) -> str | None:
    """本地降级提取（不依赖 converter 微服务）"""
    try:
        if ext == "docx":
            return _fallback_docx(file_path)
        elif ext == "doc":
            return _fallback_doc_binary(file_path)
        elif ext == "csv":
            return _fallback_csv(file_path)
        elif ext in ("txt", "md"):
            return _read_text_safe(file_path)
    except Exception as e:
        logger.warning(f"本地降级提取失败 [{ext}]: {e}")
    return None


def _local_fallback_extract_bytes(content_bytes: bytes, ext: str) -> str | None:
    """从 bytes 降级提取"""
    if ext in ("txt", "md", "csv"):
        return _decode_bytes_safe(content_bytes)

    if ext == "docx":
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        try:
            tmp.write(content_bytes)
            tmp.close()
            return _fallback_docx(Path(tmp.name))
        finally:
            try:
                Path(tmp.name).unlink(missing_ok=True)
            except Exception:
                pass

    if ext == "doc":
        tmp = tempfile.NamedTemporaryFile(suffix=".doc", delete=False)
        try:
            tmp.write(content_bytes)
            tmp.close()
            return _fallback_doc_binary(Path(tmp.name))
        finally:
            try:
                Path(tmp.name).unlink(missing_ok=True)
            except Exception:
                pass

    return None


def _fallback_docx(file_path: Path) -> str:
    """DOCX → 纯文本（使用 python-docx）

    处理策略：
      - 段落：提取纯文本
      - 表格：转为制表符分隔文本
      - 图片/嵌入对象：替换为 [图片] 占位符
    """
    try:
        import docx
        doc = docx.Document(str(file_path))
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            # 检测内联图片
            has_image = False
            try:
                for run in para.runs:
                    drawing_tags = run._element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                    )
                    pict_tags = run._element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"
                    )
                    if drawing_tags or pict_tags:
                        has_image = True
                        break
            except Exception:
                pass

            if text:
                parts.append(text)
            elif has_image:
                parts.append("[图片]")

        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)
    except ImportError:
        return None
    except Exception as e:
        logger.warning(f"DOCX 降级提取失败: {e}")
        return None


def _fallback_doc_binary(file_path: Path) -> str | None:
    """
    DOC (Word 97-2003) → 纯文本。
    尝试从 OLE2 复合文档中提取 WordDocument 流中的可读文本。
    这是一个基础的降级方案，无法处理复杂格式，但可提取大部分纯文本内容。
    """
    try:
        import olefile
    except ImportError:
        # olefile 未安装，尝试粗糙提取
        return _fallback_doc_raw_extract(file_path)

    try:
        if not olefile.isOleFile(str(file_path)):
            return _fallback_doc_raw_extract(file_path)

        ole = olefile.OleFileIO(str(file_path))
        try:
            # Word 文档的主文本流
            if ole.exists("WordDocument"):
                # 尝试读取 Word Document 流并提取可读文本
                stream = ole.openstream("WordDocument")
                data = stream.read()
                # 提取 Unicode 文本片段
                text = _extract_text_from_binary(data)
                if text and len(text.strip()) > 10:
                    return text

            # 备选：尝试所有流
            all_text_parts: list[str] = []
            for stream_path in ole.listdir():
                try:
                    stream = ole.openstream(stream_path)
                    data = stream.read()
                    text = _extract_text_from_binary(data)
                    if text and len(text.strip()) > 5:
                        all_text_parts.append(text)
                except Exception:
                    continue

            if all_text_parts:
                return "\n".join(all_text_parts)
        finally:
            ole.close()
    except Exception as e:
        logger.warning(f"DOC OLE 提取失败: {e}")

    return _fallback_doc_raw_extract(file_path)


def _extract_text_from_binary(data: bytes) -> str:
    """从二进制数据中提取可读文本片段"""
    parts: list[str] = []
    # 尝试 UTF-16-LE 解码（Word 内部编码）
    try:
        text = data.decode("utf-16-le", errors="ignore")
        # 过滤掉不可打印字符，保留中文、英文、数字、标点
        cleaned = re.sub(r"[^\u4e00-\u9fff\u3000-\u303f\uff00-\uffefA-Za-z0-9\s.,;:!?()\[\]{}'\"\-+=/\\@#$%^&*~`。\uff0c\uff1b\uff1a\uff01\uff1f\u2018\u2019\u201c\u201d\u3001\u300a\u300b\u3010\u3011]+", " ", text)
        # 合并多余空格
        cleaned = re.sub(r"\s{3,}", "\n", cleaned).strip()
        if len(cleaned) > 20:
            parts.append(cleaned)
    except Exception:
        pass

    # 尝试 GBK 解码
    try:
        text = data.decode("gbk", errors="ignore")
        cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        # 只保留足够长的文本片段
        segments = [s.strip() for s in re.split(r"\s{5,}", cleaned) if len(s.strip()) > 10]
        if segments and not parts:
            parts.extend(segments)
    except Exception:
        pass

    return "\n".join(parts)


def _fallback_doc_raw_extract(file_path: Path) -> str | None:
    """最后兜底：从 .doc 文件中求年提取可读文本"""
    try:
        data = file_path.read_bytes()
        text = _extract_text_from_binary(data)
        if text and len(text.strip()) > 20:
            return text
    except Exception as e:
        logger.warning(f"DOC 原始提取失败: {e}")
    return None


def _fallback_csv(file_path: Path) -> str:
    """CSV → 纯文本"""
    import csv as csv_mod
    text = _read_text_safe(file_path)
    reader = csv_mod.reader(text.splitlines())
    rows = list(reader)
    if not rows:
        return ""
    return "\n".join("\t".join(row) for row in rows)
