"""公文管理路由"""

import asyncio
import csv
import io
import json
import logging
import zipfile
from datetime import datetime, date, timezone
from pathlib import Path
from uuid import UUID

from fastapi import APIRouter, Depends, Query, Request, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel
from sqlalchemy import select, func, or_
from sqlalchemy.exc import IntegrityError as SAIntegrityError
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.database import get_db
from app.core.response import success, error, ErrorCode
from app.core.deps import require_permission, get_current_user
from app.core.audit import log_action
from app.services.usage_recorder import record_usage

logger = logging.getLogger(__name__)
from app.models.user import User
from app.models.document import Document, DocumentVersion
from app.models.knowledge import KBCollection
from app.schemas.document import (
    DocumentCreateRequest, DocumentUpdateRequest, DocProcessRequest,
    DocumentListItem, DocumentDetail, DocumentVersionItem, DocumentVersionDetail,
    DocumentExportRequest,
)
from app.services.dify.factory import get_dify_service
from app.services.docformat.service import DocFormatService
from app.services.doc_converter import (
    convert_bytes_to_markdown,
    convert_and_extract,
    convert_to_pdf_bytes,
    save_markdown_file,
    DOC_IMPORT_EXTENSIONS,
)
from app.core.redis import get_redis
from app.core.database import AsyncSessionLocal

router = APIRouter(prefix="/documents", tags=["Documents"])


async def _safe_update_doc(
    doc_id: UUID,
    updates: dict | None = None,
    save_version_before: bool = False,
    version_user_id: UUID | None = None,
    version_change_type: str = "",
    version_change_summary: str = "",
) -> str:
    """在独立数据库会话中持久化文档更新（SSE 断连安全）。

    SSE 生成器使用依赖注入的 db 会话，客户端断连时该会话可能回滚，
    此函数使用独立会话确保关键写入一定落库。

    Args:
        doc_id: 文档 UUID
        updates: 要更新的字段字典，如 {"content": "...", "status": "draft"}
        save_version_before: 是否在更新前保存版本快照
        version_user_id: 版本创建者 ID
        version_change_type: 版本变更类型
        version_change_summary: 版本变更说明

    Returns:
        更新后的 doc.content（供后续 SSE 事件使用）
    """
    async with AsyncSessionLocal() as s:
        async with s.begin():
            result = await s.execute(
                select(Document).where(Document.id == doc_id)
            )
            doc = result.scalar_one()

            # 可选：在更新前保存版本快照（带重试，防止并发版本号冲突）
            if save_version_before and doc.content:
                for _ver_attempt in range(3):
                    ver_result = await s.execute(
                        select(func.max(DocumentVersion.version_number))
                        .where(DocumentVersion.document_id == doc.id)
                    )
                    max_ver = ver_result.scalar() or 0
                    ver = DocumentVersion(
                        document_id=doc.id,
                        version_number=max_ver + 1,
                        content=doc.content,
                        formatted_paragraphs=doc.formatted_paragraphs,
                        change_type=version_change_type,
                        change_summary=version_change_summary,
                        created_by=version_user_id,
                    )
                    try:
                        async with s.begin_nested():
                            s.add(ver)
                            await s.flush()
                        break  # 成功，退出重试
                    except SAIntegrityError:
                        logger.warning(f"版本号冲突 (attempt {_ver_attempt+1}/3): doc={doc_id}, tried v{max_ver+1}")
                        try:
                            s.expunge(ver)
                        except Exception:
                            pass
                        continue

            if updates:
                for k, v in updates.items():
                    setattr(doc, k, v)

            return doc.content or ""


# MIME 映射
_MIME_MAP = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "doc": "application/msword",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "xls": "application/vnd.ms-excel",
    "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "csv": "text/csv",
    "txt": "text/plain",
    "md": "text/markdown",
    "html": "text/html",
    "htm": "text/html",
    "json": "application/json",
    "xml": "application/xml",
}


def _is_safe_upload_path(file_path: str | Path) -> bool:
    """校验文件路径在 UPLOAD_DIR 范围内，防止路径遍历攻击。"""
    try:
        resolved = Path(file_path).resolve()
        upload_root = Path(settings.UPLOAD_DIR).resolve()
        return resolved.is_relative_to(upload_root)
    except (ValueError, OSError):
        return False


@router.get("")
async def list_documents(
    category: str = Query(..., description="doc 或 template"),
    scope: str = Query("mine", description="mine=我的公文箱, public=公开公文箱"),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    keyword: str = Query(None),
    doc_type: str = Query(None),
    status: str = Query(None),
    security: str = Query(None),
    start_date: str = Query(None),
    end_date: str = Query(None),
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """公文列表"""
    query = select(Document).where(Document.category == category)

    # 按 scope 过滤
    if scope == "public":
        query = query.where(Document.visibility == "public")
    else:
        query = query.where(Document.creator_id == current_user.id)

    if keyword:
        query = query.where(Document.title.ilike(f"%{keyword}%"))
    if doc_type:
        query = query.where(Document.doc_type == doc_type)
    if status:
        query = query.where(Document.status == status)
    if security:
        query = query.where(Document.security == security)
    if start_date:
        query = query.where(Document.updated_at >= datetime.strptime(start_date, "%Y-%m-%d"))
    if end_date:
        query = query.where(Document.updated_at <= datetime.strptime(end_date + " 23:59:59", "%Y-%m-%d %H:%M:%S"))

    count_query = select(func.count()).select_from(query.subquery())
    total = (await db.execute(count_query)).scalar() or 0

    query = query.order_by(Document.updated_at.desc()).offset((page - 1) * page_size).limit(page_size)
    result = await db.execute(query)
    docs = result.scalars().all()

    # 批量查创建者姓名
    creator_ids = {d.creator_id for d in docs}
    creator_map = {}
    if creator_ids:
        from app.models.user import User as U
        cr = await db.execute(select(U.id, U.display_name).where(U.id.in_(creator_ids)))
        creator_map = {row[0]: row[1] for row in cr.all()}

    items = [
        {
            **DocumentListItem.model_validate(d).model_dump(mode="json"),
            "creator_name": creator_map.get(d.creator_id, ""),
        }
        for d in docs
    ]

    return success(data={"items": items, "total": total, "page": page, "page_size": page_size})


@router.post("")
async def create_document(
    body: DocumentCreateRequest,
    request: Request,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """创建公文"""
    doc = Document(
        creator_id=current_user.id,
        title=body.title,
        category=body.category,
        doc_type=body.doc_type,
        content=body.content,
        urgency=body.urgency,
        security=body.security,
    )
    db.add(doc)
    await db.flush()

    await log_action(
        db, user_id=current_user.id, user_display_name=current_user.display_name,
        action="创建公文", module="智能公文",
        detail=f"创建{'公文' if body.category == 'doc' else '模板'}: {body.title}",
        ip_address=request.client.host if request.client else None,
    )

    return success(data={"id": str(doc.id)}, message="创建成功")


@router.post("/import")
async def import_document(
    request: Request,
    file: UploadFile = File(None, description="支持 PDF/Word/Excel/PPT/TXT/HTML 等格式，可为空"),
    category: str = Form("doc"),
    doc_type: str = Form("report"),
    security: str = Form("internal"),
    title: str = Form("", description="文档标题（不上传文件时使用）"),
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """
    导入文档为公文/模板。

    通过 converter 微服务将各类文档（PDF, DOCX, DOC, XLSX, CSV, TXT, MD, PPTX, HTML 等）
    提取文本内容，并生成 PDF 预览缓存。文件参数可为空，此时创建空白文档。
    """
    # ── 校验枚举参数 ──
    VALID_CATEGORIES = {"doc", "template"}
    VALID_DOC_TYPES = {
        "request", "report", "notice", "briefing", "ai_generated",
        "official", "academic", "legal", "proposal", "lab_fund",
        "school_notice_redhead", "custom",
    }
    VALID_SECURITIES = {"public", "internal", "secret", "confidential"}

    if category not in VALID_CATEGORIES:
        return error(ErrorCode.PARAM_INVALID, f"category 必须为 {VALID_CATEGORIES} 之一，收到: '{category}'")
    if doc_type not in VALID_DOC_TYPES:
        return error(ErrorCode.PARAM_INVALID, f"doc_type 必须为 {VALID_DOC_TYPES} 之一，收到: '{doc_type}'")
    if security not in VALID_SECURITIES:
        return error(ErrorCode.PARAM_INVALID, f"security 必须为 {VALID_SECURITIES} 之一，收到: '{security}'")

    # ── 判断是否有文件上传 ──
    has_file = file is not None and file.filename
    content_bytes = b""
    content = ""
    file_name = ""
    ext = ""
    char_count = 0
    convert_result = None

    if has_file:
        file_name = file.filename or "unknown.docx"
        ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""

        if ext not in DOC_IMPORT_EXTENSIONS:
            supported = ", ".join(sorted(DOC_IMPORT_EXTENSIONS))
            return error(ErrorCode.PARAM_INVALID, f"不支持的文件格式 .{ext}，支持: {supported}")

        content_bytes = await file.read()

        # 文件大小限制: 50MB
        MAX_UPLOAD_SIZE = 50 * 1024 * 1024
        if len(content_bytes) > MAX_UPLOAD_SIZE:
            return error(
                ErrorCode.PARAM_INVALID,
                f"文件大小 ({len(content_bytes) / 1024 / 1024:.1f}MB) 超过限制，最大允许 {MAX_UPLOAD_SIZE // 1024 // 1024}MB",
            )

        if content_bytes:
            # ── 使用 converter 微服务提取文本 + 生成 PDF ──
            convert_result = await convert_and_extract(content_bytes, file_name)

            if not convert_result.success:
                return error(
                    ErrorCode.FILE_UPLOAD_ERROR,
                    f"文档解析失败: {convert_result.error_message}",
                )

            content = convert_result.markdown or ""
            char_count = convert_result.char_count

            # 安全净化：移除 null 字节，避免 PostgreSQL 存储错误
            if content:
                content = content.replace("\x00", "")

    # 提取标题：优先用传入的 title 参数，其次用文件名，最后用默认值
    doc_title = title.strip() if title and title.strip() else (
        file_name.rsplit(".", 1)[0] if file_name and "." in file_name else (
            file_name or "新建公文"
        )
    )

    # ── 创建文档记录（先 flush 获取 ID） ──
    doc = Document(
        creator_id=current_user.id,
        title=doc_title,
        category=category,
        doc_type=doc_type,
        content=content,
        urgency="normal",
        security=security,
        source_format=ext or "txt",
    )
    db.add(doc)
    await db.flush()

    # ── 持久化源文件到磁盘（仅在有文件时） ──
    upload_dir = Path(settings.UPLOAD_DIR) / "documents" / str(doc.id)
    upload_dir.mkdir(parents=True, exist_ok=True)

    if content_bytes:
        source_path = upload_dir / f"source.{ext}"
        source_path.write_bytes(content_bytes)
        doc.source_file_path = str(source_path)

    # ── 如果 converter 返回了 PDF 路径，复制为预览缓存 ──
    if convert_result and convert_result.pdf_path:
        try:
            import shutil
            pdf_src = Path(convert_result.pdf_path)
            if pdf_src.exists():
                pdf_cache = upload_dir / "preview.pdf"
                shutil.copy2(pdf_src, pdf_cache)
        except Exception as e:
            logging.getLogger(__name__).warning(f"PDF 缓存复制失败: {e}")

    # ── 持久化提取的文本到磁盘 ──
    if content:
        md_path = await save_markdown_file(content, upload_dir, "content")
        doc.md_file_path = str(md_path)
    await db.flush()

    action_detail = (
        f"导入文件: {file_name} → {doc_title} (格式: {ext}, 字符数: {char_count})"
        if has_file else f"创建空白文档: {doc_title}"
    )
    await log_action(
        db, user_id=current_user.id, user_display_name=current_user.display_name,
        action="导入公文", module="智能公文",
        detail=action_detail,
        ip_address=request.client.host if request.client else None,
    )

    return success(
        data={
            "id": str(doc.id),
            "title": doc_title,
            "format": ext or "txt",
            "char_count": char_count,
            "has_source_file": bool(content_bytes),
            "has_markdown_file": bool(content),
        },
        message="导入成功" if has_file else "空白文档创建成功",
    )


@router.post("/export")
async def export_documents(
    body: DocumentExportRequest,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """导出公文为 ZIP 压缩包（优先导出优化后的 DOCX，回退到原始文件或文本内容）"""
    # 构建查询
    if body.ids:
        query = select(Document).where(Document.id.in_(body.ids))
    else:
        query = select(Document).where(Document.category == "doc")

    query = query.order_by(Document.updated_at.desc()).limit(5000)
    result = await db.execute(query)
    docs = result.scalars().all()

    if not docs:
        return error(ErrorCode.PARAM_INVALID, "没有可导出的文档")

    # 创建 ZIP 压缩包
    buf = io.BytesIO()
    seen_names: set[str] = set()

    def _unique_name(name: str) -> str:
        """生成不重复的文件名"""
        if name not in seen_names:
            seen_names.add(name)
            return name
        base, dot_ext = (name.rsplit(".", 1) if "." in name else (name, ""))
        counter = 1
        while True:
            candidate = f"{base}_{counter}.{dot_ext}" if dot_ext else f"{base}_{counter}"
            if candidate not in seen_names:
                seen_names.add(candidate)
                return candidate
            counter += 1

    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for d in docs:
            base_name = d.title or "未命名"

            # ① 优先：有结构化排版数据 → 生成格式化 DOCX
            if d.formatted_paragraphs:
                try:
                    paragraphs = json.loads(d.formatted_paragraphs) if isinstance(d.formatted_paragraphs, str) else d.formatted_paragraphs
                    if isinstance(paragraphs, list) and len(paragraphs) > 0:
                        import asyncio as _aio_export
                        loop = _aio_export.get_event_loop()
                        docx_buf = await loop.run_in_executor(None, _build_formatted_docx, paragraphs, base_name)
                        file_name = _unique_name(f"{base_name}.docx")
                        zf.writestr(file_name, docx_buf.getvalue())
                        continue
                except Exception as e:
                    logging.getLogger(__name__).warning(f"生成格式化 DOCX 失败 [{base_name}]: {e}")

            # ② 次选：有文本内容 → 保存为 .md
            if d.content:
                file_name = _unique_name(f"{base_name}.md")
                zf.writestr(file_name, d.content)
            # ③ 兜底：有原始文件 → 添加原始文件
            elif d.source_file_path and _is_safe_upload_path(d.source_file_path) and Path(d.source_file_path).exists():
                ext = d.source_format or "txt"
                file_name = _unique_name(f"{base_name}.{ext}")
                zf.write(d.source_file_path, file_name)

    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=documents_export.zip"},
    )


# ── 结构化排版导出 DOCX ──

# 字号名 → pt 映射
_FONT_SIZE_PT = {
    "初号": 42, "小初": 36, "一号": 26, "小一": 24,
    "二号": 22, "小二": 18, "三号": 16, "小三": 15,
    "四号": 14, "小四": 12, "五号": 10.5, "小五": 9,
}

# ── style_type 归一化（与前端 StructuredDocRenderer 保持一致）──
_VALID_STYLE_TYPES = {
    "title", "subtitle", "heading1", "heading2", "heading3", "heading4",
    "body", "recipient", "signature", "date", "attachment", "closing",
}

def _normalize_style_type(raw: str | None) -> str:
    if not raw:
        return "body"
    t = raw.strip().lower()
    if t in _VALID_STYLE_TYPES:
        return t
    if "title" in t or t == "标题":
        return "title"
    if "heading" in t and "1" in t or "一级" in t:
        return "heading1"
    if "heading" in t and "2" in t or "二级" in t:
        return "heading2"
    if "heading" in t and "3" in t or "三级" in t:
        return "heading3"
    if "heading" in t and "4" in t or "四级" in t:
        return "heading4"
    if "body" in t or t == "正文":
        return "body"
    if "signature" in t or "落款" in t or "署名" in t:
        return "signature"
    if "date" in t or t == "日期":
        return "date"
    if "recipient" in t or "主送" in t:
        return "recipient"
    if "attachment" in t or "附件" in t:
        return "attachment"
    if "subtitle" in t or "副标题" in t or "子标题" in t:
        return "subtitle"
    if "closing" in t or "结束" in t:
        return "closing"
    return "body"


# ════════════════════════════════════════════════════════════
# GB/T 9704 公文样式预设（与前端 StructuredDocRenderer.STYLE_PRESETS 完全对齐）
# 每种 preset 下按 style_type 定义默认：font_family, font_size_pt, alignment,
# indent_em, line_height, bold, space_before_pt, space_after_pt
# ════════════════════════════════════════════════════════════
_STYLE_PRESETS: dict[str, dict[str, dict]] = {
    "official": {
        # space_before / space_after 单位 pt，与前端 em 换算：1em ≈ 当前字号(pt)
        # title: 首段，space_before=0；space_after 由红线逻辑覆盖
        "title":      {"font_family": "方正小标宋简体", "font_size_pt": 22, "alignment": "center", "indent_em": 0, "line_height": 2.0, "bold": False, "space_before_pt": 0,  "space_after_pt": 8},
        # recipient: 前端 marginTop 0.8em=12.8pt（由 getSpacingTop 动态规则覆盖）
        "recipient":  {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "left",   "indent_em": 0, "line_height": 2.0, "bold": False, "space_before_pt": 0,  "space_after_pt": 0},
        # heading1: 前端入场 1em=16pt, heading→heading 0.4em=6pt
        "heading1":   {"font_family": "黑体",           "font_size_pt": 16, "alignment": "left",   "indent_em": 2, "line_height": 2.0, "bold": False, "space_before_pt": 14, "space_after_pt": 2},
        "heading2":   {"font_family": "楷体_GB2312",    "font_size_pt": 16, "alignment": "left",   "indent_em": 2, "line_height": 2.0, "bold": False, "space_before_pt": 14, "space_after_pt": 2},
        "heading3":   {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "left",   "indent_em": 2, "line_height": 2.0, "bold": False, "space_before_pt": 8,  "space_after_pt": 2},
        "heading4":   {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "left",   "indent_em": 2, "line_height": 2.0, "bold": False, "space_before_pt": 8,  "space_after_pt": 2},
        "body":       {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "justify","indent_em": 2, "line_height": 2.0, "bold": False, "space_before_pt": 0,  "space_after_pt": 0},
        # signature: 前端入场 1.5em=24pt
        "signature":  {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "right",  "indent_em": 0, "line_height": 2.0, "bold": False, "space_before_pt": 20, "space_after_pt": 0},
        "date":       {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "right",  "indent_em": 0, "line_height": 2.0, "bold": False, "space_before_pt": 0,  "space_after_pt": 0},
        # attachment: 前端入场 1.2em=19.2pt
        "attachment": {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "justify","indent_em": 2, "line_height": 2.0, "bold": False, "space_before_pt": 16, "space_after_pt": 0},
        "closing":    {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "left",   "indent_em": 2, "line_height": 2.0, "bold": False, "space_before_pt": 0,  "space_after_pt": 0},
    },
    "school_notice_redhead": {
        # 红头：方正小标宋简体 32pt 红色；标题：方正小标宋简体 二号(22pt)；正文：仿宋_GB2312 三号(16pt)；行距固定28.95磅
        "title":      {"font_family": "方正小标宋简体", "font_size_pt": 32, "alignment": "center", "indent_em": 0, "line_height": 1.0, "bold": False, "space_before_pt": 0,  "space_after_pt": 10, "exact_line_spacing_pt": 40},
        "subtitle":   {"font_family": "方正小标宋简体", "font_size_pt": 22, "alignment": "center", "indent_em": 0, "line_height": 1.316, "bold": False, "space_before_pt": 4,  "space_after_pt": 10, "exact_line_spacing_pt": 28.95},
        "recipient":  {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "left",   "indent_em": 0, "line_height": 1.81, "bold": False, "space_before_pt": 8,  "space_after_pt": 0, "exact_line_spacing_pt": 28.95},
        "heading1":   {"font_family": "黑体",           "font_size_pt": 16, "alignment": "left",   "indent_em": 2, "line_height": 1.81, "bold": False, "space_before_pt": 12, "space_after_pt": 2, "exact_line_spacing_pt": 28.95},
        "heading2":   {"font_family": "楷体_GB2312",    "font_size_pt": 16, "alignment": "left",   "indent_em": 2, "line_height": 1.81, "bold": False, "space_before_pt": 10, "space_after_pt": 2, "exact_line_spacing_pt": 28.95},
        "heading3":   {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "left",   "indent_em": 2, "line_height": 1.81, "bold": False, "space_before_pt": 8,  "space_after_pt": 2, "exact_line_spacing_pt": 28.95},
        "heading4":   {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "left",   "indent_em": 2, "line_height": 1.81, "bold": False, "space_before_pt": 8,  "space_after_pt": 2, "exact_line_spacing_pt": 28.95},
        "body":       {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "justify", "indent_em": 2, "line_height": 1.81, "bold": False, "space_before_pt": 0,  "space_after_pt": 0, "exact_line_spacing_pt": 28.95},
        "closing":    {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "left",   "indent_em": 2, "line_height": 1.81, "bold": False, "space_before_pt": 0,  "space_after_pt": 0, "exact_line_spacing_pt": 28.95},
        "signature":  {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "right",  "indent_em": 0, "line_height": 1.81, "bold": False, "space_before_pt": 18, "space_after_pt": 0, "exact_line_spacing_pt": 28.95},
        "date":       {"font_family": "仿宋_GB2312",    "font_size_pt": 16, "alignment": "right",  "indent_em": 0, "line_height": 1.81, "bold": False, "space_before_pt": 0,  "space_after_pt": 0, "exact_line_spacing_pt": 28.95},
        "attachment": {"font_family": "仿宋_GB2312",    "font_size_pt": 14, "alignment": "left",   "indent_em": 0, "line_height": 1.5, "bold": False, "space_before_pt": 14, "space_after_pt": 0},
    },
    "academic": {
        "title":    {"font_family": "黑体",        "font_size_pt": 18, "alignment": "center", "indent_em": 0, "line_height": 1.8, "bold": True,  "space_before_pt": 20, "space_after_pt": 20},
        "heading1": {"font_family": "黑体",        "font_size_pt": 15, "alignment": "left",   "indent_em": 0, "line_height": 1.8, "bold": True,  "space_before_pt": 10, "space_after_pt": 6},
        "heading2": {"font_family": "黑体",        "font_size_pt": 14, "alignment": "left",   "indent_em": 0, "line_height": 1.8, "bold": True,  "space_before_pt": 8,  "space_after_pt": 4},
        "heading3": {"font_family": "楷体_GB2312", "font_size_pt": 14, "alignment": "left",   "indent_em": 0, "line_height": 1.8, "bold": True,  "space_before_pt": 6,  "space_after_pt": 4},
        "body":     {"font_family": "仿宋_GB2312", "font_size_pt": 12, "alignment": "justify","indent_em": 2, "line_height": 1.8, "bold": False, "space_before_pt": 0,  "space_after_pt": 0},
        "signature":{"font_family": "仿宋_GB2312", "font_size_pt": 12, "alignment": "right",  "indent_em": 0, "line_height": 1.8, "bold": False, "space_before_pt": 12, "space_after_pt": 0},
    },
    "legal": {
        "title":    {"font_family": "方正小标宋简体","font_size_pt": 26, "alignment": "center", "indent_em": 0, "line_height": 2.2, "bold": False, "space_before_pt": 20, "space_after_pt": 20},
        "heading1": {"font_family": "黑体",        "font_size_pt": 16, "alignment": "left",   "indent_em": 2, "line_height": 2.0, "bold": False, "space_before_pt": 10, "space_after_pt": 6},
        "body":     {"font_family": "仿宋_GB2312", "font_size_pt": 16, "alignment": "justify","indent_em": 2, "line_height": 2.0, "bold": False, "space_before_pt": 0,  "space_after_pt": 0},
        "signature":{"font_family": "仿宋_GB2312", "font_size_pt": 16, "alignment": "right",  "indent_em": 0, "line_height": 2.0, "bold": False, "space_before_pt": 12, "space_after_pt": 0},
        "date":     {"font_family": "仿宋_GB2312", "font_size_pt": 16, "alignment": "right",  "indent_em": 0, "line_height": 2.0, "bold": False, "space_before_pt": 0,  "space_after_pt": 0},
    },
}

# alignment 归一化
_ALIGN_ALIAS = {
    "居中": "center", "居右": "right", "居左": "left",
    "右对齐": "right", "左对齐": "left", "两端对齐": "justify", "两端": "justify",
}


def _resolve_font_size_pt(raw: str | None) -> float | None:
    """将 LLM 的 font_size 字符串解析为 pt 数值"""
    if not raw:
        return None
    t = raw.strip()
    if t in _FONT_SIZE_PT:
        return _FONT_SIZE_PT[t]
    # "16pt"
    import re
    m = re.match(r'^([\d.]+)\s*pt$', t, re.IGNORECASE)
    if m:
        return float(m.group(1))
    # 纯数字 → pt
    if re.match(r'^[\d.]+$', t):
        return float(t)
    return None


def _resolve_alignment(raw: str | None) -> str | None:
    if not raw:
        return None
    t = raw.strip().lower()
    if t in ("center", "right", "left", "justify"):
        return t
    return _ALIGN_ALIAS.get(t)


def _resolve_indent_em(raw: str | None) -> float | None:
    """解析首行缩进值（em），返回 None 表示未指定"""
    if raw is None:
        return None
    t = str(raw).strip()
    if t in ("", "none", "无"):
        return 0.0
    if t == "0":
        return 0.0
    import re
    m = re.match(r'^([\d.]+)\s*(em)?$', t)
    if m:
        return float(m.group(1))
    return None


def _resolve_line_height(raw: str | None) -> float | None:
    """解析行高为倍数值"""
    if not raw:
        return None
    t = str(raw).strip()
    import re
    m = re.match(r'^([\d.]+)\s*pt$', t, re.IGNORECASE)
    if m:
        return None  # pt 行高需要另行处理
    if re.match(r'^[\d.]+$', t):
        val = float(t)
        if val <= 5:
            return val  # 倍数
    return None


# ── Markdown 预处理 & 排版模板 ──────────────────────────

import re as _re

# ── 模块级正则常量（文档结构识别，多处复用） ──────────────
_RE_HEADING1 = _re.compile(r'^[一二三四五六七八九十百]+[、．.]')
_RE_HEADING2 = _re.compile(r'^[\(（][一二三四五六七八九十]+[\)）]')
_RE_HEADING3 = _re.compile(r'^\d{1,2}[\.\、](?!\d)')
_RE_HEADING4 = _re.compile(r'^[\(（]\d{1,2}[\)）]')
_RE_TITLE = _re.compile(r'^关于.{2,40}的(通知|报告|请示|批复|函|纪要|意见|决定|方案|办法|规定|计划|总结)')
_RE_RECIPIENT = _re.compile(r'^.{2,30}[：:]$')
_RE_CLOSING = _re.compile(r'^(特此(通知|报告|函复|批复)|以上(报告|意见).*[请审]|妥否.*请[批示审]|此复|当否)')
_RE_DATE = _re.compile(r'^\d{4}年\d{1,2}月\d{1,2}日$|^20\d{2}[./\-]\d{1,2}[./\-]\d{1,2}$')
_RE_ATTACHMENT = _re.compile(r'^附[件：:]')
_RE_SIGNATURE_SHORT = _re.compile(r'^.{2,25}$')  # 尾部短行辅助判定署名


def _strip_markdown_for_format(text: str) -> str:
    """
    Strip Markdown formatting symbols while preserving text content.
    Preprocesses document text before sending to the format LLM,
    ensuring #, *, > etc. don't end up in formatted output.
    """
    lines = text.split('\n')
    result: list[str] = []
    for line in lines:
        s = line.rstrip()
        if not s.strip():
            result.append('')
            continue
        # Horizontal rules: ---, ***, ___
        if _re.match(r'^\s*[-*_]{3,}\s*$', s):
            continue
        # Headings: # ## ### etc.
        s = _re.sub(r'^(\s*)#{1,6}\s+', r'\1', s)
        s = _re.sub(r'\s*#{1,6}\s*$', '', s)  # trailing ###
        # Bold+italic: ***text*** / ___text___
        s = _re.sub(r'\*{3}(.+?)\*{3}', r'\1', s)
        s = _re.sub(r'_{3}(.+?)_{3}', r'\1', s)
        # Bold: **text** / __text__
        s = _re.sub(r'\*{2}(.+?)\*{2}', r'\1', s)
        s = _re.sub(r'_{2}(.+?)_{2}', r'\1', s)
        # Italic: *text* / _text_ (avoid matching list markers)
        s = _re.sub(r'(?<!\w)\*([^*]+)\*(?!\w)', r'\1', s)
        # Strikethrough: ~~text~~
        s = _re.sub(r'~~(.+?)~~', r'\1', s)
        # Blockquotes: > text
        s = _re.sub(r'^(\s*)>\s*', r'\1', s)
        # Unordered list markers: - item, * item, + item
        s = _re.sub(r'^(\s*)[-*+]\s+', r'\1', s)
        # Ordered list markers: 1. item, 2) item
        s = _re.sub(r'^(\s*)\d+[.)\uff0e]\s+', r'\1', s)
        # Inline code: `code`
        s = _re.sub(r'`([^`]+)`', r'\1', s)
        # Code block fences: ``` or ~~~
        if _re.match(r'^\s*(`{3}|~{3})', s):
            continue
        # Links: [text](url) → text
        s = _re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', s)
        # Images: ![alt](url) → alt
        s = _re.sub(r'!\[([^\]]*)\]\([^)]*\)', r'\1', s)
        # HTML tags
        s = _re.sub(r'</?[a-zA-Z][^>]*>', '', s)
        if s.strip():
            result.append(s)
    cleaned = '\n'.join(result)
    cleaned = _re.sub(r'\n{3,}', '\n\n', cleaned)
    return cleaned.strip()


def _strip_markdown_inline(text: str) -> str:
    """Strip residual inline markdown from a single paragraph text."""
    s = text
    s = _re.sub(r'^#{1,6}\s+', '', s)
    s = _re.sub(r'\*{2,3}(.+?)\*{2,3}', r'\1', s)
    s = _re.sub(r'_{2,3}(.+?)_{2,3}', r'\1', s)
    s = _re.sub(r'(?<!\w)\*([^*]+)\*(?!\w)', r'\1', s)
    s = _re.sub(r'^>\s*', '', s)
    s = _re.sub(r'^[-*+]\s+', '', s)
    s = _re.sub(r'`([^`]+)`', r'\1', s)
    s = _re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', s)
    return s.strip()


# ── Markdown → 结构化段落解析 ──────────────────────────

def _detect_line_style(stripped: str, idx: int, total: int,
                       has_title: bool, has_closing: bool, has_signature: bool) -> str:
    """
    对单行文本识别 style_type，复用模块级正则常量。
    返回 style_type 字符串。
    """
    if not has_title and (_RE_TITLE.match(stripped) or (idx == 0 and len(stripped) < 60)):
        return "title"
    if _RE_HEADING1.match(stripped):
        return "heading1"
    if _RE_HEADING2.match(stripped):
        return "heading2"
    if _RE_HEADING3.match(stripped):
        return "heading3"
    if _RE_HEADING4.match(stripped):
        return "heading4"
    if _RE_RECIPIENT.match(stripped) and idx <= 3:
        return "recipient"
    if _RE_CLOSING.match(stripped):
        return "closing"
    if _RE_DATE.match(stripped):
        return "date"
    if _RE_ATTACHMENT.match(stripped):
        return "attachment"
    # 尾部短行可能是署名
    if idx >= total - 3 and len(stripped) < 30 and not _RE_HEADING1.match(stripped):
        if not _RE_DATE.match(stripped) and not has_signature:
            return "signature"
    return "body"


def _parse_markdown_to_paragraphs(text: str) -> list[dict]:
    """
    将 Markdown 纯文本解析为结构化段落列表。

    逐行匹配模块级正则常量，为每行分配 style_type + confidence；
    同时清除残留 Markdown 符号（#, *, > 等）。
    连续 body 行（中间无空行）会合并为一个段落，避免行距异常。

    使用 _detect_style_with_confidence() 替代 _detect_line_style()，
    输出含 confidence 信息，后续 _rules_format_paragraphs() 可跳过重复检测。

    Returns:
        list[dict]: 每项含 {"text": str, "style_type": str, "_confidence": float}
    """
    # 按空行分组，保留段落边界信息
    raw_lines = text.split('\n')
    groups: list[list[str]] = []
    current_group: list[str] = []
    for line in raw_lines:
        if not line.strip():
            if current_group:
                groups.append(current_group)
                current_group = []
        else:
            current_group.append(line)
    if current_group:
        groups.append(current_group)

    # 统计总行数（不含空行），用于位置相关的 style 检测
    total = sum(len(g) for g in groups)
    result: list[dict] = []
    has_title = False
    has_closing = False
    has_signature = False
    global_idx = 0

    for group in groups:
        pending_body: list[str] = []
        pending_body_min_conf: float = 1.0  # 累积 body 段的最小置信度
        for line in group:
            stripped = _strip_markdown_inline(line.strip())
            if not stripped:
                global_idx += 1
                continue
            style, confidence = _detect_style_with_confidence(
                stripped, global_idx, total, has_title, has_closing, has_signature,
            )
            if style == "title":
                has_title = True
            elif style == "closing":
                has_closing = True
            elif style == "signature":
                has_signature = True

            if style == "body":
                pending_body.append(stripped)
                pending_body_min_conf = min(pending_body_min_conf, confidence)
            else:
                # 遇到非 body 行，先 flush 已累积的连续 body 文本
                if pending_body:
                    result.append({"text": "".join(pending_body), "style_type": "body",
                                   "_confidence": pending_body_min_conf})
                    pending_body = []
                    pending_body_min_conf = 1.0
                result.append({"text": stripped, "style_type": style, "_confidence": confidence})
            global_idx += 1

        # 每组结束时 flush 剩余 body 文本
        if pending_body:
            result.append({"text": "".join(pending_body), "style_type": "body",
                           "_confidence": pending_body_min_conf})

    return result


# ── 行标记指令解析（增量模式） ──────────────────────────

_RE_LINE_CMD = _re.compile(
    r'^\[(REPLACE|ADD|DELETE|NEED_INFO)(?::([^\]]*))?\]\s*(.*)',
    _re.IGNORECASE,
)


def _parse_line_diff_commands(text: str) -> list[dict]:
    """
    解析增量模式 LLM 输出的行标记指令。

    行标记语法：
      [REPLACE:3|style:heading1] 新文本内容
      [ADD:after=5|style:body] 新段落
      [DELETE:7]
      [NEED_INFO] 提示文字

    Returns:
        list[dict]: 与 _apply_draft_diff() 兼容的变更指令列表
    """
    commands: list[dict] = []

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        m = _RE_LINE_CMD.match(line)
        if not m:
            continue
        op_raw = m.group(1).upper()
        params_str = m.group(2) or ""
        content = m.group(3).strip()

        # 解析参数 key=value 或 key:value 对
        params: dict[str, str] = {}
        if params_str:
            for part in params_str.split('|'):
                part = part.strip()
                if not part:
                    continue
                # 纯数字 → 当作 index
                if part.isdigit():
                    params["index"] = part
                    continue
                for sep in ('=', ':'):
                    if sep in part:
                        k, v = part.split(sep, 1)
                        params[k.strip().lower()] = v.strip()
                        break

        if op_raw == "REPLACE":
            idx = int(params.get("index", "-1"))
            cmd: dict = {"op": "replace", "index": idx, "text": _strip_markdown_inline(content)}
            if "style" in params:
                cmd["style_type"] = params["style"]
            commands.append(cmd)

        elif op_raw == "ADD":
            after_idx = int(params.get("after", params.get("index", "-1")))
            cmd = {"op": "add", "after": after_idx, "text": _strip_markdown_inline(content),
                   "style_type": params.get("style", "body")}
            commands.append(cmd)

        elif op_raw == "DELETE":
            idx = int(params.get("index", "-1"))
            if idx < 0 and params_str.strip().isdigit():
                idx = int(params_str.strip())
            commands.append({"op": "delete", "index": idx})

        elif op_raw == "NEED_INFO":
            commands.append({"op": "need_info", "text": content or "请提供更详细的指令。"})

    return commands


# ── JSON 截断检测（多轮续写用，保留兼容） ──────────────────────────

def _check_json_truncated(text: str) -> bool:
    """检测 JSON 输出是否因 token 限制而被截断（大括号/方括号不匹配）"""
    if not text or len(text) < 50:
        return False
    depth_brace = 0
    depth_bracket = 0
    in_str = False
    esc = False
    for c in text:
        if esc:
            esc = False
            continue
        if c == '\\' and in_str:
            esc = True
            continue
        if c == '"':
            in_str = not in_str
            continue
        if in_str:
            continue
        if c == '{':
            depth_brace += 1
        elif c == '}':
            depth_brace -= 1
        elif c == '[':
            depth_bracket += 1
        elif c == ']':
            depth_bracket -= 1
    return depth_brace > 0 or depth_bracket > 0


# ── 长文档分块排版 ──────────────────────────────────────

# 每块最大字符数，防止 LLM 输出 token 截断
# qwen3-32b (32k 上下文) 可处理更大块；减少分块数 → 减少风格不一致
_MAX_FORMAT_CHUNK_CHARS = 5000

# 增量模式分块默认参数（下方声明，此处提前引用）
_INCREMENTAL_MAX_PARAS_PER_CHUNK = 40   # 每块最多段落数
_INCREMENTAL_MAX_CHARS_PER_CHUNK = 4500  # 每块最多字符数


# ── 两阶段处理：Phase-1 规则化文档结构分析 ─────────────────

def _analyze_doc_structure(text: str) -> dict:
    """
    Phase-1：零延迟规则化文档结构分析。

    扫描全文（或段落列表提取的文本），识别每个段落的 style_type、
    文档整体类型、章节边界、编号体系，返回轻量大纲字典。

    返回：{
        "doc_type": "official" | "academic" | "legal" | "school_notice_redhead",
        "total_paragraphs": int,
        "sections": [{"heading": "一、总体要求", "style": "heading1", "para_range": [3, 8]}, ...],
        "outline": [{"idx": 0, "style": "title", "preview": "关于...通知"}, ...],
        "numbering": "一、→（一）→ 1. → (1)",
        "has_title": bool,
        "has_closing": bool,
        "has_signature": bool,
    }
    """
    lines = [l for l in text.split('\n') if l.strip()]
    outline: list[dict] = []
    sections: list[dict] = []
    current_section: dict | None = None
    numbering_levels: set[str] = set()

    # ── 结构识别正则（使用模块级常量） ──
    re_heading1 = _RE_HEADING1
    re_heading2 = _RE_HEADING2
    re_heading3 = _RE_HEADING3
    re_heading4 = _RE_HEADING4
    re_title = _RE_TITLE
    re_recipient = _RE_RECIPIENT
    re_closing = _RE_CLOSING
    re_date = _RE_DATE
    re_attachment = _RE_ATTACHMENT

    has_title = False
    has_closing = False
    has_signature = False

    for idx, line in enumerate(lines):
        stripped = line.strip()
        preview = stripped[:30]
        style = "body"

        if not has_title and (re_title.match(stripped) or (idx == 0 and len(stripped) < 60)):
            style = "title"
            has_title = True
        elif re_heading1.match(stripped):
            style = "heading1"
            numbering_levels.add("一、")
        elif re_heading2.match(stripped):
            style = "heading2"
            numbering_levels.add("（一）")
        elif re_heading3.match(stripped):
            style = "heading3"
            numbering_levels.add("1.")
        elif re_heading4.match(stripped):
            style = "heading4"
            numbering_levels.add("(1)")
        elif re_recipient.match(stripped) and idx <= 3:
            style = "recipient"
        elif re_closing.match(stripped):
            style = "closing"
            has_closing = True
        elif re_date.match(stripped):
            style = "date"
        elif re_attachment.match(stripped):
            style = "attachment"
        elif idx >= len(lines) - 3 and len(stripped) < 30 and not re_heading1.match(stripped):
            # 尾部短行可能是署名
            if not re_date.match(stripped) and not has_signature:
                style = "signature"
                has_signature = True

        outline.append({"idx": idx, "style": style, "preview": preview})

        # 章节追踪
        if style in ("heading1", "heading2"):
            if current_section:
                current_section["para_range"][1] = idx - 1
                sections.append(current_section)
            current_section = {
                "heading": stripped,
                "style": style,
                "para_range": [idx, idx],
            }
        elif current_section:
            current_section["para_range"][1] = idx

    if current_section:
        current_section["para_range"][1] = len(lines) - 1
        sections.append(current_section)

    # 编号体系
    level_order = ["一、", "（一）", "1.", "(1)"]
    numbering = " → ".join(l for l in level_order if l in numbering_levels) or "无明确编号"

    # 文档类型推断
    doc_type = "official"
    text_lower = text[:2000].lower()
    if any(kw in text_lower for kw in ("摘要", "abstract", "关键词", "keywords", "参考文献", "references")):
        doc_type = "academic"
    elif any(kw in text_lower for kw in ("原告", "被告", "判决", "裁定", "起诉", "法院")):
        doc_type = "legal"
    elif any(kw in text_lower for kw in ("大学", "学院", "学校", "承办单位", "联系人", "联系电话", "校办")):
        doc_type = "school_notice_redhead"

    return {
        "doc_type": doc_type,
        "total_paragraphs": len(lines),
        "sections": sections,
        "outline": outline,
        "numbering": numbering,
        "has_title": has_title,
        "has_closing": has_closing,
        "has_signature": has_signature,
    }


def _analyze_paragraphs_structure(paragraphs: list[dict]) -> dict:
    """对已有结构化段落列表做结构分析（复用 _analyze_doc_structure 逻辑）。"""
    text = "\n".join(p.get("text", "") for p in paragraphs)
    result = _analyze_doc_structure(text)
    # 用已有的 style_type 覆盖规则识别（已格式化的段落更准确）
    for i, p in enumerate(paragraphs):
        if i < len(result["outline"]):
            existing_style = p.get("style_type", "body")
            if existing_style != "body":
                result["outline"][i]["style"] = existing_style
    return result


def _build_outline_context(analysis: dict, total_chunks: int,
                           current_chunk: int,
                           chunk_para_range: tuple[int, int] | None = None) -> str:
    """
    Phase-1 → Phase-2 桥接：将结构分析结果转为注入每个分块的上下文文本。

    保持轻量（约 300-800 字符），确保 LLM 知道：
    - 全文结构大纲（哪些是标题/正文/结尾）
    - 当前块在全文中的位置
    - 前后块的内容概要
    - 编号体系（保持一致性）
    """
    sections = analysis.get("sections", [])
    outline = analysis.get("outline", [])
    total = analysis.get("total_paragraphs", 0)
    numbering = analysis.get("numbering", "")

    # 构建精简大纲（只显示非 body 段落 + 章节边界）
    key_points: list[str] = []
    for item in outline:
        if item["style"] != "body":
            key_points.append(f'  [{item["idx"]}] {item["style"]}: {item["preview"]}')
    # 大纲超过 20 行时截取首尾
    if len(key_points) > 20:
        key_points = key_points[:10] + [f"  ... (共 {len(key_points)} 个结构点) ..."] + key_points[-5:]

    outline_text = "\n".join(key_points) if key_points else "  (纯正文，无明显结构标记)"

    # 章节列表
    section_list = ""
    if sections:
        sec_items = [f'  {s["heading"]} (段落 {s["para_range"][0]}-{s["para_range"][1]})' for s in sections[:15]]
        section_list = "\n".join(sec_items)

    # 当前块在全文中的位置说明
    position_hint = ""
    if chunk_para_range:
        start, end = chunk_para_range
        position_hint = f"当前处理段落 [{start}] ~ [{end}]（全文共 {total} 段）。"
        # 找出当前块所在章节
        for sec in sections:
            if sec["para_range"][0] <= start <= sec["para_range"][1]:
                position_hint += f" 当前位于「{sec['heading']}」章节。"
                break

    context = (
        f"【全文结构大纲 — 第 {current_chunk}/{total_chunks} 部分】\n"
        f"文档共 {total} 段，编号体系：{numbering}\n"
    )
    if section_list:
        context += f"章节结构：\n{section_list}\n"
    context += f"关键段落：\n{outline_text}\n"
    if position_hint:
        context += f"{position_hint}\n"
    context += (
        "⚠ 请保持与全文一致的编号体系和段落风格，不要重复添加标题。\n"
    )

    return context


def _split_text_into_chunks(text: str, max_chars: int = _MAX_FORMAT_CHUNK_CHARS) -> list[str]:
    """将长文本按段落边界分割为多个块，每块不超过 max_chars 字符。

    改进：超大段落（如 Markdown 表格、代码块）会在单换行处进一步分割。
    """
    paragraphs = _re.split(r'\n\s*\n', text)
    if not paragraphs:
        return [text]

    chunks: list[str] = []
    current_parts: list[str] = []
    current_len = 0

    for para in paragraphs:
        para_len = len(para) + 2  # +2 for \n\n separator

        # 超大段落（如表格 / 代码块）：进一步按单换行拆分
        if para_len > max_chars:
            if current_parts:
                chunks.append('\n\n'.join(current_parts))
                current_parts = []
                current_len = 0
            sub_lines = para.split('\n')
            # 如果段落无换行（单行超长纯文本），按句号/分号等边界硬截断
            if len(sub_lines) == 1:
                _long = para
                while len(_long) > max_chars:
                    # 在 max_chars 范围内寻找最后一个句子边界
                    _cut = max_chars
                    for _sep in ('。', '；', '！', '？', '.', ';', '!', '?', '，', ','):
                        _pos = _long.rfind(_sep, 0, max_chars)
                        if _pos > max_chars // 3:  # 至少截取 1/3 以上
                            _cut = _pos + 1
                            break
                    chunks.append(_long[:_cut])
                    _long = _long[_cut:]
                if _long:
                    chunks.append(_long)
                continue
            sub_parts: list[str] = []
            sub_len = 0
            for line in sub_lines:
                line_len = len(line) + 1
                if sub_len + line_len > max_chars and sub_parts:
                    chunks.append('\n'.join(sub_parts))
                    sub_parts = []
                    sub_len = 0
                sub_parts.append(line)
                sub_len += line_len
            if sub_parts:
                chunks.append('\n'.join(sub_parts))
            continue

        if current_len + para_len > max_chars and current_parts:
            chunks.append('\n\n'.join(current_parts))
            current_parts = [para]
            current_len = para_len
        else:
            current_parts.append(para)
            current_len += para_len

    if current_parts:
        chunks.append('\n\n'.join(current_parts))

    return chunks if chunks else [text]


# 单块被截断时的最大重试次数
_MAX_CONTINUATION_RETRIES = 3

# 注意: _INCREMENTAL_MAX_PARAS_PER_CHUNK / _INCREMENTAL_MAX_CHARS_PER_CHUNK
# 已在上方"两阶段处理"代码块中声明（40 / 4500），此处不再重复声明。


def _split_paragraphs_into_chunks(
    paragraphs: list[dict],
    max_chars: int = _INCREMENTAL_MAX_CHARS_PER_CHUNK,
    max_paras: int = _INCREMENTAL_MAX_PARAS_PER_CHUNK,
) -> list[tuple[int, list[dict]]]:
    """
    将段落列表分割为多个块，返回 [(start_global_index, chunk_paragraphs), ...]。
    每块不超过 max_chars 字符或 max_paras 段落。
    """
    if not paragraphs:
        return []

    chunks: list[tuple[int, list[dict]]] = []
    current_chunk: list[dict] = []
    current_chars = 0
    chunk_start = 0

    for i, para in enumerate(paragraphs):
        para_chars = len(para.get("text", "")) + 80  # 属性开销

        if current_chunk and (current_chars + para_chars > max_chars or len(current_chunk) >= max_paras):
            chunks.append((chunk_start, current_chunk))
            current_chunk = []
            current_chars = 0
            chunk_start = i

        current_chunk.append(para)
        current_chars += para_chars

    if current_chunk:
        chunks.append((chunk_start, current_chunk))

    return chunks


async def _chunked_incremental_format_stream(
    dify,
    paragraphs: list[dict],
    doc_type: str,
    user_instruction: str,
    max_chunk_chars: int = _INCREMENTAL_MAX_CHARS_PER_CHUNK,
    max_paras: int = _INCREMENTAL_MAX_PARAS_PER_CHUNK,
):
    """
    对长文档的结构化段落逐块增量排版（带重试），实时推送进度。

    改进（v3）：
    - 顺序处理 + 逐块重试（最多 3 次），避免并行时 SSE 无心跳导致卡顿
    - 每块处理完立即推送进度事件，保持前端连接活跃
    - 失败块跳过但不中断后续块
    """
    from app.services.dify.base import SSEEvent
    import asyncio as _aio

    para_chunks = _split_paragraphs_into_chunks(paragraphs, max_chars=max_chunk_chars, max_paras=max_paras)
    total = len(para_chunks)
    logger.info(f"增量分块排版: {len(paragraphs)} 段 → {total} 块 (每块 ≤{max_chunk_chars} 字符 / {max_paras} 段)")

    # ── Phase-1：规则化结构分析，生成全局大纲 ──
    doc_analysis: dict = {}
    if total > 1:
        try:
            doc_analysis = _analyze_paragraphs_structure(paragraphs)
            logger.info(
                f"Phase-1 结构分析完成: 类型={doc_analysis.get('doc_type')}, "
                f"段落={doc_analysis.get('total_paragraphs')}, "
                f"章节={len(doc_analysis.get('sections', []))}"
            )
        except Exception as e:
            logger.warning(f"Phase-1 结构分析失败(不影响排版): {e}")

    all_modified: dict[int, dict] = {}
    all_full_output: list[dict] = []
    _conv_id = ""  # 复用 conversation_id 减少 Dify 连接开销

    for chunk_idx, (start_idx, chunk_paras) in enumerate(para_chunks):
        pct = round((chunk_idx / total) * 100)
        yield SSEEvent(event="progress", data={
            "message": f"正在格式化第 {chunk_idx + 1}/{total} 部分… ({len(chunk_paras)} 段)"
        })
        yield SSEEvent(event="format_progress", data={
            "current": chunk_idx + 1, "total": total, "percent": pct
        })

        # 构建本块的 compact listing（全局索引）
        _compact_lines = []
        for local_i, _p in enumerate(chunk_paras):
            global_i = start_idx + local_i
            _attrs = [_p.get("style_type", "body")]
            if _p.get("font_size"): _attrs.append(_p["font_size"])
            if _p.get("font_family"): _attrs.append(_p["font_family"])
            if _p.get("bold"): _attrs.append("bold")
            if _p.get("alignment") and _p["alignment"] != "left": _attrs.append(_p["alignment"])
            if _p.get("indent"): _attrs.append(f'indent={_p["indent"]}')
            if _p.get("line_height"): _attrs.append(f'lh={_p["line_height"]}')
            if _p.get("color") and _p["color"] != "#000000": _attrs.append(f'color={_p["color"]}')
            _text = _p.get("text", "")
            _compact_lines.append(f'[{global_i}] ({", ".join(_attrs)}) {_text}')
        _compact_listing = "\n".join(_compact_lines)

        # Phase-2：注入全局大纲上下文
        outline_ctx = ""
        if doc_analysis and total > 1:
            para_range = (start_idx, start_idx + len(chunk_paras) - 1)
            outline_ctx = _build_outline_context(doc_analysis, total, chunk_idx + 1, para_range)

        chunk_instr = (
            "[增量修改模式 — 仅输出被修改的段落]\n"
        )
        if outline_ctx:
            chunk_instr += f"\n{outline_ctx}\n\n"
        chunk_instr += (
            f"当前处理长文档的第 {chunk_idx + 1}/{total} 部分，"
            f"段落索引 [{start_idx}] ~ [{start_idx + len(chunk_paras) - 1}]，"
            f"共 {len(chunk_paras)} 段：\n"
            f"{_compact_listing}\n\n"
            "规则：\n"
            "1. 仅输出需要修改的段落，每个必须包含 _index + 完整 11 个属性\n"
            "2. 未修改的段落不要输出，无修改时输出 {\"paragraphs\": []}\n"
            "3. 不得擅自修改用户未要求修改的属性\n\n"
        )
        if user_instruction:
            chunk_instr += f"【用户修改要求】:\n{user_instruction}"
        else:
            chunk_instr += "请对这部分段落进行格式化排版。"

        # ── 带重试的单块处理 ──
        chunk_para_data: list[dict] = []
        max_retries = 3
        for attempt in range(max_retries):
            chunk_para_data = []
            try:
                async for event in dify.run_doc_format_stream("", doc_type, chunk_instr,
                                                               conversation_id=_conv_id):
                    if event.event == "structured_paragraph":
                        pd = dict(event.data)
                        chunk_para_data.append(pd)
                        yield event  # 实时推送段落到前端
                    elif event.event == "message_end":
                        if not _conv_id and event.data.get("conversation_id"):
                            _conv_id = event.data["conversation_id"]
                        break
                    elif event.event == "error":
                        raise RuntimeError(event.data.get("message", "Dify error"))
                    elif event.event in ("progress", "reasoning"):
                        yield event  # 转发心跳 + 思考过程
                # 成功则退出重试循环
                break
            except Exception as e:
                if attempt < max_retries - 1:
                    logger.warning(f"增量分块 {chunk_idx + 1}/{total}: 第 {attempt + 1} 次失败 ({e})，重试…")
                    yield SSEEvent(event="progress", data={
                        "message": f"第 {chunk_idx + 1} 部分重试中… ({attempt + 2}/{max_retries})"
                    })
                    await _aio.sleep(2 ** attempt)
                else:
                    logger.error(f"增量分块 {chunk_idx + 1}/{total}: {max_retries} 次重试均失败: {e}")
                    yield SSEEvent(event="progress", data={
                        "message": f"第 {chunk_idx + 1} 部分处理失败，已跳过"
                    })

        # 收集结果
        for pd in chunk_para_data:
            all_full_output.append(pd)
            idx = pd.get("_index")
            if idx is not None and isinstance(idx, int):
                all_modified[idx] = pd
        logger.info(f"增量分块 {chunk_idx + 1}/{total}: 产出 {len(chunk_para_data)} 段修改")

    # 完成进度
    yield SSEEvent(event="format_progress", data={
        "current": total, "total": total, "percent": 100
    })
    yield SSEEvent(event="message_end", data={"full_text": ""})
    logger.info(f"增量分块排版完成: 共修改 {len(all_modified)} / {len(paragraphs)} 段")


async def _chunked_format_stream(dify, doc_text: str, doc_type: str,
                                 user_instruction: str,
                                 max_chunk_chars: int = _MAX_FORMAT_CHUNK_CHARS):
    """
    对长文档逐块调用 Dify 排版，实时推送进度事件流。

    改进（v3）：
    - 顺序处理 + 逐块重试（最多 3 次），保持 SSE 连接活跃
    - 增大分块尺寸（3500 字符），减少总块数，提升上下文连续性
    - 非首块添加续接提示（含前一块末尾摘要），保证排版连贯
    - 每块产出的段落实时推送到前端，不用等全部完成
    """
    from app.services.dify.base import SSEEvent
    import asyncio as _aio

    chunks = _split_text_into_chunks(doc_text, max_chunk_chars)
    total = len(chunks)
    logger.info(f"长文档分块排版: {len(doc_text)} 字符 → {total} 块 (每块 ≤{max_chunk_chars} 字符)")

    # 单块不需要分块流程
    if total <= 1:
        async for event in dify.run_doc_format_stream(
            doc_text, doc_type, user_instruction,
        ):
            yield event
        return

    # ── Phase-1：规则化结构分析，生成全局大纲 ──
    doc_analysis: dict = {}
    try:
        doc_analysis = _analyze_doc_structure(doc_text)
        logger.info(
            f"Phase-1 结构分析完成: 类型={doc_analysis.get('doc_type')}, "
            f"段落={doc_analysis.get('total_paragraphs')}, "
            f"章节={len(doc_analysis.get('sections', []))}"
        )
    except Exception as e:
        logger.warning(f"Phase-1 结构分析失败(不影响排版): {e}")

    # 计算每块的段落偏移量（估算：按字符数比例推算段落范围）
    _cumulative_chars = 0
    _chunk_para_ranges: list[tuple[int, int]] = []
    _total_paras = doc_analysis.get("total_paragraphs", 0) if doc_analysis else 0
    _total_chars = len(doc_text) or 1
    for _ci, _ct in enumerate(chunks):
        _start_ratio = _cumulative_chars / _total_chars
        _cumulative_chars += len(_ct)
        _end_ratio = _cumulative_chars / _total_chars
        _est_start = int(_start_ratio * _total_paras)
        _est_end = min(int(_end_ratio * _total_paras), _total_paras - 1)
        _chunk_para_ranges.append((_est_start, _est_end))

    global_para_count = 0
    failed_chunks = 0
    prev_tail_text = ""  # 前一块最后段落的文本，用于续接上下文
    _conv_id = ""  # 复用 conversation_id 减少 Dify 连接开销

    for i, chunk_text in enumerate(chunks):
        pct = round((i / total) * 100)
        yield SSEEvent(
            event="progress",
            data={"message": f"正在格式化第 {i + 1}/{total} 部分… (共 {total} 部分)"},
        )
        yield SSEEvent(
            event="format_progress",
            data={"current": i + 1, "total": total, "percent": pct},
        )

        # Phase-2：注入全局大纲上下文 + 续接提示
        chunk_instr = user_instruction
        if i > 0 or (doc_analysis and total > 1):
            outline_ctx = ""
            if doc_analysis:
                para_range = _chunk_para_ranges[i] if i < len(_chunk_para_ranges) else (0, 0)
                outline_ctx = _build_outline_context(doc_analysis, total, i + 1, para_range)

            if i > 0:
                context_hint = ""
                if prev_tail_text:
                    tail_snippet = prev_tail_text[:80]
                    context_hint = f"前一部分结尾内容：「{tail_snippet}」\n"
                hint = (
                    f"（续：这是长文档的第 {i + 1}/{total} 部分，接续上文。"
                    f"{context_hint}"
                    f"请直接对这部分文本进行结构识别和排版，不要重复添加标题。"
                    f"段落序号接续前一部分。）\n"
                )
                if outline_ctx:
                    hint = outline_ctx + "\n" + hint
                chunk_instr = hint + ("\n" + user_instruction if user_instruction else "")
            elif outline_ctx:
                # 首块也注入大纲（知道全局结构）
                chunk_instr = outline_ctx + "\n" + (user_instruction or "")

        # ── 带重试的单块处理 ──
        chunk_paras: list[dict] = []
        max_retries = 3
        for attempt in range(max_retries):
            chunk_paras = []
            try:
                async for event in dify.run_doc_format_stream(
                    chunk_text, doc_type, chunk_instr,
                    conversation_id=_conv_id,
                ):
                    if event.event == "structured_paragraph":
                        chunk_paras.append(dict(event.data))
                        global_para_count += 1
                        yield event  # 实时推送段落
                    elif event.event == "message_end":
                        # 捕获 conversation_id 供后续块复用
                        if not _conv_id and event.data.get("conversation_id"):
                            _conv_id = event.data["conversation_id"]
                        break
                    elif event.event == "error":
                        raise RuntimeError(event.data.get("message", "Dify error"))
                    elif event.event in ("progress", "reasoning"):
                        yield event  # 转发 Dify 心跳 + 思考过程
                # 成功则退出重试循环
                if chunk_paras or not chunk_text.strip():
                    break
                # 空结果但有内容 → 重试
                if attempt < max_retries - 1:
                    logger.warning(f"分块 {i + 1}/{total}: 空结果，第 {attempt + 1} 次重试")
                    yield SSEEvent(event="progress", data={
                        "message": f"第 {i + 1} 部分重试中… ({attempt + 2}/{max_retries})"
                    })
                    await _aio.sleep(2 ** attempt)
            except Exception as e:
                if attempt < max_retries - 1:
                    logger.warning(f"分块 {i + 1}/{total}: 第 {attempt + 1} 次失败 ({e})，重试…")
                    yield SSEEvent(event="progress", data={
                        "message": f"第 {i + 1} 部分重试中… ({attempt + 2}/{max_retries})"
                    })
                    await _aio.sleep(2 ** attempt)
                else:
                    failed_chunks += 1
                    logger.error(f"分块 {i + 1}/{total}: {max_retries} 次重试均失败: {e}")
                    yield SSEEvent(event="progress", data={
                        "message": f"第 {i + 1}/{total} 部分处理失败，已跳过"
                    })

        # 记录末尾文本，供下一块续接
        if chunk_paras:
            prev_tail_text = chunk_paras[-1].get("text", "")
            logger.info(f"分块 {i + 1}/{total}: 产出 {len(chunk_paras)} 段")
        elif chunk_text.strip():
            failed_chunks += 1
            logger.warning(f"分块 {i + 1}/{total}: 未产出任何段落 ({len(chunk_text)} 字符)")

    # 完成
    yield SSEEvent(
        event="format_progress",
        data={"current": total, "total": total, "percent": 100},
    )
    yield SSEEvent(event="message_end", data={"full_text": ""})
    logger.info(
        f"分块排版完成: {total} 块, 成功 {total - failed_chunks}, "
        f"失败 {failed_chunks}, 共 {global_para_count} 段"
    )


# ── 排版预设模板（与 Dify 提示词中的预设保持一致） ──
_FORMAT_TEMPLATES: dict[str, dict[str, dict]] = {
    "official": {
        "title":      {"font_size": "二号", "font_family": "方正小标宋简体", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "center", "line_height": "2", "red_line": True},
        "recipient":  {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "left", "line_height": "2", "red_line": False},
        "heading1":   {"font_size": "三号", "font_family": "黑体", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "2", "red_line": False},
        "heading2":   {"font_size": "三号", "font_family": "楷体_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "2", "red_line": False},
        "heading3":   {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": True, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "2", "red_line": False},
        "heading4":   {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "2", "red_line": False},
        "body":       {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "justify", "line_height": "2", "red_line": False},
        "closing":    {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "2", "red_line": False},
        "signature":  {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "2", "red_line": False},
        "date":       {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "2", "red_line": False},
        "attachment": {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "left", "line_height": "2", "red_line": False},
    },
    "school_notice_redhead": {
        "title":      {"font_size": "32pt", "font_family": "方正小标宋简体", "bold": False, "italic": False, "color": "#CC0000", "indent": "0", "alignment": "center", "line_height": "1.25", "red_line": True, "letter_spacing": "0.6em"},
        "subtitle":   {"font_size": "二号", "font_family": "方正小标宋简体", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "center", "line_height": "1.32", "red_line": False},
        "recipient":  {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "left", "line_height": "1.81", "red_line": False},
        "heading1":   {"font_size": "三号", "font_family": "黑体", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "1.81", "red_line": False},
        "heading2":   {"font_size": "三号", "font_family": "楷体_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "1.81", "red_line": False},
        "heading3":   {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": True, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "1.81", "red_line": False},
        "heading4":   {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "1.81", "red_line": False},
        "body":       {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "justify", "line_height": "1.81", "red_line": False},
        "closing":    {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "1.81", "red_line": False},
        "signature":  {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "1.81", "red_line": False},
        "date":       {"font_size": "三号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "1.81", "red_line": False},
        "attachment": {"font_size": "四号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#333333", "indent": "0", "alignment": "left", "line_height": "1.5", "red_line": False, "footer_line": True},
    },
    "academic": {
        "title":    {"font_size": "三号", "font_family": "黑体", "bold": True, "italic": False, "color": "#000000", "indent": "0", "alignment": "center", "line_height": "1.5", "red_line": False},
        "heading1": {"font_size": "四号", "font_family": "黑体", "bold": True, "italic": False, "color": "#000000", "indent": "0", "alignment": "left", "line_height": "1.5", "red_line": False},
        "heading2": {"font_size": "小四", "font_family": "黑体", "bold": True, "italic": False, "color": "#000000", "indent": "0", "alignment": "left", "line_height": "1.5", "red_line": False},
        "body":     {"font_size": "五号", "font_family": "宋体", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "justify", "line_height": "1.5", "red_line": False},
        "signature":{"font_size": "五号", "font_family": "宋体", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "1.5", "red_line": False},
        "date":     {"font_size": "五号", "font_family": "宋体", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "1.5", "red_line": False},
    },
    "legal": {
        "title":    {"font_size": "二号", "font_family": "宋体", "bold": True, "italic": False, "color": "#000000", "indent": "0", "alignment": "center", "line_height": "2", "red_line": False},
        "heading1": {"font_size": "四号", "font_family": "黑体", "bold": True, "italic": False, "color": "#000000", "indent": "0", "alignment": "left", "line_height": "2", "red_line": False},
        "body":     {"font_size": "四号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "justify", "line_height": "2", "red_line": False},
        "signature":{"font_size": "四号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "2", "red_line": False},
        "date":     {"font_size": "四号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "2", "red_line": False},
    },
    # ── 项目建议书格式 ──
    # A4 幅面, 上下 2.5cm, 左右 2.6cm, 页眉 1.5cm, 页脚 2.0cm
    # 正文行间距固定值 25 磅, 首行缩进 2 字符
    # 一级标题黑体三号; 二级标题楷体三号; 三/四级仿宋_GB2312 四号加粗; 正文仿宋_GB2312 小四
    # 编号次序: 二→(二)→2→2.1→2.1.1→(1)→①
    "proposal": {
        "title":      {"font_size": "二号", "font_family": "方正小标宋简体", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "center", "line_height": "1.5", "red_line": False},
        "recipient":  {"font_size": "小四", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "left", "line_height": "1.5", "red_line": False},
        "heading1":   {"font_size": "三号", "font_family": "黑体", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "1.5", "red_line": False},
        "heading2":   {"font_size": "三号", "font_family": "楷体_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "1.5", "red_line": False},
        "heading3":   {"font_size": "四号", "font_family": "仿宋_GB2312", "bold": True, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "1.5", "red_line": False},
        "heading4":   {"font_size": "四号", "font_family": "仿宋_GB2312", "bold": True, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "1.5", "red_line": False},
        "body":       {"font_size": "小四", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "justify", "line_height": "1.5", "red_line": False},
        "closing":    {"font_size": "小四", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "1.5", "red_line": False},
        "signature":  {"font_size": "小四", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "1.5", "red_line": False},
        "date":       {"font_size": "小四", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "1.5", "red_line": False},
        "attachment": {"font_size": "小四", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "left", "line_height": "1.5", "red_line": False},
    },
    # ── 重点实验室基金指南格式 ──
    # 标题: 方正小标宋简体二号, 单倍行距, 居中
    # 一级标题: 黑体四号, 首行缩进 2 字符, 行间距固定值 26 磅
    # 正文: 仿宋_GB2312 四号, 首行缩进 2 字符, 行间距固定值 26 磅
    "lab_fund": {
        "title":      {"font_size": "二号", "font_family": "方正小标宋简体", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "center", "line_height": "1", "red_line": False},
        "recipient":  {"font_size": "四号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "left", "line_height": "1.5", "red_line": False},
        "heading1":   {"font_size": "四号", "font_family": "黑体", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "1.5", "red_line": False},
        "heading2":   {"font_size": "四号", "font_family": "楷体_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "1.5", "red_line": False},
        "heading3":   {"font_size": "四号", "font_family": "仿宋_GB2312", "bold": True, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "1.5", "red_line": False},
        "heading4":   {"font_size": "四号", "font_family": "仿宋_GB2312", "bold": True, "italic": False, "color": "#000000", "indent": "2em", "alignment": "left", "line_height": "1.5", "red_line": False},
        "body":       {"font_size": "四号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "2em", "alignment": "justify", "line_height": "1.5", "red_line": False},
        "closing":    {"font_size": "四号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "1.5", "red_line": False},
        "signature":  {"font_size": "四号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "1.5", "red_line": False},
        "date":       {"font_size": "四号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "right", "line_height": "1.5", "red_line": False},
        "attachment": {"font_size": "四号", "font_family": "仿宋_GB2312", "bold": False, "italic": False, "color": "#000000", "indent": "0", "alignment": "left", "line_height": "1.5", "red_line": False},
    },
}


def _apply_format_template(para: dict, doc_type: str) -> dict:
    """
    Fill in missing formatting attributes from the preset template.
    If LLM already specified an attribute (e.g. user override), keep it.
    """
    templates = _FORMAT_TEMPLATES.get(doc_type, _FORMAT_TEMPLATES["official"])
    style = para.get("style_type", "body")
    defaults = templates.get(style, templates.get("body", {}))
    for key, default_val in defaults.items():
        if key not in para or para[key] is None:
            para[key] = default_val
    return para


# ── 规则引擎排版（Phase-1：零 LLM 调用） ──────────────────

def _detect_style_with_confidence(
    text: str, idx: int, total: int,
    has_title: bool = False, has_closing: bool = False, has_signature: bool = False,
) -> tuple[str, float]:
    """
    对单段落文本返回 (style_type, confidence)。

    - 正则精确匹配 → confidence = 0.95
    - 启发式匹配 → confidence = 0.6
    - 兜底 body → confidence = 0.3

    Args:
        text: 段落文本（已 strip）
        idx: 段落在文档中的位置
        total: 文档总段落数
        has_title / has_closing / has_signature: 前面已识别的标记

    Returns:
        (style_type, confidence)
    """
    if not text.strip():
        return ("body", 0.1)

    stripped = text.strip()

    # ── 精确正则匹配（confidence = 0.95） ──
    if not has_title and (_RE_TITLE.match(stripped) or (idx == 0 and len(stripped) < 60)):
        return ("title", 0.95)
    # subtitle：title 之后的"关于…的请示/通知/报告"行
    if has_title and idx <= 2 and _re.match(r'^关于.{2,}的(请示|通知|报告|函|批复|决定|意见)', stripped):
        return ("subtitle", 0.90)
    if _RE_HEADING1.match(stripped):
        return ("heading1", 0.95)
    if _RE_HEADING2.match(stripped):
        return ("heading2", 0.95)
    if _RE_HEADING3.match(stripped):
        return ("heading3", 0.95)
    if _RE_HEADING4.match(stripped):
        return ("heading4", 0.95)
    if _RE_RECIPIENT.match(stripped) and idx <= 3:
        return ("recipient", 0.95)
    if _RE_CLOSING.match(stripped):
        return ("closing", 0.95)
    if _RE_DATE.match(stripped):
        return ("date", 0.95)
    if _RE_ATTACHMENT.match(stripped):
        return ("attachment", 0.95)

    # ── 尾部署名启发式（confidence = 0.85） ──
    if idx >= total - 3 and len(stripped) < 30 and not _RE_HEADING1.match(stripped):
        if not _RE_DATE.match(stripped) and not has_signature:
            return ("signature", 0.85)

    # ── 启发式匹配（confidence = 0.6） ──
    # 短行 + 无句尾标点 → 可能是标题/署名
    if len(stripped) < 25 and not _re.search(r'[。！？；：，]$', stripped):
        # 首段可能是标题
        if idx == 0:
            return ("title", 0.6)
        # 包含编号特征但不完全匹配
        if _re.match(r'^第[一二三四五六七八九十]+[章节条]', stripped):
            return ("heading1", 0.7)

    # ── 正文段落启发式（confidence = 0.85） ──
    # 长段落 + 句末中文标点 → 几乎可以确定是正文
    if len(stripped) > 20 and _re.search(r'[。！？）；：]$', stripped):
        return ("body", 0.85)
    # 中等长度段落，不匹配任何特殊模式 → 大概率正文
    if len(stripped) > 15:
        return ("body", 0.6)

    return ("body", 0.3)


def _rules_format_paragraphs(
    paragraphs: list[dict], doc_type: str = "official",
) -> tuple[list[dict], list[int]]:
    """
    规则引擎排版：对所有段落做 style_type 检测 + 模板属性填充。

    高置信度（>= 0.8）的段落直接应用模板，
    低置信度的收集到 llm_needed_indices 由 LLM 降级处理。

    Args:
        paragraphs: 原始段落列表（每项至少含 text）
        doc_type: 文档类型，用于选择 _FORMAT_TEMPLATES

    Returns:
        (formatted_paras, llm_needed_indices)
        formatted_paras: 完整段落列表（高置信度已填充排版属性）
        llm_needed_indices: 需要 LLM 处理的段落下标
    """
    total = len(paragraphs)
    formatted: list[dict] = []
    llm_needed: list[int] = []
    has_title = False
    has_closing = False
    has_signature = False

    for idx, para in enumerate(paragraphs):
        text = para.get("text", "").strip()
        out = dict(para)  # 保留原有属性

        # 如果段落已有明确 style_type（非 body），保留不覆盖
        existing_style = para.get("style_type", "")
        if existing_style and existing_style != "body":
            # 已有明确样式 → 仅补全模板缺失属性
            _apply_format_template(out, doc_type)
            out["_rule_formatted"] = True
            formatted.append(out)
            if existing_style == "title":
                has_title = True
            elif existing_style == "closing":
                has_closing = True
            elif existing_style == "signature":
                has_signature = True
            continue

        # 如果上游（_parse_markdown_to_paragraphs）已附带置信度，直接复用，避免重复正则匹配
        pre_confidence = para.get("_confidence")
        if pre_confidence is not None and existing_style:
            style = existing_style
            confidence = pre_confidence
        else:
            # 规则引擎检测
            style, confidence = _detect_style_with_confidence(
                text, idx, total, has_title, has_closing, has_signature,
            )

        if confidence >= 0.8:
            out["style_type"] = style
            _apply_format_template(out, doc_type)
            out["_rule_formatted"] = True
            if style == "title":
                has_title = True
            elif style == "closing":
                has_closing = True
            elif style == "signature":
                has_signature = True
        else:
            # 低置信度：设置初步 style_type 但标记需要 LLM 确认
            out["style_type"] = style
            out["_rule_formatted"] = False
            llm_needed.append(idx)

        formatted.append(out)

    return formatted, llm_needed


def _build_formatted_docx(paragraphs: list[dict], title: str, preset: str = "official"):
    """
    根据 StructuredDocRenderer 的 STYLE_PRESETS 逻辑生成高质量 DOCX。
    与前端实时预览保持像素级一致：
    - 先取 preset 默认值，再用 LLM 显式输出的属性覆盖
    - 固定行距（Exactly）精确匹配 CSS lineHeight（font_size × multiplier）
    - 四槽字体设置（ascii / hAnsi / eastAsia / cs）
    - GB/T 9704 红头文件红线直接附在标题段落底边框上（无多余空段落）
    - 署名右缩进、页码等细节
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re as _re

    doc = DocxDocument()

    # ── 全局默认样式：清理 Normal 样式避免干扰 ──
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(16)
    # 通过 XML 统一设置段落间距 + 固定行距，避免 python-docx API 生成重复节点
    nPr = normal_style.element.get_or_add_pPr()
    existing_spacing = nPr.find(qn('w:spacing'))
    if existing_spacing is not None:
        nPr.remove(existing_spacing)
    nSpacing = OxmlElement('w:spacing')
    nSpacing.set(qn('w:line'), '640')     # 16pt * 2.0 * 20 twips = 640
    nSpacing.set(qn('w:lineRule'), 'exact')
    nSpacing.set(qn('w:before'), '0')
    nSpacing.set(qn('w:after'), '0')
    nPr.append(nSpacing)
    normal_style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # ── 清除所有内置样式的编号属性，防止段落出现黑色项目符号 ──
    for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
                        'List Bullet', 'List Number', 'List Paragraph']:
        try:
            s = doc.styles[style_name]
            sPr = s.element.get_or_add_pPr()
            for numPr in sPr.findall(qn('w:numPr')):
                sPr.remove(numPr)
        except KeyError:
            pass

    # ── 页面设置（GB/T 9704-2012 标准） ──
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        # 页码（居中，— n — 格式）
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 页码前缀
        run_prefix = fp.add_run("— ")
        run_prefix.font.size = Pt(9)
        run_prefix.font.name = 'Times New Roman'
        # PAGE 域: begin
        run_begin = fp.add_run()
        run_begin.font.size = Pt(9)
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run_begin._element.append(fld_begin)
        # PAGE 域: instrText
        run_instr = fp.add_run()
        run_instr.font.size = Pt(9)
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        run_instr._element.append(instr)
        # PAGE 域: end
        run_end = fp.add_run()
        run_end.font.size = Pt(9)
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run_end._element.append(fld_end)
        # 页码后缀
        run_suffix = fp.add_run(" —")
        run_suffix.font.size = Pt(9)
        run_suffix.font.name = 'Times New Roman'

    alignment_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    # 字体别名归一化
    _FONT_ALIAS = {
        "仿宋": "仿宋_GB2312", "fangsong": "仿宋_GB2312", "FangSong": "仿宋_GB2312",
        "华文仿宋": "仿宋_GB2312", "STFangsong": "仿宋_GB2312",
        "楷体": "楷体_GB2312", "kaiti": "楷体_GB2312", "KaiTi": "楷体_GB2312",
        "华文楷体": "楷体_GB2312", "STKaiti": "楷体_GB2312",
        "黑体": "黑体", "SimHei": "黑体", "heiti": "黑体",
        "宋体": "宋体", "SimSun": "宋体", "songti": "宋体",
        "方正小标宋简体": "方正小标宋简体", "FZXiaoBiaoSong": "方正小标宋简体",
        "方正小标宋": "方正小标宋简体",
        "微软雅黑": "微软雅黑", "Microsoft YaHei": "微软雅黑",
        "华文中宋": "华文中宋", "STZhongsong": "华文中宋",
    }

    def _normalize_font(raw: str) -> str:
        return _FONT_ALIAS.get(raw.strip(), raw.strip())

    def _set_run_font(run, cn_font: str, en_font: str | None = None):
        """精确设置 run 的四槽字体（ASCII/Latin 统一用 Times New Roman）"""
        if en_font is None:
            en_font = 'Times New Roman'
        run.font.name = en_font
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), en_font)
        rFonts.set(qn('w:hAnsi'), en_font)
        rFonts.set(qn('w:eastAsia'), cn_font)
        rFonts.set(qn('w:cs'), en_font)

    def _set_exact_line_spacing(para, font_size_pt: float, multiplier: float):
        """
        设置固定行距以精确匹配 CSS lineHeight。
        CSS lineHeight: N = N × font-size。
        Word Exactly: line 值单位为 twips（1pt = 20 twips）。
        exact_pt = font_size_pt × multiplier → twips = exact_pt × 20
        """
        exact_twips = int(font_size_pt * multiplier * 20)
        pPr = para._element.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:line'), str(exact_twips))
        spacing.set(qn('w:lineRule'), 'exact')

    def _add_bottom_border_to_para(para, color: str = 'CC0000', sz: str = '18'):
        """直接给段落添加底边框（红线），不创建额外空段落"""
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), sz)      # 1/8pt 为单位, 18 = 2.25pt
        bottom.set(qn('w:space'), '6')  # 边框与文字间距（pt）
        bottom.set(qn('w:color'), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_top_double_border_to_para(para):
        """为段落顶部添加双线边框（版记线）"""
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'double')
        top.set(qn('w:sz'), '24')    # 1/8pt, 24 = 3pt
        top.set(qn('w:space'), '8')  # 边框与文字间距 (pt)
        top.set(qn('w:color'), '000000')
        pBdr.append(top)
        pPr.append(pBdr)

    def _set_run_letter_spacing(run, spacing_str: str):
        """设置 Run 的字符间距（letter-spacing），支持 '0.6em' / '10pt' 等"""
        m = _re.match(r'^([\d.]+)\s*(em|pt|px)?$', spacing_str.strip())
        if not m:
            return
        val = float(m.group(1))
        unit = m.group(2) or 'em'
        if unit == 'em':
            font_size_pt = run.font.size.pt if run.font.size else 16
            twips = int(val * font_size_pt * 20)
        elif unit == 'pt':
            twips = int(val * 20)
        elif unit == 'px':
            twips = int(val * 0.75 * 20)
        else:
            return
        rPr = run._element.get_or_add_rPr()
        spacing_el = OxmlElement('w:spacing')
        spacing_el.set(qn('w:val'), str(twips))
        rPr.append(spacing_el)

    def _clear_numPr(para):
        """清除段落的 numPr 编号属性，防止出现项目符号黑点"""
        pPr = para._element.get_or_add_pPr()
        for numPr in pPr.findall(qn('w:numPr')):
            pPr.remove(numPr)

    preset_styles = _STYLE_PRESETS.get(preset, _STYLE_PRESETS["official"])
    body_default = preset_styles.get("body", _STYLE_PRESETS["official"]["body"])

    prev_style_type = None

    # 将段内换行拆分为多个段落，避免 Word 内部换行导致格式混乱
    expanded_paragraphs: list[dict] = []
    for _p in paragraphs:
        _text = str(_p.get("text", ""))
        if "\n" in _text:
            for _line in _text.splitlines():
                _np = dict(_p)
                _np["text"] = _line
                expanded_paragraphs.append(_np)
        else:
            expanded_paragraphs.append(_p)

    # 预过滤空段落（与 HTML 导出保持一致，避免空行累积产生间距）
    expanded_paragraphs = [p for p in expanded_paragraphs if str(p.get("text", "")).strip()]

    for _para_idx, para_data in enumerate(expanded_paragraphs):
        text = para_data.get("text", "")

        # 1) 归一化 style_type
        style_type = _normalize_style_type(para_data.get("style_type"))

        # 2) 取 preset 默认值
        defaults = preset_styles.get(style_type, body_default)

        # 3) 解析 LLM 显式属性
        llm_font_family = para_data.get("font_family") or None
        llm_font_size_pt = _resolve_font_size_pt(para_data.get("font_size"))
        llm_alignment = _resolve_alignment(para_data.get("alignment"))
        llm_indent_em = _resolve_indent_em(para_data.get("indent"))
        llm_line_height = _resolve_line_height(para_data.get("line_height"))
        llm_bold = para_data.get("bold")
        llm_italic = para_data.get("italic")
        llm_color = para_data.get("color")

        # 4) 合并：LLM → preset → global default
        final_cn_font = _normalize_font(llm_font_family) if llm_font_family else defaults["font_family"]
        final_font_size_pt = llm_font_size_pt if llm_font_size_pt is not None else defaults["font_size_pt"]
        final_alignment = llm_alignment or defaults["alignment"]
        final_indent_em = llm_indent_em if llm_indent_em is not None else defaults["indent_em"]
        final_line_height = llm_line_height if llm_line_height is not None else defaults["line_height"]
        final_bold = llm_bold if llm_bold is not None else defaults.get("bold", False)
        final_italic = llm_italic if llm_italic is not None else False

        # 解析颜色
        final_color = None
        if llm_color:
            c = str(llm_color).strip()
            _COLOR_NAME_MAP = {
                "黑色": "#000000", "红色": "#CC0000", "蓝色": "#0033CC",
                "绿色": "#006600", "紫色": "#800080", "深灰": "#333333",
                "灰色": "#666666",
                "black": "#000000", "red": "#CC0000", "blue": "#0033CC",
                "green": "#006600", "purple": "#800080",
            }
            if c.lower() in _COLOR_NAME_MAP:
                c = _COLOR_NAME_MAP[c.lower()]
            elif c in _COLOR_NAME_MAP:
                c = _COLOR_NAME_MAP[c]
            if not c.startswith("#"):
                c = "#" + c
            if _re.match(r'^#[0-9A-Fa-f]{6}$', c):
                final_color = c

        # ── 构建段落 ──
        p = doc.add_paragraph(style='Normal')
        _clear_numPr(p)

        # 对齐：短行使用左对齐，避免 Word 分散对齐导致字间距异常
        _effective_alignment = final_alignment
        if final_alignment == "justify":
            _t = text.strip()
            if len(_t) <= 20 or _t.endswith("：") or _t.endswith(":"):
                _effective_alignment = "left"
        p.alignment = alignment_map.get(_effective_alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

        # 首行缩进
        if final_indent_em and final_indent_em > 0:
            p.paragraph_format.first_line_indent = Pt(final_font_size_pt * final_indent_em)

        # ── 固定行距（精确匹配 CSS lineHeight） ──
        _exact_ls_pt = defaults.get("exact_line_spacing_pt")
        if _exact_ls_pt:
            # 预设指定了精确固定行距（如 28.95pt），直接使用
            _set_exact_line_spacing(p, _exact_ls_pt, 1.0)
        elif final_line_height and final_line_height > 0:
            _set_exact_line_spacing(p, final_font_size_pt, final_line_height)

        # ── 段前段后间距 ──
        space_before = defaults.get("space_before_pt", 0)
        space_after = defaults.get("space_after_pt", 0)

        # 动态间距规则（与前端 getSpacingTop 完全对齐）
        if prev_style_type:
            if style_type == "title":
                pass
            elif style_type == "subtitle" and prev_style_type == "title":
                space_before = max(space_before, 4)
            elif style_type == "recipient" and prev_style_type in ("title", "subtitle"):
                space_before = max(space_before, 8)
            elif style_type.startswith("heading") and not prev_style_type.startswith("heading"):
                space_before = max(space_before, 12)
            elif style_type.startswith("heading") and prev_style_type.startswith("heading"):
                space_before = max(space_before, 4)
            elif style_type in ("signature", "date") and prev_style_type not in ("signature", "date"):
                space_before = max(space_before, 18)
            elif style_type == "attachment" and prev_style_type != "attachment":
                space_before = max(space_before, 14)

        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)

        # heading / title 保持与下段不分页
        # 仅当下一个非空段落不是 heading/title 时才设置 keep_with_next，
        # 避免连续 heading 形成长链导致 Word 将大量内容推到下一页。
        if style_type.startswith("heading") or style_type in ("title", "subtitle"):
            _next_st = None
            for _nj in range(_para_idx + 1, len(expanded_paragraphs)):
                _nt = str(expanded_paragraphs[_nj].get("text", "")).strip()
                if _nt:
                    _next_st = _normalize_style_type(expanded_paragraphs[_nj].get("style_type"))
                    break
            if not _next_st or not (_next_st.startswith("heading") or _next_st in ("title", "subtitle")):
                p.paragraph_format.keep_with_next = True

        # 署名/日期右缩进
        if style_type in ("signature", "date") and final_alignment == "right":
            p.paragraph_format.right_indent = Cm(2.0)

        # ── 添加文本 Run ──
        run = p.add_run(text)

        # 字号
        run.font.size = Pt(final_font_size_pt)

        # 四槽字体（ASCII/Latin 也使用中文字体，避免细微差异）
        _set_run_font(run, final_cn_font)

        # 加粗
        run.font.bold = True if final_bold else False

        # 斜体
        if final_italic:
            run.font.italic = True

        # letter_spacing: 通过低层 XML <w:rPr><w:spacing w:val="N"/> 设置
        llm_letter_spacing = para_data.get("letter_spacing")
        if llm_letter_spacing:
            _set_run_letter_spacing(run, llm_letter_spacing)

        # 颜色
        if final_color:
            try:
                r_val = int(final_color[1:3], 16)
                g_val = int(final_color[3:5], 16)
                b_val = int(final_color[5:7], 16)
                run.font.color.rgb = RGBColor(r_val, g_val, b_val)
            except (ValueError, TypeError):
                pass

        # ── 标题段落直接加底边框作红色分隔线（不创建多余空段落） ──
        # red_line 字段由 AI 控制：None/True → 显示, False → 隐藏
        para_red_line = para_data.get("red_line")
        if style_type == "title" and preset in ("official", "school_notice_redhead") and para_red_line is not False:
            _add_bottom_border_to_para(p)
            # 标题底边框后需要额外段后间距让红线与正文拉开
            p.paragraph_format.space_after = Pt(14)

        # ── 版记反线（仅 date/signature → attachment 过渡处） ──
        para_footer_line = para_data.get("footer_line")
        if style_type == "attachment" and para_footer_line is True and prev_style_type in ("date", "signature"):
            _add_top_double_border_to_para(p)

        prev_style_type = style_type

    # 输出到内存
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


class _ExportDocxRequest(BaseModel):
    paragraphs: list[dict]
    title: str = "排版文档"
    preset: str = "official"


@router.post("/{doc_id}/export-docx")
async def export_formatted_docx(
    doc_id: UUID,
    body: _ExportDocxRequest,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """将结构化段落数据导出为带格式的 DOCX 文件（与前端 StructuredDocRenderer 效果对齐）"""
    import asyncio as _aio_export
    loop = _aio_export.get_event_loop()
    buf = await loop.run_in_executor(None, _build_formatted_docx, body.paragraphs, body.title, body.preset)

    safe_title = body.title.replace("/", "_").replace("\\", "_")[:100]
    from urllib.parse import quote
    encoded_name = quote(f"{safe_title}.docx")

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=\"document.docx\"; filename*=UTF-8''{encoded_name}"
        },
    )


@router.post("/{doc_id}/export-pdf")
async def export_formatted_pdf(
    doc_id: UUID,
    body: _ExportDocxRequest,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """将结构化段落数据导出为 PDF（HTML → WeasyPrint → PDF，与前端渲染保持一致）"""
    from app.services.html_export import render_export_html

    try:
        html_str = render_export_html(body.paragraphs, body.title, body.preset)
    except Exception as e:
        logger.error(f"HTML 渲染失败: {e}")
        raise HTTPException(status_code=500, detail=f"HTML 渲染失败: {str(e)}")

    html_bytes = html_str.encode("utf-8")
    safe_title = body.title.replace("/", "_").replace("\\", "_")[:100]
    pdf_bytes = await convert_to_pdf_bytes(html_bytes, f"{safe_title}.html")
    if not pdf_bytes:
        # 降级：尝试 DOCX → PDF
        logger.warning("HTML→PDF 失败，尝试 DOCX 降级方案")
        try:
            import asyncio as _aio_export
            loop = _aio_export.get_event_loop()
            buf = await loop.run_in_executor(None, _build_formatted_docx, body.paragraphs, body.title, body.preset)
            pdf_bytes = await convert_to_pdf_bytes(buf.getvalue(), f"{safe_title}.docx")
        except Exception:
            pass
    if not pdf_bytes:
        raise HTTPException(status_code=502, detail="PDF 导出失败：converter 服务不可用或转换出错，请稍后重试")

    pdf_buf = io.BytesIO(pdf_bytes)
    from urllib.parse import quote
    encoded_name = quote(f"{safe_title}.pdf")

    return StreamingResponse(
        pdf_buf,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename=\"document.pdf\"; filename*=UTF-8''{encoded_name}"
        },
    )


# ── 源文件下载 & Markdown 预览 ──


@router.get("/{doc_id}/source")
async def download_document_source(
    doc_id: UUID,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """
    下载公文的原始上传文件（PDF/DOCX/XLSX 等）。

    仅对通过 /import 导入的公文有效。
    """
    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return error(ErrorCode.NOT_FOUND, "公文不存在")

    if not doc.source_file_path:
        return error(ErrorCode.NOT_FOUND, "此公文没有关联的源文件（可能是手动创建的公文）")

    source_path = Path(doc.source_file_path)
    if not source_path.exists():
        return error(ErrorCode.NOT_FOUND, "源文件已被删除或不可用")

    ext = doc.source_format or "bin"
    media_type = _MIME_MAP.get(ext, "application/octet-stream")
    # 文件名: 标题.原始扩展名
    download_name = f"{doc.title}.{ext}"

    return FileResponse(
        path=str(source_path),
        media_type=media_type,
        filename=download_name,
    )


@router.get("/{doc_id}/markdown")
async def get_document_markdown(
    doc_id: UUID,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """
    获取公文的 Markdown 文件内容（用于预览/对比）。

    返回磁盘上保存的 .md 文件内容。
    """
    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return error(ErrorCode.NOT_FOUND, "公文不存在")

    if not doc.md_file_path:
        return error(ErrorCode.NOT_FOUND, "此公文没有关联的 Markdown 文件")

    md_path = Path(doc.md_file_path)
    if not md_path.exists():
        return error(ErrorCode.NOT_FOUND, "Markdown 文件已被删除或不可用")

    md_content = md_path.read_text(encoding="utf-8")
    return success(data={
        "document_id": str(doc_id),
        "title": doc.title,
        "source_format": doc.source_format,
        "markdown": md_content,
        "char_count": len(md_content),
    })


@router.get("/{doc_id}/preview-pdf")
async def preview_document_pdf(
    doc_id: UUID,
    token: str = Query(None, description="JWT token（用于 iframe 内嵌预览）"),
    db: AsyncSession = Depends(get_db),
    request: Request = None,
):
    """
    获取公文的 PDF 预览。
    支持 Bearer header 或 ?token= 查询参数（iframe 场景）。

    策略:
      1. 如果源文件本身就是 PDF → 直接返回
      2. 如果有源文件（DOCX/XLSX 等）→ 通过 converter 微服务转为 PDF → 返回
      3. 如果仅有文本内容 → 返回 404（前端降级到 Markdown 预览）
    """
    # ── 手动鉴权：优先 header，降级到 query param ──
    from app.core.security import decode_access_token
    from app.core.redis import get_redis
    jwt_token = None
    auth_header = request.headers.get("authorization", "") if request else ""
    if auth_header.lower().startswith("bearer "):
        jwt_token = auth_header[7:]
    elif token:
        jwt_token = token
    if not jwt_token:
        return error(ErrorCode.TOKEN_INVALID, "未提供认证令牌")
    r = await get_redis()
    if await r.get(f"token_blacklist:{jwt_token}"):
        return error(ErrorCode.TOKEN_INVALID, "令牌已失效")
    user_id_str = decode_access_token(jwt_token)
    if not user_id_str:
        return error(ErrorCode.TOKEN_EXPIRED, "令牌已过期或无效")

    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return error(ErrorCode.NOT_FOUND, "公文不存在")

    # 检查是否有源文件
    if not doc.source_file_path:
        return error(ErrorCode.NOT_FOUND, "此公文没有源文件，无法生成 PDF 预览")

    source_path = Path(doc.source_file_path)
    if not _is_safe_upload_path(source_path):
        return error(ErrorCode.PERMISSION_DENIED, "文件路径不合法")
    if not source_path.exists():
        return error(ErrorCode.NOT_FOUND, "源文件已被删除")

    ext = doc.source_format or ""

    # 如果源文件就是 PDF，直接返回
    if ext == "pdf":
        return FileResponse(
            path=str(source_path),
            media_type="application/pdf",
            filename=f"{doc.title}.pdf",
        )

    # 检查是否已有缓存的 PDF
    pdf_cache_path = source_path.parent / "preview.pdf"
    if pdf_cache_path.exists():
        return FileResponse(
            path=str(pdf_cache_path),
            media_type="application/pdf",
            filename=f"{doc.title}.pdf",
        )

    # 调用 converter 微服务转换
    source_bytes = source_path.read_bytes()
    pdf_bytes = await convert_to_pdf_bytes(source_bytes, f"{doc.title}.{ext}")

    if not pdf_bytes:
        raise HTTPException(status_code=502, detail="PDF 转换失败，converter 微服务不可用")

    # 缓存 PDF 到磁盘
    pdf_cache_path.write_bytes(pdf_bytes)

    return FileResponse(
        path=str(pdf_cache_path),
        media_type="application/pdf",
        filename=f"{doc.title}.pdf",
    )


@router.get("/{doc_id}")
async def get_document(
    doc_id: UUID,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """公文详情"""
    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return error(ErrorCode.NOT_FOUND, "公文不存在")

    # 访问控制：非创建者只能查看公开文档
    if doc.creator_id != current_user.id and getattr(doc, 'visibility', 'private') != 'public':
        return error(ErrorCode.PERMISSION_DENIED, "无权访问此文档")

    # 查创建者姓名
    creator_name = ""
    cr = await db.execute(select(User.display_name).where(User.id == doc.creator_id))
    row = cr.scalar_one_or_none()
    if row:
        creator_name = row

    data = {
        **DocumentDetail.model_validate(doc).model_dump(mode="json"),
        "creator_name": creator_name,
        "has_source_file": bool(doc.source_file_path),
        "has_markdown_file": bool(doc.md_file_path),
    }
    return success(data=data)


@router.put("/{doc_id}")
async def update_document(
    doc_id: UUID,
    body: DocumentUpdateRequest,
    request: Request,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """更新公文"""
    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return error(ErrorCode.NOT_FOUND, "公文不存在")
    if doc.creator_id != current_user.id:
        return error(ErrorCode.PERMISSION_DENIED, "只有创建者才能修改公文")

    update_data = body.model_dump(exclude_unset=True)

    # 内容变更时自动保存版本快照
    if "content" in update_data and update_data["content"] != doc.content and doc.content:
        await _save_version(db, doc, current_user.id, change_summary="手动编辑")

    for field, value in update_data.items():
        setattr(doc, field, value)
    await db.flush()

    await log_action(
        db, user_id=current_user.id, user_display_name=current_user.display_name,
        action="更新公文", module="智能公文",
        detail=f"更新公文: {doc.title}, 字段: {list(update_data.keys())}",
        ip_address=request.client.host if request.client else None,
    )

    return success(message="更新成功")


class VisibilityRequest(BaseModel):
    visibility: str


@router.patch("/{doc_id}/visibility")
async def toggle_doc_visibility(
    doc_id: UUID,
    body: VisibilityRequest,
    request: Request,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """切换公文可见性（私密/公开）"""
    if body.visibility not in ("private", "public"):
        return error(ErrorCode.PARAM_ERROR, "visibility 只能是 private 或 public")

    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return error(ErrorCode.NOT_FOUND, "公文不存在")

    # 只有创建者才能修改可见性
    if doc.creator_id != current_user.id:
        return error(ErrorCode.PERMISSION_DENIED, "只有创建者才能修改公文可见性")

    # 设为公开需要 app:doc:public 权限
    if body.visibility == "public":
        from app.models.user import Role, RolePermission
        role_result = await db.execute(select(Role).where(Role.id == current_user.role_id))
        role = role_result.scalar_one_or_none()
        has_perm = False
        if role and role.is_system:
            has_perm = True
        else:
            perm_result = await db.execute(
                select(RolePermission).where(
                    RolePermission.role_id == current_user.role_id,
                    RolePermission.permission_key == "app:doc:public",
                )
            )
            has_perm = perm_result.scalar_one_or_none() is not None
        if not has_perm:
            return error(ErrorCode.PERMISSION_DENIED, "当前角色无发布公开公文权限")

    doc.visibility = body.visibility
    doc.updated_at = datetime.now(timezone.utc)
    await db.flush()

    await log_action(
        db, user_id=current_user.id, user_display_name=current_user.display_name,
        action="修改公文可见性", module="智能公文",
        detail=f"公文「{doc.title}」设为{'公开' if body.visibility == 'public' else '私密'}",
        ip_address=request.client.host if request.client else None,
    )

    return success(message=f"已设为{'公开' if body.visibility == 'public' else '私密'}")


async def _delete_one_document(doc_id: UUID, db: AsyncSession):
    """删除单个公文的文件和数据库记录（内部辅助函数，不做权限校验）"""
    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return None

    title = doc.title

    # 清理本地文件（源文件 + Markdown 文件 + 整个目录）
    doc_upload_dir = Path(settings.UPLOAD_DIR) / "documents" / str(doc_id)
    for path_str in (doc.source_file_path, doc.md_file_path):
        if path_str:
            try:
                Path(path_str).unlink(missing_ok=True)
            except Exception:
                pass
    # 尝试清理空目录
    try:
        if doc_upload_dir.exists():
            import shutil
            shutil.rmtree(doc_upload_dir, ignore_errors=True)
    except Exception:
        pass

    # 删除版本历史
    versions = await db.execute(select(DocumentVersion).where(DocumentVersion.document_id == doc_id))
    for v in versions.scalars().all():
        await db.delete(v)

    await db.delete(doc)
    return title


@router.delete("/{doc_id}")
async def delete_document(
    doc_id: UUID,
    request: Request,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """删除公文（仅创建者可删除）"""
    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return error(ErrorCode.NOT_FOUND, "公文不存在")
    if doc.creator_id != current_user.id:
        return error(ErrorCode.PERMISSION_DENIED, "只能删除自己创建的公文")

    title = await _delete_one_document(doc_id, db)
    await db.flush()

    await log_action(
        db, user_id=current_user.id, user_display_name=current_user.display_name,
        action="删除公文", module="智能公文",
        detail=f"删除公文: {title}",
        ip_address=request.client.host if request.client else None,
    )

    return success(message="删除成功")


class BatchDeleteRequest(BaseModel):
    ids: list[str]


@router.post("/batch-delete")
async def batch_delete_documents(
    body: BatchDeleteRequest,
    request: Request,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """批量删除公文（仅删除属于当前用户的公文）"""
    if not body.ids:
        return error(ErrorCode.PARAM_ERROR, "请选择要删除的公文")

    deleted_titles = []
    skipped = 0
    for id_str in body.ids:
        try:
            doc_id = UUID(id_str)
        except ValueError:
            skipped += 1
            continue
        result = await db.execute(select(Document).where(Document.id == doc_id))
        doc = result.scalar_one_or_none()
        if not doc or doc.creator_id != current_user.id:
            skipped += 1
            continue
        title = await _delete_one_document(doc_id, db)
        if title:
            deleted_titles.append(title)

    await db.flush()

    if deleted_titles:
        await log_action(
            db, user_id=current_user.id, user_display_name=current_user.display_name,
            action="批量删除公文", module="智能公文",
            detail=f"批量删除 {len(deleted_titles)} 篇公文: {', '.join(deleted_titles[:5])}{'...' if len(deleted_titles) > 5 else ''}",
            ip_address=request.client.host if request.client else None,
        )

    msg = f"成功删除 {len(deleted_titles)} 篇公文"
    if skipped:
        msg += f"，跳过 {skipped} 篇（不存在或无权限）"
    return success(message=msg)


@router.post("/{doc_id}/archive")
async def archive_document(
    doc_id: UUID,
    request: Request,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """归档公文"""
    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return error(ErrorCode.NOT_FOUND, "公文不存在")
    if doc.creator_id != current_user.id:
        return error(ErrorCode.PERMISSION_DENIED, "只有创建者才能归档公文")

    doc.status = "archived"
    await db.flush()

    await log_action(
        db, user_id=current_user.id, user_display_name=current_user.display_name,
        action="归档公文", module="智能公文",
        detail=f"归档公文: {doc.title}",
        ip_address=request.client.host if request.client else None,
    )

    return success(message="归档成功")


@router.post("/{doc_id}/process")
async def process_document(
    doc_id: UUID,
    body: DocProcessRequest,
    request: Request,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """AI 公文处理（起草/检查/优化）"""
    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return error(ErrorCode.NOT_FOUND, "公文不存在")
    if doc.creator_id != current_user.id:
        return error(ErrorCode.PERMISSION_DENIED, "只有创建者才能处理公文")

    dify = get_dify_service()

    try:
        if body.process_type == "draft":
            wf_result = await dify.run_doc_draft(
                title=doc.title,
                outline=doc.content or "",
                doc_type=doc.doc_type,
            )
            # 保存旧版本
            if doc.content:
                await _save_version(db, doc, current_user.id, change_type="draft", change_summary="AI起草前版本")

            doc.content = wf_result.output_text
            doc.status = "draft"
            await db.flush()

            return success(data={
                "document_id": str(doc_id),
                "process_type": "draft",
                "content": wf_result.output_text,
                "new_status": "draft",
                "review_result": None,
            })

        elif body.process_type == "check":
            if not doc.content:
                return error(ErrorCode.PARAM_INVALID, "公文内容为空，无法检查")

            review = await dify.run_doc_check(doc.content)

            doc.status = "checked"
            await db.flush()

            return success(data={
                "document_id": str(doc_id),
                "process_type": "check",
                "content": doc.content,
                "new_status": "checked",
                "review_result": {
                    "typos": [{"text": i.text, "suggestion": i.suggestion, "context": i.context} for i in review.typos],
                    "grammar": [{"text": i.text, "suggestion": i.suggestion, "context": i.context} for i in review.grammar],
                    "sensitive": [{"text": i.text, "suggestion": i.suggestion, "context": i.context} for i in review.sensitive],
                },
            })

        elif body.process_type == "optimize":
            if not doc.content:
                return error(ErrorCode.PARAM_INVALID, "公文内容为空，无法优化")

            # 保存旧版本
            await _save_version(db, doc, current_user.id, change_type="optimize", change_summary="AI优化前版本")

            wf_result = await dify.run_doc_optimize(doc.content)
            doc.content = wf_result.output_text
            doc.status = "optimized"
            await db.flush()

            return success(data={
                "document_id": str(doc_id),
                "process_type": "optimize",
                "content": wf_result.output_text,
                "new_status": "optimized",
                "review_result": None,
            })

        elif body.process_type == "format":
            # 格式化处理 — 使用本地 python-docx，不走 Dify
            if not doc.source_file_path:
                return error(ErrorCode.PARAM_INVALID, "该文档无原始文件，无法进行格式化")
            source_path = Path(doc.source_file_path)
            if not source_path.exists():
                return error(ErrorCode.PARAM_INVALID, "原始文件已丢失")
            if source_path.suffix.lower() != '.docx':
                return error(ErrorCode.PARAM_INVALID, f"格式化仅支持 .docx 文件，当前格式: {source_path.suffix}")

            try:
                output_path, stats = DocFormatService.smart_format(
                    str(source_path), preset_name="official"
                )
                # 保存格式化后的文件覆盖原始文件路径
                import shutil
                formatted_dir = source_path.parent / "formatted"
                formatted_dir.mkdir(parents=True, exist_ok=True)
                formatted_path = formatted_dir / source_path.name
                shutil.move(str(output_path), str(formatted_path))

                # 保存版本
                if doc.content:
                    await _save_version(db, doc, current_user.id, change_type="format", change_summary="格式化前版本")

                doc.status = "formatted"
                await db.flush()

                return success(data={
                    "document_id": str(doc_id),
                    "process_type": "format",
                    "content": doc.content,
                    "new_status": "formatted",
                    "review_result": None,
                    "format_stats": stats,
                    "formatted_file": str(formatted_path),
                })
            except Exception as fmt_err:
                return error(ErrorCode.INTERNAL_ERROR, f"格式化处理失败: {str(fmt_err)}")

        else:
            return error(ErrorCode.PARAM_INVALID, f"不支持的处理类型: {body.process_type}")

    except Exception as e:
        return error(ErrorCode.DIFY_ERROR, f"AI处理失败: {str(e)}")

    finally:
        await log_action(
            db, user_id=current_user.id, user_display_name=current_user.display_name,
            action=f"AI公文{body.process_type}", module="智能公文",
            detail=f"{body.process_type} 公文: {doc.title}",
            ip_address=request.client.host if request.client else None,
        )


# ── 对话式 AI 处理（SSE 流式） ──

class AiProcessRequest(BaseModel):
    """对话式 AI 处理请求"""
    stage: str  # draft / check / optimize / format
    user_instruction: str = ""  # 用户对话式指令
    existing_paragraphs: list[dict] | None = None  # 已有格式化段落（增量修改时传入）
    kb_collection_ids: list[UUID] | None = None  # 引用知识库集合 ID（起草时可选）
    confirmed_outline: str | None = None  # #18: 已确认的大纲（两步起草第二步传入）


@router.post("/{doc_id}/ai-process")
async def ai_process_document(
    doc_id: UUID,
    body: AiProcessRequest,
    request: Request,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """
    对话式 AI 公文处理（SSE 流式）。

    每个阶段(draft/check/optimize/format)都有独立的对话输入框，
    用户描述需求后，AI 流式输出处理结果。

    安全机制:
      - Redis 分布式锁：同一文档同一时间只允许一个 AI 处理任务
      - 独立事务：SSE 流内的 DB 写操作使用独立会话并显式 commit
      - 异常保护：客户端断开时安全释放锁和回滚

    SSE 事件格式:
      data: {"type":"status","message":"正在处理..."}
      data: {"type":"text","text":"增量文本片段"}
      data: {"type":"done","full_content":"完整结果文本"}
      data: [DONE]
    """
    _logger = logging.getLogger(__name__)

    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return error(ErrorCode.NOT_FOUND, "公文不存在")
    if doc.creator_id != current_user.id:
        return error(ErrorCode.PERMISSION_DENIED, "只有创建者才能处理公文")

    valid_stages = {"draft", "review", "format", "format_suggest"}
    if body.stage not in valid_stages:
        return error(ErrorCode.PARAM_INVALID, f"stage 必须为 {valid_stages} 之一")

    # ── 并发锁：同一文档同时只允许一个 AI 处理任务 ──
    r = await get_redis()
    lock_key = f"doc_ai_lock:{doc_id}"
    lock_acquired = await r.set(lock_key, f"{current_user.id}:{body.stage}", nx=True, ex=120)
    if not lock_acquired:
        lock_info = await r.get(lock_key)
        return error(
            ErrorCode.CONFLICT,
            f"该文档正在被 AI 处理中，请稍后再试（{lock_info}）",
        )

    try:
        dify = get_dify_service()
    except Exception:
        # get_dify_service 失败时释放锁
        try:
            await r.delete(lock_key)
        except Exception:
            pass
        raise

    def _para_to_dict(p) -> dict:
        """将 StructuredParagraph 转为 SSE 字典，包含富格式属性"""
        d = {"text": p.text, "style_type": p.style_type}
        for key in ("font_size", "font_family", "bold", "italic", "indent", "alignment", "line_height"):
            val = getattr(p, key, None)
            if val is not None:
                d[key] = val
        return d

    def _compute_para_diff(
        old_paras: list[dict] | None,
        new_paras: list[dict],
    ) -> list[dict]:
        """
        Copilot-style 段落级 diff 计算。

        比对 old_paras 与 new_paras：
        - 完全一致 → 无标记
        - old 中有而 new 中无 → 插入 _change=deleted 的段落
        - new 中有而 old 中无 → _change=added
        - 同位置文本不同 → _change=modified + _original_text
        """
        if not old_paras:
            # 全部为新增
            for p in new_paras:
                p["_change"] = "added"
            return new_paras

        # 构建旧段落文本 → 索引映射
        old_texts = [p.get("text", "").strip() for p in old_paras]
        old_used = [False] * len(old_paras)
        result: list[dict] = []

        for new_p in new_paras:
            new_text = new_p.get("text", "").strip()
            # 精确匹配查找
            matched_idx = None
            for i, ot in enumerate(old_texts):
                if not old_used[i] and ot == new_text:
                    matched_idx = i
                    break
            if matched_idx is not None:
                # 完全一致
                old_used[matched_idx] = True
                result.append(new_p)  # 无 _change
            else:
                # 尝试模糊匹配（同位置或文本相似）
                # 按顺序查找第一个未使用且文本有交集的旧段落
                fuzzy_idx = None
                for i, ot in enumerate(old_texts):
                    if old_used[i]:
                        continue
                    # 简单相似度：共同子串 > 30%
                    shorter = min(len(ot), len(new_text))
                    if shorter == 0:
                        continue
                    common = sum(1 for a, b in zip(ot, new_text) if a == b)
                    if common / max(shorter, 1) > 0.3:
                        fuzzy_idx = i
                        break
                if fuzzy_idx is not None:
                    old_used[fuzzy_idx] = True
                    new_p["_change"] = "modified"
                    new_p["_original_text"] = old_texts[fuzzy_idx]
                    result.append(new_p)
                else:
                    new_p["_change"] = "added"
                    result.append(new_p)

        # 未匹配的旧段落 → deleted
        for i, used in enumerate(old_used):
            if not used and old_texts[i]:
                deleted_p = {**old_paras[i], "_change": "deleted"}
                # 插入到结果中合适的位置（尽量保持原始顺序）
                # 简化处理：追加到末尾
                result.append(deleted_p)

        return result

    def _apply_draft_diff(existing_paras: list[dict], changes: list[dict]) -> list[dict]:
        """
        将 AI 输出的增量 diff 变更指令应用到现有段落上。

        支持的 op：
          - {"op": "replace", "index": N, "text": "..."}
          - {"op": "insert_after", "index": N, "text": "...", "style_type": "..."}
          - {"op": "add", "after": N, "text": "...", "style_type": "..."}  (等同 insert_after)
          - {"op": "delete", "index": N}

        返回完整段落列表，变更的段落带 _change / _original_text / _change_reason 标记。
        """
        from collections import defaultdict

        deleted: set[int] = set()
        replaced: dict[int, dict] = {}
        inserts: dict[int, list[dict]] = defaultdict(list)

        for c in changes:
            op = c.get("op", "")
            idx = c.get("index", -999)
            if op == "delete" and 0 <= idx < len(existing_paras):
                deleted.add(idx)
            elif op == "replace" and 0 <= idx < len(existing_paras):
                replaced[idx] = c
            elif op == "insert_after":
                if idx >= -1:
                    inserts[idx].append(c)
            elif op == "add":
                # "add" 使用 "after" 字段（等同 insert_after 使用 "index"）
                after_idx = c.get("after", -999)
                if after_idx == -999:
                    after_idx = c.get("index", -999)
                if after_idx >= -1:
                    inserts[after_idx].append(c)

        result: list[dict] = []

        # 在文档最前面插入（index = -1）
        for ins in inserts.get(-1, []):
            result.append({
                "text": ins.get("text", ""),
                "style_type": ins.get("style_type", "body"),
                "_change": "added",
                "_change_reason": ins.get("reason", "AI 新增"),
            })

        for i, para in enumerate(existing_paras):
            if i in deleted:
                result.append({
                    **para,
                    "_change": "deleted",
                    "_change_reason": "AI 删除",
                })
            elif i in replaced:
                r = replaced[i]
                result.append({
                    **para,
                    "text": r.get("text", para.get("text", "")),
                    "style_type": r.get("style_type", para.get("style_type", "body")),
                    "_change": "modified",
                    "_original_text": para.get("text", ""),
                    "_change_reason": r.get("reason", "AI 修改"),
                })
            else:
                result.append(dict(para))

            # 在这个段落后面插入新段落
            for ins in inserts.get(i, []):
                result.append({
                    "text": ins.get("text", ""),
                    "style_type": ins.get("style_type", "body"),
                    "_change": "added",
                    "_change_reason": ins.get("reason", "AI 新增"),
                })

        return result

    async def event_generator():
        def _sse(data: dict) -> str:
            return f"data: {json.dumps(data, ensure_ascii=False)}\n\n"

        import asyncio
        import time as _time_usage
        _stage_start = _time_usage.time()
        _accumulated_usage: dict = {}  # 累积各阶段的 Dify usage

        def _capture_usage(event_data: dict):
            """从 SSEEvent.data 中提取 usage 信息"""
            nonlocal _accumulated_usage
            usage = event_data.get("usage", {})
            if usage:
                _accumulated_usage["prompt_tokens"] = _accumulated_usage.get("prompt_tokens", 0) + (usage.get("prompt_tokens", 0) or 0)
                _accumulated_usage["completion_tokens"] = _accumulated_usage.get("completion_tokens", 0) + (usage.get("completion_tokens", 0) or 0)
                _accumulated_usage["total_tokens"] = _accumulated_usage.get("total_tokens", 0) + (usage.get("total_tokens", 0) or 0)
                _accumulated_usage["model"] = usage.get("model") or _accumulated_usage.get("model")

        def _record_stage_usage(stage: str, status: str = "success", error_msg: str | None = None):
            """异步记录本次 AI 处理的用量"""
            asyncio.create_task(record_usage(
                user_id=current_user.id,
                user_display_name=current_user.display_name,
                function_type=f"doc_{stage}",
                tokens_input=_accumulated_usage.get("prompt_tokens", 0),
                tokens_output=_accumulated_usage.get("completion_tokens", 0),
                tokens_total=_accumulated_usage.get("total_tokens", 0),
                duration_ms=int((_time_usage.time() - _stage_start) * 1000),
                model_name=_accumulated_usage.get("model"),
                status=status,
                error_message=error_msg,
            ))

        _all_para_data: list[dict] = []  # 排版阶段收集的段落（供断连保存）

        try:
            yield _sse({"type": "status", "message": f"正在执行{_STAGE_NAMES.get(body.stage, body.stage)}..."})

            if body.stage == "draft":
                # 起草 — NDJSON 流式变更模式（实时渲染，不暴露 JSON）
                _logger.info(f"[draft-trace] 开始起草 doc={doc_id}, confirmed_outline={bool(body.confirmed_outline)}, "
                             f"user_instruction_len={len(body.user_instruction or '')}, "
                             f"doc_content_len={len(doc.content or '')}, "
                             f"kb_ids={body.kb_collection_ids}")
                if doc.content:
                    # 使用独立会话保存版本，避免与后续 _safe_update_doc 的版本号死锁
                    await _safe_update_doc(
                        doc.id, save_version_before=True,
                        version_user_id=current_user.id,
                        version_change_type="draft",
                        version_change_summary="AI对话起草前版本",
                    )
                    _logger.info("[draft-trace] _save_version 完成")

                full_text = ""

                # ── 判断是否已有内容（决定"增量修改" vs "新建文档"模式）──
                has_structured = body.existing_paragraphs and len(body.existing_paragraphs) > 0
                _existing_paras: list[dict] | None = None

                if has_structured:
                    # 过滤掉所有文本为空的占位段落
                    _non_empty = [dict(p) for p in body.existing_paragraphs if p.get("text", "").strip()]
                    if _non_empty:
                        _existing_paras = [dict(p) for p in body.existing_paragraphs]
                    else:
                        _logger.info(f"existing_paragraphs 全部为空占位 ({len(body.existing_paragraphs)} 个)，视为新建文档")

                if not _existing_paras and doc.content and doc.content.strip():
                    # 文档已有内容但无结构化段落 → 按行拆分为基础段落
                    _existing_paras = []
                    for _line in doc.content.split("\n"):
                        _line = _line.strip()
                        if _line:
                            _existing_paras.append({"text": _line, "style_type": "body"})

                _has_existing = bool(_existing_paras)

                # ── 意图检测：判断用户是要"局部修改"还是"重写/另起一篇" ──
                _user_instr = (body.user_instruction or "").strip()
                _rewrite_keywords = (
                    "写一份", "写一篇", "起草一份", "起草一篇", "重新写", "重新起草",
                    "另写", "另起草", "改写成", "改写为", "换一篇", "换成",
                    "写一个", "帮我写", "帮我起草", "请写", "请起草",
                    "生成一份", "生成一篇", "撰写一份", "撰写一篇",
                )
                _is_rewrite = _has_existing and any(kw in _user_instr for kw in _rewrite_keywords)
                if _is_rewrite:
                    _logger.info(f"检测到重写意图，切换到新建文档模式: '{_user_instr[:80]}'")
                    _has_existing = False
                    _existing_paras = None

                # 多模态：只有在没有已有内容时，才读取源文件
                draft_file_bytes: bytes | None = None
                draft_file_name: str = ""
                if not _has_existing and doc.source_file_path and _is_safe_upload_path(doc.source_file_path):
                    try:
                        source_path = Path(doc.source_file_path)
                        if source_path.exists():
                            draft_file_bytes = source_path.read_bytes()
                            ext = doc.source_format or source_path.suffix.lstrip(".")
                            draft_file_name = f"{doc.title}.{ext}" if ext else source_path.name
                            _logger.info(f"多模态起草：读取源文件 {source_path.name} ({len(draft_file_bytes)} bytes)")
                    except Exception as e:
                        _logger.warning(f"源文件读取失败，降级为纯文本模式: {e}")

                # ── 知识库检索（起草参考） ── 优化版：检索最相关的完整文档 + 碎片补充
                _kb_context = ""
                _kb_ref_docs: list[dict] = []  # 用于前端显示参考了哪些文档
                if body.kb_collection_ids:
                    _logger.info(f"[draft-trace] 开始知识库检索, kb_ids={body.kb_collection_ids}")
                    import httpx as _httpx
                    # 构建更精准的检索 query: 标题 + 用户指令
                    _title_part = (doc.title or "").strip()
                    _instr_part = (body.user_instruction or "").strip()
                    if _title_part and _instr_part:
                        _kb_query = f"{_title_part} {_instr_part}"
                    else:
                        _kb_query = _title_part or _instr_part
                    if _kb_query:
                        yield _sse({"type": "status", "message": f"正在检索 {len(body.kb_collection_ids)} 个知识库..."})
                        # 查找 dify_dataset_id 和关联 KBFile
                        _coll_result = await db.execute(
                            select(KBCollection)
                            .where(
                                KBCollection.id.in_(body.kb_collection_ids),
                                KBCollection.dify_dataset_id.isnot(None),
                            )
                        )
                        _kb_records = []
                        _coll_map: dict[str, str] = {}  # dify_dataset_id → collection_name
                        _colls_all = _coll_result.scalars().all()
                        for _coll in _colls_all:
                            _coll_map[_coll.dify_dataset_id] = _coll.name

                        # 并行检索所有知识库集合（优化：原逐个顺序检索 → asyncio.gather 并行）
                        async def _retrieve_one_kb(coll):
                            try:
                                _ret_url = f"{settings.DIFY_BASE_URL}/datasets/{coll.dify_dataset_id}/retrieve"
                                _ret_headers = {"Authorization": f"Bearer {settings.DIFY_DATASET_API_KEY}"}
                                _ret_body = {
                                    "query": _kb_query[:500],
                                    "retrieval_model": {
                                        "search_method": "hybrid_search",
                                        "reranking_enable": True,
                                        "reranking_mode": "reranking_model",
                                        "reranking_model": {
                                            "reranking_provider_name": "langgenius/tongyi/tongyi",
                                            "reranking_model_name": "gte-rerank",
                                        },
                                        "top_k": 8,
                                        "score_threshold_enabled": True,
                                        "score_threshold": 0.3,
                                    },
                                }
                                async with _httpx.AsyncClient(timeout=_httpx.Timeout(30.0, connect=5.0)) as _hc:
                                    _ret_resp = await _hc.post(_ret_url, headers=_ret_headers, json=_ret_body)
                                    _records = []
                                    if _ret_resp.status_code < 400:
                                        for _r in _ret_resp.json().get("records", []):
                                            _seg = _r.get("segment", {})
                                            _doc_info = _seg.get("document", {})
                                            _records.append({
                                                "content": _seg.get("content", ""),
                                                "document_name": _doc_info.get("name", ""),
                                                "dify_document_id": _doc_info.get("id", ""),
                                                "dify_dataset_id": coll.dify_dataset_id,
                                                "collection_name": coll.name,
                                                "score": _r.get("score", 0),
                                            })
                                    return _records
                            except Exception as _e:
                                _logger.warning(f"知识库 {coll.name} 检索失败: {_e}")
                                return []

                        _gather_results = await asyncio.gather(
                            *[_retrieve_one_kb(c) for c in _colls_all]
                        )
                        for _recs in _gather_results:
                            _kb_records.extend(_recs)
                        yield _sse({"type": "status", "message": f"知识库检索完成，共 {len(_kb_records)} 条结果"})

                        # 按 score 排序
                        _kb_records.sort(key=lambda x: x.get("score", 0), reverse=True)

                        if _kb_records:
                            # ── 策略：找到最相关文档的完整内容 ──
                            # 1. 从检索结果中提取 top-1 最相关的源文档名
                            _best_rec = _kb_records[0]
                            _best_doc_name = _best_rec.get("document_name", "")
                            _best_dify_doc_id = _best_rec.get("dify_document_id", "")
                            _best_score = _best_rec.get("score", 0)

                            # 2. 尝试从本地 KB 文件中读取完整 Markdown 内容
                            _full_doc_content = ""
                            if _best_doc_name:
                                try:
                                    from app.models.knowledge import KBFile
                                    # 根据 dify_document_id 或文件名查找 KBFile
                                    _kb_file_q = select(KBFile).where(KBFile.status == "indexed")
                                    if _best_dify_doc_id:
                                        _kb_file_q = _kb_file_q.where(KBFile.dify_document_id == _best_dify_doc_id)
                                    else:
                                        _kb_file_q = _kb_file_q.where(KBFile.name.ilike(f"%{_best_doc_name.rsplit('.', 1)[0]}%"))
                                    _kb_file_result = await db.execute(_kb_file_q.limit(1))
                                    _kb_file = _kb_file_result.scalar_one_or_none()
                                    if _kb_file and _kb_file.md_file_path:
                                        _md_path = Path(_kb_file.md_file_path)
                                        if _md_path.exists():
                                            _full_doc_content = _md_path.read_text(encoding="utf-8")
                                            _logger.info(
                                                f"读取最相关文档完整内容: {_best_doc_name} "
                                                f"({len(_full_doc_content)} 字符, score={_best_score:.2f})"
                                            )
                                except Exception as _e:
                                    _logger.warning(f"读取完整文档失败: {_e}")

                            # 3. 构建 KB 上下文：完整最相关文档 + 其他碎片补充
                            _context_parts = []

                            if _full_doc_content:
                                # 完整文档截取前 10000 字符（保留尽可能多的参考内容）
                                _context_parts.append(
                                    f"【最相关参考文档 — 《{_best_doc_name}》(相关度: {_best_score:.2f})】\n"
                                    f"以下是与你要起草的文档最相似的参考范文，请仔细学习其结构、用语和行文风格，"
                                    f"并在起草时参考借鉴：\n\n"
                                    f"{_full_doc_content[:10000]}"
                                )
                                _kb_ref_docs.append({
                                    "name": _best_doc_name,
                                    "score": round(_best_score, 2),
                                    "type": "full_document",
                                    "char_count": len(_full_doc_content),
                                })
                                yield _sse({
                                    "type": "status",
                                    "message": f"找到最相关参考文档：《{_best_doc_name}》(相关度 {_best_score:.0%})，正在参考起草..."
                                })
                            else:
                                yield _sse({"type": "status", "message": f"检索到 {len(_kb_records)} 条相关参考片段"})

                            # 补充其他相关片段（去重：排除已作为完整文档引入的内容）
                            _seen_doc_names = {_best_doc_name} if _full_doc_content else set()
                            _extra_parts = []
                            for _i, _rec in enumerate(_kb_records[:10], 1):
                                _rec_doc_name = _rec.get("document_name", "")
                                if _rec_doc_name in _seen_doc_names and _full_doc_content:
                                    continue  # 已有完整文档，跳过其片段
                                _extra_parts.append(
                                    f"[{_i}] 来源: {_rec_doc_name} "
                                    f"(集合: {_rec.get('collection_name', '未知')}, 相关度: {_rec.get('score', 0):.2f})\n"
                                    f"{_rec['content']}"
                                )
                                if _rec_doc_name and _rec_doc_name not in _seen_doc_names:
                                    _kb_ref_docs.append({
                                        "name": _rec_doc_name,
                                        "score": round(_rec.get("score", 0), 2),
                                        "type": "segment",
                                    })
                                    _seen_doc_names.add(_rec_doc_name)

                            if _extra_parts:
                                _context_parts.append(
                                    "【其他参考片段】\n" + "\n\n".join(_extra_parts)
                                )

                            _kb_context = "\n\n".join(_context_parts)
                            _logger.info(
                                f"知识库检索完成: {len(_kb_records)} 条结果, "
                                f"完整文档={'有' if _full_doc_content else '无'}, "
                                f"context={len(_kb_context)} 字符, "
                                f"参考文档: {[d['name'] for d in _kb_ref_docs]}"
                            )

                            # 向前端推送参考文档信息
                            if _kb_ref_docs:
                                yield _sse({
                                    "type": "kb_references",
                                    "references": _kb_ref_docs,
                                })

                # ── #18: 大纲两步流程：新建文档且无已确认大纲时，先生成大纲 ──
                _outline_confirmed = body.confirmed_outline
                _logger.info(f"[draft-trace] 大纲检查: has_existing={_has_existing}, "
                             f"outline_confirmed={bool(_outline_confirmed)}, kb_context_len={len(_kb_context)}")
                if not _has_existing and not _outline_confirmed:
                    # 第一步：生成大纲（不生成正文）
                    _outline_instruction = (body.user_instruction or f"请起草一份{doc.doc_type}文档。")
                    if doc.title:
                        _outline_instruction = f"[文档标题]: {doc.title}\n\n[起草要求]: {_outline_instruction}"
                    if _kb_context:
                        _outline_instruction += f"\n\n[参考资料]:\n{_kb_context[:6000]}"
                    _outline_instruction += (
                        '\n\n【输出格式 — 最高优先级】\n'
                        '请只生成文档的大纲结构，不要生成正文内容。\n'
                        '大纲格式要求：\n'
                        '# 文档标题\n\n'
                        '## 一、第一部分标题\n'
                        '- 要点1\n'
                        '- 要点2\n\n'
                        '## 二、第二部分标题\n'
                        '- 要点1\n\n'
                        '...以此类推。\n\n'
                        '用 ## 表示一级标题（中文编号），- 表示该部分的主要内容要点。\n'
                        '⚠️ 只输出大纲结构，不要展开任何正文！\n'
                    )
                    yield _sse({"type": "status", "message": "正在生成文档大纲…"})
                    _outline_text = ""
                    _outline_error = False
                    _MAX_OUTLINE_RETRIES = 2
                    for _outline_attempt in range(_MAX_OUTLINE_RETRIES):
                        _outline_text = ""
                        _outline_error = False
                        if _outline_attempt > 0:
                            yield _sse({"type": "status", "message": f"AI 服务响应超时，正在重试… ({_outline_attempt + 1}/{_MAX_OUTLINE_RETRIES})"})
                            await asyncio.sleep(2)
                        async for sse_event in dify.run_doc_draft_stream(
                            title=doc.title,
                            outline="",
                            doc_type=doc.doc_type,
                            user_instruction=_outline_instruction,
                            file_bytes=draft_file_bytes,
                            file_name=draft_file_name,
                        ):
                            if sse_event.event == "text_chunk":
                                _outline_text += sse_event.data.get("text", "")
                                yield _sse({"type": "text", "text": sse_event.data.get("text", "")})
                            elif sse_event.event == "reasoning":
                                yield _sse({"type": "reasoning", "delta": sse_event.data.get("delta", ""), "text": sse_event.data.get("text", ""), "partial": sse_event.data.get("partial", False)})
                            elif sse_event.event == "progress":
                                yield _sse({"type": "status", "message": sse_event.data.get("message", "AI 正在思考…")})
                            elif sse_event.event == "message_end":
                                _outline_text = sse_event.data.get("full_text", "") or _outline_text
                                _capture_usage(sse_event.data)
                            elif sse_event.event == "error":
                                _outline_error = True
                                _logger.warning(f"大纲生成失败(第{_outline_attempt+1}次): {sse_event.data.get('message', '')}")
                        if _outline_text.strip() and not _outline_error:
                            break  # 成功
                    if not _outline_text.strip():
                        yield _sse({"type": "error", "message": "AI 服务暂时无法响应，请稍后重试"})
                        yield "data: [DONE]\n\n"
                        return
                    # 发送大纲事件，等待前端确认
                    yield _sse({"type": "outline", "outline_text": _outline_text.strip()})
                    yield _sse({"type": "done", "full_content": doc.content or ""})
                    _record_stage_usage("draft_outline")
                    yield "data: [DONE]\n\n"
                    return

                # 如果用户确认了大纲，将大纲嵌入起草指令
                if _outline_confirmed:
                    _logger.info(f"使用已确认大纲起草（{len(_outline_confirmed)} 字符）")

                # ── 构造起草指令（Markdown 纯文本 / 行标记指令） ──
                draft_instruction = body.user_instruction or ""

                _MD_FORMAT = (
                    '\n\n【输出格式 — 最高优先级，必须严格遵守】\n'
                    '请直接输出公文的完整正文内容，使用 Markdown 格式。\n'
                    '一篇完整公文通常包含标题、主送单位、正文各段、结束语、署名和日期。\n'
                    '要求：\n'
                    '1. 标题用 # 开头（仅一个 #），居中\n'
                    '2. 一级标题用中文编号（一、二、三、）\n'
                    '3. 二级标题用（一）（二）（三）\n'
                    '4. 三级标题用 1. 2. 3.\n'
                    '5. 四级标题用 (1) (2) (3)\n'
                    '6. 正文段落直接写，首行缩进由系统处理\n'
                    '7. 主送单位格式：XX单位：（以冒号结尾）\n'
                    '8. 结束语如"特此通知。"独立成段\n'
                    '9. 署名和日期分别独立成段\n'
                    '信息不足时，只输出一行: [NEED_INFO] 请提供XX信息\n'
                    '⚠️ 只输出公文正文，不要输出任何解释、说明或代码块包裹！'
                )

                _DIFF_FORMAT = (
                    '\n\n【输出格式 — 最高优先级，必须严格遵守】\n'
                    '你必须使用行标记指令格式输出变更，每条指令占一行：\n'
                    '替换段落: [REPLACE:段落编号|style:样式] 修改后的完整文本\n'
                    '新增段落: [ADD:after=段落编号|style:样式] 新段落文本\n'
                    '删除段落: [DELETE:段落编号]\n'
                    '信息不足: [NEED_INFO] 请提供XX信息\n\n'
                    '段落编号为 0-based。只输出需要修改的段落。\n'
                    'style 可选: title, recipient, heading1, heading2, heading3, heading4, '
                    'body, closing, signature, date, attachment\n'
                    'style 为可选参数，如果不需要修改样式可省略。\n\n'
                    '【示例1】用户要求"将XX替换为50台"：\n'
                    '[REPLACE:7] 目前共有电脑50台。\n\n'
                    '【示例2】用户要求"在第5段后新增一段"：\n'
                    '[ADD:after=5|style:body] 新增的段落内容。\n\n'
                    '【示例3】用户要求"删掉第5段"：\n'
                    '[DELETE:5]\n\n'
                    '⚠️ 只输出行标记指令，不要输出任何解释文字！'
                )

                if _has_existing:
                    # ── 增量修改模式（行标记指令） ──
                    _MAX_PARA_PREVIEW = 200
                    _compact_lines = []
                    _total = len(_existing_paras)

                    if _total > 80:
                        for _i in range(min(15, _total)):
                            _text = _existing_paras[_i].get("text", "")[:_MAX_PARA_PREVIEW]
                            _st = _existing_paras[_i].get("style_type", "body")
                            _compact_lines.append(f"[{_i}]({_st}) {_text}")
                        _compact_lines.append(f"  ... (中间省略 {_total - 30} 个段落) ...")
                        for _i in range(max(_total - 15, 15), _total):
                            _text = _existing_paras[_i].get("text", "")[:_MAX_PARA_PREVIEW]
                            _st = _existing_paras[_i].get("style_type", "body")
                            _compact_lines.append(f"[{_i}]({_st}) {_text}")
                    else:
                        for _i, _p in enumerate(_existing_paras):
                            _text = _p.get("text", "")
                            if len(_text) > _MAX_PARA_PREVIEW:
                                _text = _text[:_MAX_PARA_PREVIEW] + "…"
                            _st = _p.get("style_type", "body")
                            _compact_lines.append(f"[{_i}]({_st}) {_text}")

                    _compact_listing = "\n".join(_compact_lines)
                    _user_req = draft_instruction or "请在此基础上优化文字内容。"

                    draft_instruction = (
                        '你是文档修改专家，必须严格执行用户要求。\n\n'
                        f'以下是待修改的文档（共 {_total} 个段落）：\n'
                        f'{_compact_listing}\n\n'
                        '─────────────────\n'
                        f'⚠️ 用户要求（必须严格执行）：{_user_req}\n\n'
                        '请仔细检查每一个段落，所有符合用户修改条件的段落都必须输出对应指令。'
                        + _DIFF_FORMAT
                    )
                    if _kb_context:
                        draft_instruction += (
                            '\n\n【参考资料 — 可结合以下知识库内容进行修改】\n'
                            f'{_kb_context[:8000]}'
                        )
                else:
                    # ── 新建文档模式（Markdown 纯文本） ──
                    _user_req = draft_instruction or "请起草公文。"
                    # #18: 嵌入已确认的大纲
                    _outline_section = ""
                    if _outline_confirmed:
                        _outline_section = (
                            '\n\n【已确认的文档大纲 — 请严格按照此结构展开正文】\n'
                            f'{_outline_confirmed}\n\n'
                            '请严格按照上述大纲的标题和要点展开完整正文，'
                            '不要改变大纲中的章节标题和结构。\n'
                        )
                    if _kb_context:
                        draft_instruction = (
                            '【参考资料 — 请务必结合以下知识库内容进行起草，学习其结构和用语风格】\n'
                            f'{_kb_context[:12000]}\n\n'
                            f'{_outline_section}'
                            f'【起草要求】\n{_user_req}'
                            + _MD_FORMAT
                        )
                    else:
                        draft_instruction = _user_req + _outline_section + _MD_FORMAT

                # ── 流式接收 + Markdown/行标记解析 ──
                _logger.info(f"起草模式: has_existing={_has_existing}, instruction_len={len(draft_instruction)}")
                if len(draft_instruction) < 1000:
                    _logger.info(f"起草指令: {repr(draft_instruction)}")
                else:
                    _logger.info(f"起草指令(前500): {repr(draft_instruction[:500])}")

                _acc_text = ""          # 累积全部 LLM 输出文本
                _last_newline_pos = 0   # 上次解析到的换行位置
                _streamed_paras: list[dict] = []   # 新建模式已推送段落
                _parsed_cmds: list[dict] = []      # 增量模式已解析指令
                import time as _time_mod
                _last_progress_ts = _time_mod.monotonic()
                _is_needs_more_info = False
                _prev_content = doc.content
                _completion_tokens = 0   # Dify 返回的实际输出 token 数

                yield _sse({"type": "status", "message": "正在连接 AI 起草服务…"})
                _logger.info(f"[draft-trace] 准备调用 Dify, outline_for_dify_len={len('' if _has_existing else (doc.content or ''))}")

                # 增量修改模式下 draft_instruction 已包含完整段落列表，
                # 不要再传 outline 以避免重复内容干扰 LLM
                _outline_for_dify = "" if _has_existing else (doc.content or "")

                # ── Markdown 续写状态（仅新建模式） ──
                _MAX_CONTINUATION_ROUNDS = 3
                _conversation_id = ""
                _round_error = False

                # 新建模式实时解析状态
                _md_has_title = False
                _md_has_closing = False
                _md_has_signature = False

                for _round_num in range(_MAX_CONTINUATION_ROUNDS):
                    if _round_num > 0:
                        yield _sse({"type": "status", "message": f"文档内容较长，正在续写第 {_round_num + 1} 轮…（已生成 {len(_streamed_paras)} 段落）"})
                        # 续写指令
                        _last_texts = [p.get("text", "")[:150] for p in _streamed_paras[-3:]] if _streamed_paras else []
                        _continuation_instruction = (
                            f"请继续写。上文已输出 {len(_streamed_paras)} 个段落，"
                            f"最后一段是：{_last_texts[-1] if _last_texts else ''}\n\n"
                            f"请从此处继续输出后续内容（Markdown 格式），确保文档有完整的结尾"
                            f"（结束语、署名、日期）。不要重复已输出的内容。"
                        )
                        _acc_text = ""
                        _last_newline_pos = 0

                    _current_instruction = _continuation_instruction if _round_num > 0 else draft_instruction

                    # ── 带重试的 Dify 调用（首轮支持重试，续写轮不重试） ──
                    _MAX_DRAFT_RETRIES = 2 if _round_num == 0 else 1
                    _draft_stream_ok = False
                    for _draft_attempt in range(_MAX_DRAFT_RETRIES):
                        if _draft_attempt > 0:
                            yield _sse({"type": "status", "message": f"AI 服务响应超时，正在重试… ({_draft_attempt + 1}/{_MAX_DRAFT_RETRIES})"})
                            await asyncio.sleep(2)
                        _draft_had_error = False
                        async for sse_event in dify.run_doc_draft_stream(
                            title=doc.title,
                            outline=_outline_for_dify if _round_num == 0 else "",
                            doc_type=doc.doc_type,
                            user_instruction=_current_instruction,
                            file_bytes=draft_file_bytes if _round_num == 0 else None,
                            file_name=draft_file_name if _round_num == 0 else "",
                            conversation_id=_conversation_id if _round_num > 0 else "",
                        ):
                            if sse_event.event == "text_chunk":
                                _chunk_text = sse_event.data.get("text", "")
                                _acc_text += _chunk_text

                                if not _has_existing:
                                    # ── 新建模式：逐行解析 Markdown，实时推送段落 ──
                                    while '\n' in _acc_text[_last_newline_pos:]:
                                        _nl_idx = _acc_text.index('\n', _last_newline_pos)
                                        _line = _acc_text[_last_newline_pos:_nl_idx].strip()
                                        _last_newline_pos = _nl_idx + 1

                                        if not _line:
                                            continue

                                        # 清除 Markdown 符号并识别 style
                                        _clean = _strip_markdown_inline(_line)
                                        if not _clean:
                                            continue

                                        _total_so_far = len(_streamed_paras)
                                        _style = _detect_line_style(
                                            _clean, _total_so_far, _total_so_far + 10,
                                            _md_has_title, _md_has_closing, _md_has_signature,
                                        )
                                        if _style == "title":
                                            _md_has_title = True
                                        elif _style == "closing":
                                            _md_has_closing = True
                                        elif _style == "signature":
                                            _md_has_signature = True

                                        _para = {"text": _clean, "style_type": _style}
                                        _streamed_paras.append(_para)
                                        yield _sse({"type": "structured_paragraph", "paragraph": _para})

                                else:
                                    # ── 增量模式：逐行解析行标记指令 ──
                                    while '\n' in _acc_text[_last_newline_pos:]:
                                        _nl_idx = _acc_text.index('\n', _last_newline_pos)
                                        _line = _acc_text[_last_newline_pos:_nl_idx].strip()
                                        _last_newline_pos = _nl_idx + 1

                                        if not _line:
                                            continue

                                        _line_cmds = _parse_line_diff_commands(_line)
                                        for _cmd in _line_cmds:
                                            if _cmd.get("op") == "need_info":
                                                _is_needs_more_info = True
                                                yield _sse({"type": "needs_more_info", "suggestions": [_cmd.get("text", "请提供更详细的指令。")]})
                                            else:
                                                _parsed_cmds.append(_cmd)

                                # 定期进度
                                _now = _time_mod.monotonic()
                                if _now - _last_progress_ts >= 2.0:
                                    if _has_existing:
                                        yield _sse({"type": "status", "message": f"AI 正在分析变更…（已解析 {len(_parsed_cmds)} 条指令）"})
                                    else:
                                        _round_label = f"（第 {_round_num + 1} 轮）" if _round_num > 0 else ""
                                        yield _sse({"type": "status", "message": f"正在生成文档{_round_label}…（已完成 {len(_streamed_paras)} 个段落）"})
                                    _last_progress_ts = _now

                            elif sse_event.event == "message_end":
                                full_text = sse_event.data.get("full_text", "") or _acc_text
                                _conversation_id = sse_event.data.get("conversation_id", "") or _conversation_id
                                _capture_usage(sse_event.data)

                                # 提取 completion_tokens 用于续写判断
                                _usage = sse_event.data.get("usage", {})
                                _completion_tokens = _usage.get("completion_tokens", 0) or 0

                                _logger.info(
                                    f"起草流结束(第{_round_num+1}轮): "
                                    f"acc_text={len(_acc_text)} chars, "
                                    f"completion_tokens={_completion_tokens}, "
                                    f"paras={len(_streamed_paras)}, "
                                    f"cmds={len(_parsed_cmds)}, "
                                    f"has_existing={_has_existing}"
                                )
                                if len(_acc_text) < 2000:
                                    _logger.info(f"起草AI完整输出(第{_round_num+1}轮): {repr(_acc_text)}")
                                else:
                                    _logger.info(f"起草AI输出(第{_round_num+1}轮,前500): {repr(_acc_text[:500])}")

                                # ── 处理最后一行（可能没有 \n 结尾） ──
                                _remaining = _acc_text[_last_newline_pos:].strip()
                                if _remaining:
                                    if not _has_existing:
                                        _clean = _strip_markdown_inline(_remaining)
                                        if _clean:
                                            _total_so_far = len(_streamed_paras)
                                            _style = _detect_line_style(
                                                _clean, _total_so_far, _total_so_far + 1,
                                                _md_has_title, _md_has_closing, _md_has_signature,
                                            )
                                            if _style == "title":
                                                _md_has_title = True
                                            elif _style == "closing":
                                                _md_has_closing = True
                                            elif _style == "signature":
                                                _md_has_signature = True
                                            _para = {"text": _clean, "style_type": _style}
                                            _streamed_paras.append(_para)
                                            yield _sse({"type": "structured_paragraph", "paragraph": _para})
                                    else:
                                        _line_cmds = _parse_line_diff_commands(_remaining)
                                        for _cmd in _line_cmds:
                                            if _cmd.get("op") == "need_info":
                                                _is_needs_more_info = True
                                                yield _sse({"type": "needs_more_info", "suggestions": [_cmd.get("text", "")]})
                                            else:
                                                _parsed_cmds.append(_cmd)

                            elif sse_event.event == "reasoning":
                                yield _sse({"type": "reasoning", "delta": sse_event.data.get("delta", ""), "text": sse_event.data.get("text", ""), "partial": sse_event.data.get("partial", False)})
                            elif sse_event.event == "progress":
                                yield _sse({"type": "status", "message": sse_event.data.get("message", "生成中…")})
                            elif sse_event.event == "error":
                                _draft_had_error = True
                                if _draft_attempt == _MAX_DRAFT_RETRIES - 1:
                                    yield _sse({"type": "error", "message": sse_event.data.get("message", "起草失败")})
                                break

                        # ── 重试判断 ──
                        if not _draft_had_error:
                            break  # 成功，退出重试循环
                        _logger.warning(f"起草第 {_draft_attempt + 1}/{_MAX_DRAFT_RETRIES} 次尝试失败")

                    _round_error = _draft_had_error
                    if _round_error:
                        return

                    _logger.info(f"[draft-trace] Dify流完成, 进入续写/保存判断 (has_existing={_has_existing}, paras={len(_streamed_paras)})")
                    # ── Markdown 续写判断（仅新建文档模式） ──
                    if not _has_existing and not _is_needs_more_info:
                        # 判断是否输出被截断：completion_tokens 接近上限 + 文档结构不完整
                        _doc_complete = _md_has_closing or _md_has_signature
                        _token_near_limit = _completion_tokens >= settings.DRAFT_MAX_COMPLETION_TOKENS
                        _has_new_content = len(_streamed_paras) > 0

                        if _token_near_limit and not _doc_complete and _has_new_content and _conversation_id:
                            _logger.info(
                                f"起草第 {_round_num + 1} 轮输出可能被截断 "
                                f"(tokens={_completion_tokens}, paras={len(_streamed_paras)}, "
                                f"has_closing={_md_has_closing}, has_signature={_md_has_signature})，"
                                f"将发起第 {_round_num + 2} 轮续写"
                            )
                            continue  # 进入下一轮续写
                        elif _token_near_limit and not _has_new_content:
                            _logger.warning(
                                f"起草第 {_round_num + 1} 轮 token 接近上限但无新内容，"
                                f"终止续写以防死循环"
                            )
                    break  # 输出完整或增量模式不续写，退出循环

                if _round_num > 0 and not _round_error:
                    _logger.info(f"起草多轮续写完成：共 {_round_num + 1} 轮，最终 {len(_streamed_paras)} 段落")

                # ── 结果处理 & 保存 ──
                if _is_needs_more_info:
                    yield _sse({"type": "done", "needs_more_info": True})

                elif _has_existing and _parsed_cmds:
                    # ── 增量模式：应用行标记变更到现有段落 ──
                    _applied = _apply_draft_diff(_existing_paras, _parsed_cmds)
                    _plain = "\n".join(
                        p.get("text", "") for p in _applied
                        if p.get("_change") != "deleted"
                    )
                    await _safe_update_doc(doc.id, {"content": _plain, "status": "draft"})
                    # 保存"AI起草完成"版本快照
                    await _safe_update_doc(
                        doc.id, save_version_before=True,
                        version_user_id=current_user.id,
                        version_change_type="draft",
                        version_change_summary="AI起草完成（增量修改）",
                    )

                    _change_count = len(_parsed_cmds)
                    _logger.info(f"起草阶段(diff)：成功应用 {_change_count} 处变更")
                    yield _sse({
                        "type": "draft_result",
                        "paragraphs": _applied,
                        "summary": f"共 {_change_count} 处变更",
                        "change_count": _change_count,
                    })
                    yield _sse({"type": "done", "full_content": _plain})

                elif _has_existing and not _parsed_cmds:
                    # ── 增量模式但未解析出指令 → JSON 降级兜底 ──
                    _logger.warning(f"增量模式未解析出行标记指令，尝试 JSON 降级 (acc={len(_acc_text)} chars)")
                    _fallback_text = full_text or _acc_text
                    _fallback_done = False

                    # 尝试 JSON 整体解析（向后兼容旧 prompt）
                    _stripped = _fallback_text.strip()
                    if "```" in _stripped:
                        _stripped = _stripped.split("```json")[-1].split("```")[0].strip() if "```json" in _stripped else _stripped.split("```")[1].split("```")[0].strip() if _stripped.count("```") >= 2 else _stripped
                    if _stripped.startswith("{") and _stripped.endswith("}"):
                        try:
                            from json_repair import loads as jr_loads
                            _parsed = jr_loads(_stripped)
                            if isinstance(_parsed, dict):
                                _ai_paras = _parsed.get("paragraphs", _parsed.get("changes", []))
                                if isinstance(_ai_paras, list) and _ai_paras:
                                    _applied = _apply_draft_diff(_existing_paras, _ai_paras)
                                    _plain = "\n".join(p.get("text", "") for p in _applied if p.get("_change") != "deleted")
                                    await _safe_update_doc(doc.id, {"content": _plain, "status": "draft"})
                                    # 保存"AI起草完成"版本快照
                                    await _safe_update_doc(
                                        doc.id, save_version_before=True,
                                        version_user_id=current_user.id,
                                        version_change_type="draft",
                                        version_change_summary="AI起草完成（JSON降级）",
                                    )
                                    yield _sse({"type": "draft_result", "paragraphs": _applied, "summary": "", "change_count": len(_ai_paras)})
                                    _fallback_done = True
                        except Exception as _e:
                            _logger.debug(f"增量降级 JSON 解析失败: {_e}")

                    if not _fallback_done:
                        yield _sse({"type": "needs_more_info", "suggestions": ["AI未能生成修改内容，请尝试更具体的修改要求。"]})
                        yield _sse({"type": "done", "needs_more_info": True})

                elif not _has_existing and _streamed_paras:
                    # ── 新文档模式：段落已实时推送 ──
                    _plain = "\n".join(p.get("text", "") for p in _streamed_paras)
                    # 自动提取标题：从第一个 style_type=title 的段落提取
                    _auto_title = ""
                    for _p in _streamed_paras:
                        if _p.get("style_type") == "title" and _p.get("text", "").strip():
                            _auto_title = _p["text"].strip()
                            break
                    _should_rename = _auto_title and (not doc.title or doc.title in ("新建公文", "新建文档", ""))
                    if _should_rename:
                        _logger.info(f"起草自动命名: '{_auto_title}'")
                    # ── 先发送 done 事件，确保前端立即收到完成通知和标题 ──
                    _done_data: dict = {"type": "done", "full_content": _plain}
                    if _should_rename:
                        _done_data["new_title"] = _auto_title
                    if _round_num > 0:
                        _done_data["continuation_rounds"] = _round_num + 1
                    _logger.info(f"[draft-trace] 发送done事件: new_title={'有' if 'new_title' in _done_data else '无'}, keys={list(_done_data.keys())}")
                    yield _sse(_done_data)
                    # ── 然后保存到 DB（不阻塞前端 done 通知） ──
                    _logger.info(f"[draft-trace] 保存起草结果到DB... ({len(_plain)} chars, {len(_streamed_paras)} paras)")
                    _update_fields: dict = {"content": _plain, "status": "draft"}
                    if _should_rename:
                        _update_fields["title"] = _auto_title
                    await _safe_update_doc(doc.id, _update_fields)
                    _logger.info("[draft-trace] 内容保存完成，开始保存版本快照...")
                    await _safe_update_doc(
                        doc.id, save_version_before=True,
                        version_user_id=current_user.id,
                        version_change_type="draft",
                        version_change_summary="AI起草完成（新建文档）",
                    )
                    _logger.info("[draft-trace] 版本快照保存完成")

                elif not _has_existing and not _streamed_paras:
                    # ── 新建模式但无段落 → Markdown 整体兜底解析 ──
                    _logger.warning(f"新建模式未实时解析出段落，尝试整体解析 (acc={len(_acc_text)} chars)")
                    _fallback_text = (full_text or _acc_text).strip()

                    # 去除可能的代码块包裹
                    if _fallback_text.startswith("```") and _fallback_text.endswith("```"):
                        _fallback_text = _fallback_text.strip("`").strip()
                        if _fallback_text.startswith("markdown\n"):
                            _fallback_text = _fallback_text[len("markdown\n"):]

                    if _fallback_text:
                        # 尝试 Markdown 解析
                        _streamed_paras = _parse_markdown_to_paragraphs(_fallback_text)
                        if _streamed_paras:
                            for _p in _streamed_paras:
                                yield _sse({"type": "structured_paragraph", "paragraph": _p})
                            _plain = "\n".join(p["text"] for p in _streamed_paras)
                            await _safe_update_doc(doc.id, {"content": _plain, "status": "draft"})
                            await _safe_update_doc(
                                doc.id, save_version_before=True,
                                version_user_id=current_user.id,
                                version_change_type="draft",
                                version_change_summary="AI起草完成（Markdown解析）",
                            )
                            yield _sse({"type": "done", "full_content": _plain})
                        else:
                            # 纯文本兜底
                            await _safe_update_doc(doc.id, {"content": _fallback_text, "status": "draft"})
                            await _safe_update_doc(
                                doc.id, save_version_before=True,
                                version_user_id=current_user.id,
                                version_change_type="draft",
                                version_change_summary="AI起草完成（纯文本）",
                            )
                            yield _sse({"type": "replace_streaming_text", "text": _fallback_text})
                            yield _sse({"type": "done", "full_content": _fallback_text})
                    else:
                        yield _sse({"type": "done", "full_content": doc.content or ""})

                _record_stage_usage("draft")

            elif body.stage == "review":
                # 审查&优化（合并版） — 流式调用，逐条推送建议
                # 审查始终基于前端当前内容（existing_paragraphs 优先，否则 doc.content），不上传源文件
                has_structured = body.existing_paragraphs and len(body.existing_paragraphs) > 0

                if not has_structured and not doc.content:
                    yield _sse({"type": "error", "message": "公文内容为空，无法审查"})
                    yield "data: [DONE]\n\n"
                    return

                await _safe_update_doc(
                    doc.id, save_version_before=True,
                    version_user_id=current_user.id,
                    version_change_type="review",
                    version_change_summary="AI审查优化前版本",
                )

                # 审查内容：优先用前端当前结构化段落的文本，否则用 doc.content
                review_content = doc.content or ""
                review_instruction = body.user_instruction or ""
                if has_structured:
                    # 按段落拼接文本，保留段落边界（便于 Dify 返回的 original 精确匹配）
                    _text_lines = []
                    for _p in body.existing_paragraphs:
                        _text = _p.get("text", "").strip()
                        if _text:
                            _text_lines.append(_text)
                    review_content = "\n\n".join(_text_lines)

                _logger.info(f"审查优化：内容长度 {len(review_content)} 字符, 结构化={has_structured}")

                # ── 长文档分块审查（>8000 字符时自动分块） ──
                _MAX_REVIEW_CHUNK = 8000
                all_suggestions: list[dict] = []
                all_summaries: list[str] = []

                if len(review_content) > _MAX_REVIEW_CHUNK:
                    review_chunks = _split_text_into_chunks(review_content, _MAX_REVIEW_CHUNK)
                    _logger.info(f"审查分块: {len(review_content)} 字符 → {len(review_chunks)} 块")

                    # Phase-1：规则化结构分析，为审查提供全局大纲
                    _review_analysis: dict = {}
                    try:
                        _review_analysis = _analyze_doc_structure(review_content)
                    except Exception:
                        pass

                    for chunk_idx, chunk_text in enumerate(review_chunks):
                        if len(review_chunks) > 1:
                            yield _sse({"type": "status", "message": f"正在审查第 {chunk_idx+1}/{len(review_chunks)} 部分…"})

                        chunk_instr = review_instruction
                        if chunk_idx > 0 or (_review_analysis and len(review_chunks) > 1):
                            _outline = ""
                            if _review_analysis:
                                _rtp = _review_analysis.get("total_paragraphs", 0)
                                _rtc = len(review_content) or 1
                                _cum = sum(len(review_chunks[j]) for j in range(chunk_idx))
                                _s_ratio = _cum / _rtc
                                _e_ratio = (_cum + len(chunk_text)) / _rtc
                                _pr = (int(_s_ratio * _rtp), min(int(_e_ratio * _rtp), _rtp - 1))
                                _outline = _build_outline_context(_review_analysis, len(review_chunks), chunk_idx + 1, _pr)
                            if chunk_idx > 0:
                                chunk_instr = (
                                    f"（续：这是文档的第 {chunk_idx+1}/{len(review_chunks)} 部分）\n"
                                    + (_outline + "\n" if _outline else "")
                                    + review_instruction
                                )
                            elif _outline:
                                chunk_instr = _outline + "\n" + review_instruction

                        # ── 单块重试（最多 2 次） ──
                        _chunk_ok = False
                        for _retry in range(2):
                            try:
                                async for sse_event in dify.run_doc_review_stream(
                                    content=chunk_text,
                                    user_instruction=chunk_instr,
                                ):
                                    if sse_event.event == "review_suggestion":
                                        all_suggestions.append(sse_event.data)
                                        yield _sse({
                                            "type": "review_suggestion",
                                            "suggestion": sse_event.data,
                                        })
                                    elif sse_event.event == "review_result":
                                        _capture_usage(sse_event.data)
                                        chunk_suggestions = sse_event.data.get("suggestions", [])
                                        chunk_summary = sse_event.data.get("summary", "")
                                        for s in chunk_suggestions:
                                            if s not in all_suggestions:
                                                all_suggestions.append(s)
                                                yield _sse({"type": "review_suggestion", "suggestion": s})
                                        if chunk_summary:
                                            all_summaries.append(f"[第{chunk_idx+1}部分] {chunk_summary}")
                                        _logger.info(f"审查分块 {chunk_idx+1}/{len(review_chunks)}: {len(chunk_suggestions)} 条建议")
                                    elif sse_event.event == "reasoning":
                                        yield _sse({"type": "reasoning", "delta": sse_event.data.get("delta", ""), "text": sse_event.data.get("text", ""), "partial": sse_event.data.get("partial", False)})
                                    elif sse_event.event == "progress":
                                        yield _sse({"type": "status", "message": sse_event.data.get("message", "审查中…")})
                                    elif sse_event.event == "error":
                                        raise RuntimeError(sse_event.data.get("message", "Dify 审查错误"))
                                _chunk_ok = True
                                break  # 成功，跳出重试循环
                            except Exception as _chunk_err:
                                if _retry < 1:
                                    _logger.warning(f"审查分块 {chunk_idx+1} 第 {_retry+1} 次失败，重试: {_chunk_err}")
                                    yield _sse({"type": "status", "message": f"第 {chunk_idx+1} 部分审查失败，正在重试…"})
                                    import asyncio as _aio
                                    await _aio.sleep(1)
                                else:
                                    _logger.error(f"审查分块 {chunk_idx+1} 重试用尽: {_chunk_err}")
                        if not _chunk_ok:
                            yield _sse({"type": "status", "message": f"第 {chunk_idx+1} 部分审查失败，已跳过"})

                    # 合并所有分块结果
                    combined_summary = "\n".join(all_summaries) if all_summaries else "审查完成"
                    yield _sse({
                        "type": "review_suggestions",
                        "suggestions": all_suggestions,
                        "summary": combined_summary,
                    })
                    await _safe_update_doc(doc.id, {"status": "reviewed"})
                else:
                    # 单块审查（短文档）
                    async for sse_event in dify.run_doc_review_stream(
                        content=review_content,
                        user_instruction=review_instruction,
                    ):
                        if sse_event.event == "review_suggestion":
                            yield _sse({
                                "type": "review_suggestion",
                                "suggestion": sse_event.data,
                            })
                        elif sse_event.event == "review_result":
                            _capture_usage(sse_event.data)
                            yield _sse({
                                "type": "review_suggestions",
                                "suggestions": sse_event.data.get("suggestions", []),
                                "summary": sse_event.data.get("summary", ""),
                            })
                            await _safe_update_doc(doc.id, {"status": "reviewed"})
                        elif sse_event.event == "reasoning":
                            yield _sse({"type": "reasoning", "delta": sse_event.data.get("delta", ""), "text": sse_event.data.get("text", ""), "partial": sse_event.data.get("partial", False)})
                        elif sse_event.event == "progress":
                            yield _sse({"type": "status", "message": sse_event.data.get("message", "审查中…")})
                        elif sse_event.event == "error":
                            yield _sse({"type": "error", "message": sse_event.data.get("message", "审查失败")})
                            return

                yield _sse({"type": "done", "full_content": doc.content})
                _record_stage_usage("review")

            elif body.stage == "format":
                # 格式化 — Dify 流式返回结构化段落（支持文件上传到 Dify 文档提取器）
                if not doc.content:
                    yield _sse({"type": "error", "message": "公文内容为空，无法排版"})
                    yield "data: [DONE]\n\n"
                    return

                doc_text = doc.content

                # ── Markdown 预处理：去除 #、*、> 等格式符号，避免被当成正文 ──
                _raw_len = len(doc_text)
                doc_text = _strip_markdown_for_format(doc_text)
                if len(doc_text) != _raw_len:
                    _logger.info(f"Markdown 预处理: {_raw_len} → {len(doc_text)} 字符")

                # 读取源文件字节（如果有），用于上传到 Dify 文档提取器
                # 如果已有结构化段落，不再上传源文件
                format_file_bytes = None
                format_file_name = ""
                has_structured = body.existing_paragraphs and len(body.existing_paragraphs) > 0
                if not has_structured and doc.source_file_path and _is_safe_upload_path(doc.source_file_path):
                    source_path = Path(doc.source_file_path)
                    if source_path.exists():
                        try:
                            format_file_bytes = source_path.read_bytes()
                            format_file_name = source_path.name
                            _logger.info(f"排版阶段读取源文件: {format_file_name} ({len(format_file_bytes)} bytes)")
                        except Exception as e:
                            _logger.warning(f"排版阶段读取源文件失败: {e}")
                        # 同时尝试用 docx 提取纯文本兜底
                        if source_path.suffix.lower() == ".docx":
                            try:
                                from app.api.docformat import _extract_docx_text
                                doc_text = _extract_docx_text(str(source_path))
                            except Exception:
                                pass

                doc_type = "official"
                user_format_instruction = ""
                if body.user_instruction:
                    # 智能识别文档类型
                    instruction_lower = body.user_instruction.strip().lower()
                    if instruction_lower in ("official", "academic", "legal", "proposal", "lab_fund", "school_notice_redhead"):
                        doc_type = instruction_lower
                    else:
                        # 通过关键词推断 doc_type
                        if any(kw in body.user_instruction for kw in ("学术", "论文", "期刊", "毕业论文", "academic")):
                            doc_type = "academic"
                        elif any(kw in body.user_instruction for kw in ("法律", "法规", "判决", "裁定", "起诉", "legal")):
                            doc_type = "legal"
                        elif any(kw in body.user_instruction for kw in ("项目建议书", "建议书", "proposal")):
                            doc_type = "proposal"
                        elif any(kw in body.user_instruction for kw in ("实验室基金", "基金指南", "基金课题", "lab_fund")):
                            doc_type = "lab_fund"
                        elif any(kw in body.user_instruction for kw in ("大学", "学院", "学校", "校名红头", "高校红头", "承办单位", "联系人", "电话")):
                            doc_type = "school_notice_redhead"
                        # 将完整的用户指令传给 Dify
                        user_format_instruction = body.user_instruction

                # ── Phase-1：规则引擎排版（毫秒级，无 LLM 调用） ──
                _rule_paras: list[dict] = []
                _llm_needed_indices: list[int] = []
                _use_rule_engine = True  # 默认启用规则引擎
                _rule_formatted_count = 0

                # 如果用户有明确排版指令（非类型选择），说明需要 LLM 理解意图
                _has_modification_instruction = False
                if user_format_instruction and user_format_instruction.strip():
                    # 仅当指令包含"修改动词"时标记需要LLM
                    _has_modification_instruction = any(
                        kw in user_format_instruction
                        for kw in ("修改", "改成", "调整", "设为", "设置", "换成", "改为",
                                   "去掉", "删掉", "添加", "不要", "不需要", "移除")
                    )
                    # 规则引擎始终运行（用于分类）；是否跳过LLM由后续覆盖率决定

                if _use_rule_engine and has_structured:
                    # 已有结构化段落 → 对每段做规则引擎排版
                    _rule_paras, _llm_needed_indices = _rules_format_paragraphs(
                        [dict(p) for p in body.existing_paragraphs], doc_type,
                    )
                    _rule_formatted_count = sum(1 for p in _rule_paras if p.get("_rule_formatted"))
                    _logger.info(
                        f"规则引擎排版: {_rule_formatted_count}/{len(_rule_paras)} 段高置信度, "
                        f"{len(_llm_needed_indices)} 段需要 LLM"
                    )
                elif _use_rule_engine and not has_structured and doc_text:
                    # 无结构化段落 → 先从纯文本解析段落，再做规则引擎排版
                    _text_paras = _parse_markdown_to_paragraphs(doc_text)
                    if _text_paras:
                        _rule_paras, _llm_needed_indices = _rules_format_paragraphs(_text_paras, doc_type)
                        _rule_formatted_count = sum(1 for p in _rule_paras if p.get("_rule_formatted"))
                        _logger.info(
                            f"规则引擎排版(文本解析): {_rule_formatted_count}/{len(_rule_paras)} 段高置信度, "
                            f"{len(_llm_needed_indices)} 段需要 LLM"
                        )

                # 如果规则引擎覆盖率足够高（>= 60%）且无修改指令，直接输出，低置信度段落也用规则兜底
                _skip_llm = False
                if _rule_paras and not _has_modification_instruction and len(_llm_needed_indices) <= len(_rule_paras) * 0.4:
                    _skip_llm = True
                    _logger.info(f"规则引擎覆盖率足够高 ({_rule_formatted_count}/{len(_rule_paras)}), 跳过 LLM")
                    # 对低置信度段落也用规则引擎兜底填充模板属性
                    for _idx in _llm_needed_indices:
                        _apply_format_template(_rule_paras[_idx], doc_type)
                        _rule_paras[_idx]["_rule_formatted"] = True

                if _skip_llm and _rule_paras:
                    # ── 纯规则引擎排版，直接输出全部段落 ──
                    _low_conf_count = len(_llm_needed_indices)
                    _high_conf_count = len(_rule_paras) - _low_conf_count
                    yield _sse({"type": "format_stats", "rule_count": len(_rule_paras), "llm_count": 0,
                                "high_confidence": _high_conf_count, "low_confidence": _low_conf_count})
                    yield _sse({"type": "status", "message": f"规则引擎排版完成（{len(_rule_paras)} 段，其中 {_low_conf_count} 段低置信度）"})
                    _format_paragraphs: list[str] = []
                    _all_para_data: list[dict] = []
                    _low_conf_set = set(_llm_needed_indices)
                    for _pi, _p in enumerate(_rule_paras):
                        _out = {k: v for k, v in _p.items() if k != "_rule_formatted"}
                        if _pi in _low_conf_set:
                            _out["_confidence"] = "low"
                        _all_para_data.append(_out)
                        yield _sse({"type": "structured_paragraph", "paragraph": _out})
                        if _out.get("text"):
                            _format_paragraphs.append(_out["text"])

                    # 跳过 LLM 排版流，直接跳到保存阶段
                    _use_incremental = has_structured
                    _skip_llm_format = True

                else:
                    _skip_llm_format = False

                    # 如果规则引擎已处理大部分段落，只将低置信度段落发给 LLM
                    if _rule_paras and _llm_needed_indices and not _skip_llm:
                        # 构建 LLM 所需的段落子集（带 _index 标记，便于合并回去）
                        _llm_subset = []
                        for _idx in _llm_needed_indices:
                            _p = dict(_rule_paras[_idx])
                            _p["_index"] = _idx
                            _llm_subset.append(_p)

                        # 将低置信度段落连同锚点上下文一起发给 LLM（紧凑格式，减少 token）
                        if not user_format_instruction:
                            user_format_instruction = ""
                        _llm_needed_set = set(_llm_needed_indices)
                        _llm_prefix = (
                            f"[部分段落排版] 以下文档共 {len(_rule_paras)} 段，"
                            f"标记 ★ 的 {len(_llm_subset)} 个段落需要你确定 style_type 并排版，"
                            f"其余为已确定的锚点（仅供参考上下文，不要输出）。\n"
                            f"请为每个 ★ 段落输出完整的 11 个属性 + _index。\n\n"
                        )
                        # 构建紧凑列表：已确定的段落仅显示 [idx:style]，待分类段落显示完整文本
                        for _pi, _rp in enumerate(_rule_paras):
                            if _pi in _llm_needed_set:
                                _lp_text = _rp.get("text", "")[:200]
                                _llm_prefix += f"★[{_pi}] {_lp_text}\n"
                            else:
                                # 锚点：仅输出 style + 文本摘要（≤30字），大幅减少 token
                                _anchor_text = _rp.get("text", "")[:30]
                                _anchor_style = _rp.get("style_type", "body")
                                _llm_prefix += f"  [{_pi}:{_anchor_style}] {_anchor_text}\n"
                        user_format_instruction = _llm_prefix + "\n" + user_format_instruction
                        # 不发送全文，LLM 只需处理低置信度段落
                        doc_text = ""

                # ── 分块 & 增量模式策略（仅 LLM 排版路径） ──
                if not _skip_llm_format:
                    _chunk_size = _MAX_FORMAT_CHUNK_CHARS
                    # 计算有效文本长度
                    _use_incremental = has_structured  # 默认：有结构化段落→增量
                    if has_structured:
                        _total_para_chars = sum(len(p.get("text", "")) for p in body.existing_paragraphs)
                    else:
                        _total_para_chars = len(doc_text)

                    # ── 长文档策略：自动选择最优排版路径 ──
                    # 阈值：增量模式下，如果段落太多且文本太长，降级为分块排版
                    _INCREMENTAL_THRESHOLD_CHARS = _chunk_size * 2  # 超过 2 倍分块大小
                    _INCREMENTAL_THRESHOLD_PARAS = 100              # 超过 100 段
                    _force_full_reformat = False

                    if _use_incremental and (
                        _total_para_chars > _INCREMENTAL_THRESHOLD_CHARS
                        and len(body.existing_paragraphs) > _INCREMENTAL_THRESHOLD_PARAS
                    ):
                        # 判断是否需要全量重排（大部分段落都是 body = 未格式化 → 全量排版更合适）
                        _body_count = sum(
                            1 for p in body.existing_paragraphs
                            if p.get("style_type", "body") == "body"
                        )
                        _body_ratio = _body_count / max(len(body.existing_paragraphs), 1)

                        if _body_ratio > 0.7:
                            # 70%+ 段落都是 body → 文档基本未格式化，使用全量排版
                            _force_full_reformat = True
                            _use_incremental = False
                            # 从段落中提取纯文本
                            _para_texts = [p.get("text", "").strip() for p in body.existing_paragraphs if p.get("text", "").strip()]
                            doc_text = "\n\n".join(_para_texts)
                            _logger.info(
                                f"长文档降级全量排版: {_total_para_chars} 字符, "
                                f"{len(body.existing_paragraphs)} 段 (body占比 {_body_ratio:.0%})"
                            )
                        else:
                            # 已有格式化信息，使用分块增量排版
                            _logger.info(
                                f"长文档分块增量排版: {_total_para_chars} 字符, "
                                f"{len(body.existing_paragraphs)} 段"
                            )

                    # ── 增量模式：构建紧凑索引列表（仅在不分块增量时使用） ──
                    if _use_incremental and _total_para_chars <= _INCREMENTAL_THRESHOLD_CHARS \
                            and len(body.existing_paragraphs) <= _INCREMENTAL_THRESHOLD_PARAS:
                        # 短文档：单次增量调用
                        _compact_lines = []
                        for _i, _p in enumerate(body.existing_paragraphs):
                            _attrs = [_p.get("style_type", "body")]
                            if _p.get("font_size"): _attrs.append(_p["font_size"])
                            if _p.get("font_family"): _attrs.append(_p["font_family"])
                            if _p.get("bold"): _attrs.append("bold")
                            if _p.get("alignment") and _p["alignment"] != "left": _attrs.append(_p["alignment"])
                            if _p.get("indent"): _attrs.append(f'indent={_p["indent"]}')
                            if _p.get("line_height"): _attrs.append(f'lh={_p["line_height"]}')
                            if _p.get("color") and _p["color"] != "#000000": _attrs.append(f'color={_p["color"]}')
                            if _p.get("red_line") is not None: _attrs.append(f'red_line={"true" if _p["red_line"] else "false"}')
                            _text = _p.get("text", "")
                            _compact_lines.append(f'[{_i}] ({", ".join(_attrs)}) {_text}')
                        _compact_listing = "\n".join(_compact_lines)

                        incremental_prefix = (
                            "[增量修改模式 — 仅输出被修改的段落]\n"
                            f"当前文档共 {len(body.existing_paragraphs)} 个段落，索引与当前属性如下：\n"
                            f"{_compact_listing}\n\n"
                            "规则：\n"
                            "1. 仅输出需要修改的段落，每个必须包含 _index + 完整 11 个属性\n"
                            "2. 未修改的段落不要输出，无修改时输出 {\"paragraphs\": []}\n"
                            "3. 不得擅自修改用户未要求修改的属性\n"
                            "4. red_line 必填：用户说删掉/去掉红线时设 false，说加上红线时设 true\n\n"
                        )
                        if user_format_instruction:
                            user_format_instruction = incremental_prefix + f"【用户修改要求】:\n{user_format_instruction}"
                        else:
                            user_format_instruction = incremental_prefix + "请保持当前排版不变。"

                    # 收集所有结构化段落
                    _format_paragraphs: list[str] = []
                    _all_para_data: list[dict] = []

                    # ── 选择排版流 + 多轮续写支持 ──
                    _is_chunked_path = False
                    _MAX_FMT_CONTINUATION = 50  # 安全上限；实际轮数由截断检测驱动，可适配任意 max_tokens
                    _fmt_conv_id = ""
                    _fmt_full_text = ""
                    _fmt_round_error = False

                    for _fmt_round in range(_MAX_FMT_CONTINUATION):
                        _round_para_start = len(_all_para_data)

                        if _fmt_round == 0:
                            # ── 首轮：按原有 4 路策略选择排版流 ──
                            if _use_incremental and (
                                _total_para_chars > _INCREMENTAL_THRESHOLD_CHARS
                                and len(body.existing_paragraphs) > _INCREMENTAL_THRESHOLD_PARAS
                            ):
                                # ★ 长文档分块增量排版
                                _is_chunked_path = True
                                _logger.info(f"排版路径: 分块增量 ({len(body.existing_paragraphs)} 段, {_total_para_chars} 字符)")
                                _format_stream = _chunked_incremental_format_stream(
                                    dify, body.existing_paragraphs, doc_type,
                                    user_format_instruction,
                                    max_chunk_chars=min(_chunk_size, _INCREMENTAL_MAX_CHARS_PER_CHUNK),
                                    max_paras=_INCREMENTAL_MAX_PARAS_PER_CHUNK,
                                )
                            elif not _use_incremental and len(doc_text) > _chunk_size:
                                # ★ 长文档全量分块排版（含降级后的全量路径）
                                _is_chunked_path = True
                                _logger.info(f"排版路径: 全量分块 ({len(doc_text)} 字符 > {_chunk_size})")
                                _format_stream = _chunked_format_stream(
                                    dify, doc_text, doc_type, user_format_instruction,
                                    max_chunk_chars=_chunk_size,
                                )
                            elif _use_incremental:
                                # ★ 短文档单次增量排版
                                _logger.info(f"排版路径: 单次增量 ({len(body.existing_paragraphs)} 段)")
                                _format_stream = dify.run_doc_format_stream(
                                    "", doc_type, user_format_instruction,
                                )
                            else:
                                # ★ 短文档单次全量排版
                                _logger.info(f"排版路径: 单次全量 ({len(doc_text)} 字符)")
                                _format_stream = dify.run_doc_format_stream(
                                    doc_text, doc_type, user_format_instruction,
                                    file_bytes=format_file_bytes, file_name=format_file_name,
                                )
                        else:
                            # ── 续写轮次：构造上下文连贯的续写指令 ──
                            yield _sse({"type": "status", "message": f"AI 输出被截断，正在续写第 {_fmt_round + 1} 轮…（已获得 {len(_all_para_data)} 段）"})

                            if _use_incremental:
                                # ★ 增量模式续写：基于 _index 精准衔接
                                _received_indices = sorted(set(
                                    p.get("_index") for p in _all_para_data
                                    if p.get("_index") is not None
                                ))
                                if _received_indices:
                                    _max_idx = max(_received_indices)
                                    _continuation_query = (
                                        f"请继续输出。上一轮因长度限制在段落索引 {_max_idx} 处被截断。\n"
                                        f"已收到修改的段落索引：{_received_indices}\n"
                                        f"请继续检查索引 {_max_idx + 1} 到 {len(body.existing_paragraphs) - 1} 的段落，"
                                        f"只输出需要修改的段落，每个必须包含 _index + 完整 11 个属性。\n"
                                        f"未修改的段落不要输出。\n"
                                        f'输出格式：{{"paragraphs":[{{"_index": N, "text":"...","style_type":"...","font_size":"...","font_family":"...","bold":false,"italic":false,"color":"#000000","indent":"","alignment":"left","line_height":"1.5","red_line":false}},…]}}'
                                    )
                                else:
                                    # 回退：LLM 未使用 _index，按段落序号续写
                                    _last_para_texts = [p.get("text", "")[:100] for p in _all_para_data[-3:]] if _all_para_data else []
                                    _continuation_query = (
                                        f"请继续输出。上一轮因长度限制被截断，已输出 {len(_all_para_data)} 个段落。\n"
                                        f"最后几段内容：\n" + "\n".join(f"  第{len(_all_para_data) - len(_last_para_texts) + i + 1}段: {t}" for i, t in enumerate(_last_para_texts)) +
                                        f"\n\n请紧接上文继续输出后续段落，不要重复已输出的内容。"
                                        f"保持完全相同的排版风格和 JSON 格式。\n"
                                        f'{{"paragraphs":[{{"text":"...","style_type":"..."}},…]}}'
                                    )
                            else:
                                # ★ 全量模式续写：提供最后几段的格式上下文以保持风格一致
                                _last_paras_info = _all_para_data[-3:] if _all_para_data else []
                                _context_lines = []
                                for _ci, _cp in enumerate(_last_paras_info):
                                    _cp_text = _cp.get("text", "")[:100]
                                    _cp_style = _cp.get("style_type", "body")
                                    _para_num = len(_all_para_data) - len(_last_paras_info) + _ci + 1
                                    _context_lines.append(f"  第{_para_num}段 [{_cp_style}]: {_cp_text}")
                                _continuation_query = (
                                    f"请继续输出。上一轮因长度限制被截断，已排版 {len(_all_para_data)} 个段落。\n"
                                    f"最后几段及其格式：\n" + "\n".join(_context_lines) +
                                    f"\n\n请紧接上文继续输出后续段落，不要重复已输出的内容。"
                                    f"保持完全相同的排版风格和 JSON 格式。\n"
                                    f'{{"paragraphs":[{{"text":"...","style_type":"..."}},…]}}'
                                )

                            _format_stream = dify.run_doc_format_stream(
                                "", doc_type, _continuation_query,
                                conversation_id=_fmt_conv_id,
                            )

                        async for sse_event in _format_stream:
                            if sse_event.event == "structured_paragraph":
                                para_data = {
                                    "text": sse_event.data.get("text", ""),
                                    "style_type": sse_event.data.get("style_type", "body"),
                                }
                                # 透传富格式属性（含 color, red_line, _index）
                                for key in ("font_size", "font_family", "bold", "italic", "color", "indent", "alignment", "line_height", "red_line", "_index"):
                                    if key in sse_event.data and sse_event.data[key] is not None:
                                        para_data[key] = sse_event.data[key]

                                # ── 后处理：清除残留 Markdown 符号 + 补全模板默认值 ──
                                para_data["text"] = _strip_markdown_inline(para_data["text"])
                                _apply_format_template(para_data, doc_type)

                                _all_para_data.append(para_data)

                                # 实时推送每个解析到的段落（全量 + 增量均推送）
                                yield _sse({"type": "structured_paragraph", "paragraph": para_data})
                                await asyncio.sleep(0)  # 让出事件循环，强制 ASGI 刷新 SSE 缓冲
                                if not _use_incremental and para_data["text"]:
                                    _format_paragraphs.append(para_data["text"])
                            elif sse_event.event == "progress":
                                yield _sse({"type": "status", "message": sse_event.data.get("message", "排版中…")})
                            elif sse_event.event == "reasoning":
                                yield _sse({"type": "reasoning", "delta": sse_event.data.get("delta", ""), "text": sse_event.data.get("text", ""), "partial": sse_event.data.get("partial", False)})
                            elif sse_event.event == "format_progress":
                                yield _sse({"type": "format_progress", **sse_event.data})
                            elif sse_event.event == "text_chunk":
                                yield _sse({"type": "text", "text": sse_event.data.get("text", "")})
                            elif sse_event.event == "message_end":
                                _fmt_full_text = sse_event.data.get("full_text", "")
                                _fmt_conv_id = sse_event.data.get("conversation_id", "") or _fmt_conv_id
                                _capture_usage(sse_event.data)
                            elif sse_event.event == "error":
                                yield _sse({"type": "error", "message": sse_event.data.get("message", "排版失败")})
                                _fmt_round_error = True
                                break

                        if _fmt_round_error:
                            break

                        # ── 续写判定：增量 / 全量均支持多轮续写 ──
                        _new_para_count = len(_all_para_data) - _round_para_start
                        if not _is_chunked_path and _new_para_count > 0 and _fmt_conv_id:
                            if _check_json_truncated(_fmt_full_text):
                                _mode_label = "增量" if _use_incremental else "全量"
                                _logger.info(
                                    f"{_mode_label}排版第 {_fmt_round + 1} 轮输出被截断 "
                                    f"(paras={len(_all_para_data)}, text={len(_fmt_full_text)} chars)，"
                                    f"将发起第 {_fmt_round + 2} 轮续写"
                                )
                                continue  # → 下一轮续写
                        break  # 输出完整或分块路径，退出循环

                    if _fmt_round_error:
                        yield "data: [DONE]\n\n"
                        return

                    if _fmt_round > 0:
                        _logger.info(f"排版多轮续写完成：共 {_fmt_round + 1} 轮，最终 {len(_all_para_data)} 段")
                        yield _sse({"type": "status", "message": f"排版续写完成（共 {_fmt_round + 1} 轮，{len(_all_para_data)} 段）"})

                    # ── 后处理：红线删除关键词检测 ──
                    _user_instr_lower = (body.user_instruction or "").lower()
                    _want_remove_redline = any(
                        kw in _user_instr_lower
                        for kw in ("删掉红线", "去掉红线", "移除红线", "删除红线",
                                   "删掉横线", "去掉横线", "移除横线", "删除横线",
                                   "删掉红色横线", "去掉红色横线", "移除红色横线", "删除红色横线",
                                   "删掉红色分隔线", "去掉红色分隔线", "删掉分隔线", "去掉分隔线",
                                   "不要红线", "不需要红线", "不要横线", "不需要横线")
                    )

                    # ── 增量模式：基于 _index 合并或回退全量 diff ──
                    if _use_incremental:
                        # 通知前端清空预览段落，准备接收最终合并结果
                        yield _sse({"type": "format_clear"})
                        _has_index = any(p.get("_index") is not None for p in _all_para_data) if _all_para_data else False
                        if _has_index:
                            # ★ 快速路径：AI 仅返回了被修改的段落（带 _index）
                            _modified_map: dict[int, dict] = {}
                            _skipped_indices: list[int] = []  # #22: 记录越界被跳过的索引
                            for p in _all_para_data:
                                idx = p.pop("_index", None)
                                if idx is not None and isinstance(idx, int) and 0 <= idx < len(body.existing_paragraphs):
                                    _modified_map[idx] = p
                                elif idx is not None:
                                    _skipped_indices.append(idx)
                                    _logger.warning(f"增量排版：AI 返回越界 _index={idx}（段落总数={len(body.existing_paragraphs)}），已跳过")
                            if _skipped_indices:
                                yield _sse({"type": "status", "message": f"⚠ AI 返回了 {len(_skipped_indices)} 个无效段落索引（{_skipped_indices[:5]}），已自动跳过"})
                            _format_paragraphs = []
                            for i, old_p in enumerate(body.existing_paragraphs):
                                if i in _modified_map:
                                    new_p = _modified_map[i]
                                    if _want_remove_redline and new_p.get("style_type") == "title":
                                        new_p["red_line"] = False
                                    # ── 仅当文本或关键属性有实际变化时才标记 modified ──
                                    _old_text = old_p.get("text", "").strip()
                                    _new_text = new_p.get("text", "").strip()
                                    _attrs_changed = any(
                                        new_p.get(k) != old_p.get(k)
                                        for k in ("style_type", "font_size", "font_family", "bold",
                                                  "italic", "color", "alignment", "indent",
                                                  "line_height", "red_line")
                                        if new_p.get(k) is not None
                                    )
                                    if _old_text != _new_text or _attrs_changed:
                                        new_p["_change"] = "modified"
                                        new_p["_original_text"] = old_p.get("text", "")
                                    yield _sse({"type": "structured_paragraph", "paragraph": new_p})
                                    _format_paragraphs.append(new_p.get("text", ""))
                                else:
                                    out_p = {k: v for k, v in old_p.items() if k not in ("_change", "_original_text", "_change_reason")}
                                    if _want_remove_redline and out_p.get("style_type") == "title":
                                        out_p["red_line"] = False
                                    yield _sse({"type": "structured_paragraph", "paragraph": out_p})
                                    _format_paragraphs.append(out_p.get("text", ""))
                            _logger.info(f"增量排版完成（索引模式）: 修改 {len(_modified_map)} / {len(body.existing_paragraphs)} 段")
                        elif _all_para_data:
                            # 回退路径：AI 输出了全部段落（无 _index），使用传统 diff
                            if _want_remove_redline:
                                for pd in _all_para_data:
                                    if pd.get("style_type") == "title":
                                        pd["red_line"] = False
                            diffed = _compute_para_diff(body.existing_paragraphs, _all_para_data)
                            _format_paragraphs = []
                            for dp in diffed:
                                yield _sse({"type": "structured_paragraph", "paragraph": dp})
                                if dp.get("text"):
                                    _format_paragraphs.append(dp["text"])
                            _logger.info(f"增量排版完成（回退全量diff）: 共 {len(_all_para_data)} 段")
                        else:
                            # AI 无输出，保持原样
                            _format_paragraphs = [p.get("text", "") for p in body.existing_paragraphs if p.get("text")]
                            for old_p in body.existing_paragraphs:
                                out_p = {k: v for k, v in old_p.items() if k not in ("_change", "_original_text", "_change_reason")}
                                yield _sse({"type": "structured_paragraph", "paragraph": out_p})
                    else:
                        # 非增量模式：红线兜底
                        if _want_remove_redline and _all_para_data:
                            for pd in _all_para_data:
                                if pd.get("style_type") == "title":
                                    pd["red_line"] = False

                # #19: 混合可视化 — 发送规则引擎 vs LLM 统计信息
                if not _skip_llm_format and _rule_paras:
                    _llm_count = len(_llm_needed_indices)
                    _rule_only_count = len(_rule_paras) - _llm_count
                    yield _sse({"type": "format_stats", "rule_count": _rule_only_count, "llm_count": _llm_count,
                                "high_confidence": _rule_only_count, "low_confidence": 0})

                # 保存格式化前版本快照 + 更新文档（独立事务，断连安全）
                _final_content = "\n\n".join(_format_paragraphs) if _format_paragraphs else None
                _updates = {"status": "formatted"}
                if _final_content is not None:
                    _updates["content"] = _final_content
                try:
                    _saved_content = await asyncio.wait_for(
                        _safe_update_doc(
                            doc.id, _updates,
                            save_version_before=True,
                            version_user_id=current_user.id,
                            version_change_type="format",
                            version_change_summary="格式化前版本",
                        ),
                        timeout=30.0,
                    )
                except (asyncio.TimeoutError, Exception) as _save_err:
                    _logger.warning(f"排版保存失败（不影响前端显示）: {_save_err}")
                    _saved_content = _final_content or ""
                yield _sse({"type": "done", "full_content": _saved_content})

                _record_stage_usage("format")

            elif body.stage == "format_suggest":
                # 排版建议 — 分析文档给出详细排版格式建议
                if not doc.content:
                    yield _sse({"type": "error", "message": "公文内容为空，无法分析"})
                    yield "data: [DONE]\n\n"
                    return

                suggest_content = doc.content
                has_structured = body.existing_paragraphs and len(body.existing_paragraphs) > 0
                if has_structured:
                    _text_lines = []
                    for _p in body.existing_paragraphs:
                        _st = _p.get("style_type", "body")
                        _text = _p.get("text", "").strip()
                        if _text:
                            _text_lines.append(f"[{_st}] {_text}")
                    suggest_content = "\n".join(_text_lines)

                _logger.info(f"排版建议：内容长度 {len(suggest_content)} 字符")

                async for sse_event in dify.run_format_suggest_stream(
                    content=suggest_content,
                    user_instruction=body.user_instruction or "",
                ):
                    if sse_event.event == "format_suggestion":
                        yield _sse({"type": "format_suggestion", "suggestion": sse_event.data})
                    elif sse_event.event == "format_suggest_result":
                        _capture_usage(sse_event.data)
                        yield _sse({"type": "format_suggest_result", **sse_event.data})
                    elif sse_event.event == "reasoning":
                        yield _sse({"type": "reasoning", "delta": sse_event.data.get("delta", ""), "text": sse_event.data.get("text", ""), "partial": sse_event.data.get("partial", False)})
                    elif sse_event.event == "progress":
                        yield _sse({"type": "status", "message": sse_event.data.get("message", "分析中…")})
                    elif sse_event.event == "error":
                        yield _sse({"type": "error", "message": sse_event.data.get("message", "排版建议失败")})
                        return

                # 如果有结构化段落，额外运行规则引擎生成格式化预览段落（允许前端一键应用）
                if has_structured and body.existing_paragraphs:
                    _suggest_paras, _suggest_llm_indices = _rules_format_paragraphs(
                        [dict(p) for p in body.existing_paragraphs],
                        doc.doc_type or "official",
                    )
                    # 对低置信度段落也用模板兜底
                    for _idx in _suggest_llm_indices:
                        _apply_format_template(_suggest_paras[_idx], doc.doc_type or "official")
                    # 比较并标记变更
                    _changed_paras = []
                    for _i, (_new_p, _old_p) in enumerate(zip(_suggest_paras, body.existing_paragraphs)):
                        _out_p = {k: v for k, v in _new_p.items() if k not in ("_rule_formatted",)}
                        _old_dict = dict(_old_p)
                        # 检测是否有实质变更（排版属性差异）
                        _has_diff = False
                        for _attr in ("font_size", "font_family", "bold", "alignment", "indent",
                                      "line_height", "color", "letter_spacing", "style_type"):
                            if _out_p.get(_attr) != _old_dict.get(_attr) and _out_p.get(_attr) is not None:
                                _has_diff = True
                                break
                        if _has_diff:
                            _out_p["_change"] = "modified"
                            _out_p["_change_reason"] = "规则引擎排版建议"
                        _changed_paras.append(_out_p)
                    _change_count = sum(1 for p in _changed_paras if p.get("_change"))
                    if _change_count > 0:
                        yield _sse({
                            "type": "format_suggest_paragraphs",
                            "paragraphs": _changed_paras,
                            "change_count": _change_count,
                        })
                        _logger.info(f"排版建议附带格式化段落预览: {_change_count}/{len(_changed_paras)} 段有变更")

                yield _sse({"type": "done", "full_content": doc.content})
                _record_stage_usage("format_suggest")

            yield "data: [DONE]\n\n"

        except asyncio.CancelledError:
            # 客户端断开连接 — 释放锁，尝试保存已获取的部分排版结果
            _logger.warning(f"AI处理被取消（客户端断开）[{body.stage}] doc={doc_id}")
            if body.stage == "format" and _all_para_data:
                try:
                    _partial_texts = [p.get("text", "") for p in _all_para_data if p.get("text")]
                    if _partial_texts:
                        _partial_content = "\n\n".join(_partial_texts)
                        await _safe_update_doc(
                            doc.id,
                            {"content": _partial_content, "status": "formatted"},
                            save_version_before=True,
                            version_user_id=current_user.id,
                            version_change_type="format",
                            version_change_summary="格式化中断（客户端断开）- 部分结果",
                        )
                        _logger.info(f"客户端断开：已保存 {len(_partial_texts)} 段部分排版结果")
                except Exception as _save_ex:
                    _logger.warning(f"客户端断开：保存部分排版结果失败: {_save_ex}")
        except Exception as e:
            _logger.exception(f"AI对话处理异常 [{body.stage}]")
            yield _sse({"type": "error", "message": f"AI处理异常: {str(e)}"})
            yield "data: [DONE]\n\n"
        finally:
            # 无论成功/失败/断开，都释放并发锁
            try:
                await r.delete(lock_key)
            except Exception:
                _logger.warning(f"释放 AI 处理锁失败: {lock_key}")
            # 主动关闭 DB session，防止连接池泄漏（CancelledError 时依赖注入清理可能不执行）
            try:
                await db.close()
            except Exception:
                pass

    try:
        return StreamingResponse(
            event_generator(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no",
            },
        )
    except Exception:
        # 如果 event_generator 定义/启动前出错，确保释放锁
        try:
            await r.delete(lock_key)
        except Exception:
            pass
        raise


_STAGE_NAMES = {
    "draft": "公文起草",
    "review": "审查优化",
    "format": "格式规范",
}


@router.delete("/{doc_id}/ai-lock")
async def release_ai_lock(
    doc_id: UUID,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """手动释放 AI 处理锁（当锁卡住时使用）"""
    _logger = logging.getLogger(__name__)
    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return error(ErrorCode.NOT_FOUND, "公文不存在")
    if doc.creator_id != current_user.id:
        return error(ErrorCode.PERMISSION_DENIED, "只有创建者才能释放锁")

    r = await get_redis()
    lock_key = f"doc_ai_lock:{doc_id}"
    deleted = await r.delete(lock_key)
    if deleted:
        _logger.info(f"用户 {current_user.display_name} 手动释放了 AI 锁: {lock_key}")
    return success(data={"released": bool(deleted)})


@router.get("/{doc_id}/versions")
async def list_document_versions(
    doc_id: UUID,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """公文版本历史"""
    # 验证公文存在
    doc_result = await db.execute(select(Document).where(Document.id == doc_id))
    if not doc_result.scalar_one_or_none():
        return error(ErrorCode.NOT_FOUND, "公文不存在")

    result = await db.execute(
        select(DocumentVersion)
        .where(DocumentVersion.document_id == doc_id)
        .order_by(DocumentVersion.version_number.desc())
    )
    versions = result.scalars().all()

    # 批量查创建者姓名
    user_ids = {v.created_by for v in versions}
    user_map = {}
    if user_ids:
        ur = await db.execute(select(User.id, User.display_name).where(User.id.in_(user_ids)))
        user_map = {row[0]: row[1] for row in ur.all()}

    items = [
        {
            **DocumentVersionItem.model_validate(v).model_dump(mode="json"),
            "created_by_name": user_map.get(v.created_by, ""),
            "has_format": bool(v.formatted_paragraphs),
        }
        for v in versions
    ]

    return success(data=items)


@router.get("/{doc_id}/versions/{version_id}")
async def get_document_version(
    doc_id: UUID,
    version_id: UUID,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """获取指定版本详情"""
    result = await db.execute(
        select(DocumentVersion).where(
            DocumentVersion.id == version_id,
            DocumentVersion.document_id == doc_id,
        )
    )
    version = result.scalar_one_or_none()
    if not version:
        return error(ErrorCode.NOT_FOUND, "版本不存在")

    # 查创建者姓名
    ur = await db.execute(select(User.display_name).where(User.id == version.created_by))
    created_by_name = ur.scalar_one_or_none() or ""

    data = {
        **DocumentVersionDetail.model_validate(version).model_dump(mode="json"),
        "created_by_name": created_by_name,
        "has_format": bool(version.formatted_paragraphs),
    }
    return success(data=data)


@router.post("/{doc_id}/versions/{version_id}/restore")
async def restore_document_version(
    doc_id: UUID,
    version_id: UUID,
    current_user: User = Depends(require_permission("app:doc:write")),
    db: AsyncSession = Depends(get_db),
):
    """恢复到指定版本（回退），先保存当前内容为新版本快照。
    使用 Redis 锁防止与自动保存/AI 处理并发冲突。"""
    import logging
    logger = logging.getLogger("govai.restore")

    doc_result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = doc_result.scalar_one_or_none()
    if not doc:
        return error(ErrorCode.NOT_FOUND, "公文不存在")
    if doc.creator_id != current_user.id:
        return error(ErrorCode.PERMISSION_DENIED, "只有创建者才能恢复版本")

    ver_result = await db.execute(
        select(DocumentVersion).where(
            DocumentVersion.id == version_id,
            DocumentVersion.document_id == doc_id,
        )
    )
    version = ver_result.scalar_one_or_none()
    if not version:
        return error(ErrorCode.NOT_FOUND, "版本不存在")

    # 先把目标值保存到局部变量（避免 ORM 对象状态干扰）
    restore_content = version.content or ""
    restore_version_number = version.version_number
    restore_formatted = version.formatted_paragraphs

    # ── Redis 锁：防止与自动保存/AI 处理并发覆盖 ──
    redis = None
    lock_key = f"govai:doc_lock:{doc_id}"
    lock_acquired = False
    try:
        redis = await get_redis()
    except Exception:
        pass  # Redis 不可用时降级为无锁

    if redis:
        lock_acquired = await redis.set(lock_key, "restore", ex=10, nx=True)
        if not lock_acquired:
            return error(ErrorCode.CONFLICT, "文档正在被修改，请稍后再试")

    try:
        # 先把当前内容保存为快照
        if doc.content:
            await _save_version(db, doc, current_user.id, change_type="restore", change_summary="回退前备份")

        # 恢复内容 + 结构化排版段落
        doc.content = restore_content
        doc.formatted_paragraphs = restore_formatted
        await db.flush()

        # 保存恢复后的版本记录
        await _save_version(db, doc, current_user.id, change_type="restore", change_summary=f"恢复到版本 v{restore_version_number}")

        logger.info(f"版本恢复成功: doc={doc_id}, version={version_id}, v{restore_version_number}")
    except Exception as e:
        logger.error(f"版本恢复失败: doc={doc_id}, version={version_id}, error={e}", exc_info=True)
        raise
    finally:
        if redis and lock_acquired:
            await redis.delete(lock_key)

    # 不显式 commit — 由 get_db 依赖统一提交事务
    return success(data={
        "content": restore_content,
        "version_number": restore_version_number,
        "formatted_paragraphs": restore_formatted,
    })


# ── 辅助函数 ──


async def _save_version(
    db: AsyncSession,
    doc: Document,
    user_id: UUID,
    change_type: str | None = None,
    change_summary: str | None = None,
):
    """保存公文版本快照（带重试，防止并发版本号冲突）"""
    for attempt in range(3):
        # 获取最新版本号
        result = await db.execute(
            select(func.max(DocumentVersion.version_number))
            .where(DocumentVersion.document_id == doc.id)
        )
        max_ver = result.scalar() or 0

        version = DocumentVersion(
            document_id=doc.id,
            version_number=max_ver + 1,
            content=doc.content or "",
            formatted_paragraphs=doc.formatted_paragraphs,
            change_type=change_type,
            change_summary=change_summary,
            created_by=user_id,
        )
        try:
            async with db.begin_nested():
                db.add(version)
                await db.flush()
            return  # 成功
        except SAIntegrityError:
            logger.warning(f"版本号冲突 (attempt {attempt+1}/3): doc={doc.id}, tried v{max_ver+1}")
            try:
                db.expunge(version)
            except Exception:
                pass
            continue
    # 3 次均失败 → 记录错误但不中断业务流程
    logger.error(f"保存版本失败: doc={doc.id}, 3次重试均因版本号冲突失败")
